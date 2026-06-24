"""FundTrail — cybercrime fund-trail tracking for the TN Cyber Crime Wing.

Single-module Flask application. Sections, in order:

     1. Logging & request tracing
     2. Application setup, middleware & security headers
     3. IFSC bank/state lookup & case access
     4. Auth, session & validation helpers
     5. Schema bootstrap & legacy data migration
     6. Routes: authentication & session
     7. Routes: dashboards & Excel upload
     8. Routes: fund-trail graph & transaction views
     9. Routes: KYC, hold/refund & letter generation
    10. Routes: case & officer administration
    11. Error handlers
    12. Routes: reports, analytics & exports
    13. Routes: search, case workflow & health
    14. Entrypoint

Run locally with `python app.py`, or in production with
`gunicorn --preload "app:app"` (see Dockerfile).
"""

import secrets
import string
import time
from functools import wraps

from dotenv import load_dotenv
from flask import (
    Flask,
    abort,
    flash,
    g,
    jsonify,
    make_response,
    redirect,
    render_template,
    request,
    send_file,
    send_from_directory,
    session,
    url_for,
)
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from werkzeug.utils import secure_filename

load_dotenv()
import contextlib
import io
import json
import os
import zipfile
from collections import defaultdict
from urllib.parse import urlparse

if os.name == "nt":
    import msvcrt
else:
    import fcntl
from datetime import datetime, timedelta, timezone

import pandas as pd  # pyright: ignore[reportMissingImports]
import pymysql  # pyright: ignore[reportMissingModuleSource]
from email_utils import send_integration_test_email, send_password_reset
from flask_login import (  # pyright: ignore[reportMissingImports]
    LoginManager,
    current_user,
    login_required,
    login_user,
    logout_user,
)
from flask_migrate import Migrate  # pyright: ignore[reportMissingModuleSource]
from flask_wtf.csrf import CSRFError, CSRFProtect
from itsdangerous import BadSignature, SignatureExpired, URLSafeTimedSerializer
from models import (
    CASE_STATUSES,
    CaseNote,
    Complaint,
    KYCDetails,
    MRMStatusLog,
    MRMTracking,
    POHRefundDetails,
    Transaction,
    UploadedFile,
    UsageLog,
    User,
    db,
)
from sqlalchemy import desc, func, inspect, or_, text
from sqlalchemy.orm import defer
from werkzeug.security import check_password_hash, generate_password_hash

# Only patch pymysql as MySQLdb when actually using MySQL. The default path is
# SQLite, which needs no patch. (DATABASE_URL is loaded by load_dotenv() above.)
if os.environ.get("DATABASE_URL", "").startswith("mysql"):
    pymysql.install_as_MySQLdb()
import concurrent.futures
import re

# TOTP two-factor auth. Optional — the app runs fine without pyotp,
# 2FA setup is simply unavailable until it is installed.
try:
    import pyotp
except ImportError:  # pragma: no cover
    pyotp = None

try:
    import qrcode
except ImportError:  # pragma: no cover
    qrcode = None
import logging
import sys
from decimal import Decimal, InvalidOperation
from logging.handlers import RotatingFileHandler

import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape

# ReportLab imports for PDF generation
from reportlab.pdfgen import canvas
from xhtml2pdf import pisa

# ---------------------------------------------------------------------------
# Shared string constants (deduplicated literals — SonarCloud python:S1192)
# ---------------------------------------------------------------------------
INITIAL_CREDENTIALS_FILENAME = "INITIAL_CREDENTIALS.txt"
STATE_ANDHRA_PRADESH = "Andhra Pradesh"
STATE_TAMIL_NADU = "Tamil Nadu"
STATE_ANDAMAN_NICOBAR = "Andaman and Nicobar Islands"
ROLE_INVESTIGATIVE_OFFICER = "Investigative Officer"
STATUS_PARTIALLY_REFUNDED = "Partially Refunded"
DATETIME_DISPLAY_FORMAT = "%d %b %Y, %H:%M"
DDL_VARCHAR_100 = "VARCHAR(100)"
DDL_VARCHAR_50 = "VARCHAR(50)"
DDL_VARCHAR_20 = "VARCHAR(20)"
DDL_INT_NULL = "INT NULL"
ROUTE_LOGIN = "/login"
ROUTE_INDEX = "/index"
TEMPLATE_FORGOT_PASSWORD = "forgot_password.html"
TXN_WITHDRAWAL_ATM = "Withdrawal through ATM"
MSG_INTERNAL_SERVER_ERROR = "Internal server error"
MSG_DOC_GENERATION_ERROR = "Internal error while generating the document"
MIME_PDF = "application/pdf"
BANK_UNKNOWN = "Unknown Bank"
LABEL_SUSPECT_ACCOUNT_NUMBER = "Suspect Account Number"
MIME_XLSX = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

# "Transaction ID / UTR Number2" column-name handling (python:S1192).
# Detection stage — exact column-name spellings find_utr2_column() matches against.
UTR_COLUMN_VARIANTS = [
    "Transaction ID / UTR Number2",
    "Transaction Id / UTR Number2",
    "Transaction ID/ UTR Number2",
    "Transaction ID/UTR Number2",
    "Transaction ID / UTR Number 2",
    "Transaction Id / UTR Number 2",
    "Txn ID / UTR Number2",
    "Txn Id / UTR Number2",
    "Txn ID / UTR Number 2",
    "Txn Id / UTR Number 2",
    "UTR Number2",
    "Txn ID/UTR Number2",
    "Transaction ID/UTR Number 2",
    "Transaction ID / UTR",
    "Transaction ID/UTR",
    "Txn ID/UTR",
    "Transaction ID / UTR Number",
]
# Extraction stage — the detection list PLUS three broader UTR-only fallbacks that
# are deliberately used only when pulling a value from a row, never for column
# detection. Built by extending UTR_COLUMN_VARIANTS so the subset relationship
# stays visible in code (investigation: A/D = B/C + these 3).
UTR_EXTRACTION_FALLBACKS = ["Txn ID / UTR Number", "UTR Number", "UTR"]
UTR_TXN_ID_VARIANTS = UTR_COLUMN_VARIANTS + UTR_EXTRACTION_FALLBACKS
# Column-name normalization regex + fuzzy "number 2" substring, shared by every
# norm()/has_number2 site in the UTR matching code.
COLUMN_NAME_NORM_RE = r"[\s/_\-\.]+"
UTR_FUZZY_NUMBER2 = "number 2"

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# log file location
if os.environ.get("FUNDTRAIL_DATA_DIR"):
    log_dir = os.path.join(os.environ["FUNDTRAIL_DATA_DIR"], "logs")
elif os.name == "nt":
    log_dir = os.path.join(os.environ.get("APPDATA"), "FundTrailTool")
else:
    log_dir = os.path.join(os.path.expanduser("~"), ".fundtrailtool")

try:
    os.makedirs(log_dir, exist_ok=True)
except OSError:
    log_dir = None
log_file_path = os.path.join(log_dir, "app.log") if log_dir else None

try:
    handler = RotatingFileHandler(log_file_path, maxBytes=10000000, backupCount=5)
except OSError:
    handler = logging.StreamHandler()
handler.setLevel(logging.INFO)


# ---------------------------------------------------------------------------
# 1. Logging & request tracing
# ---------------------------------------------------------------------------


class RequestIdFilter(logging.Filter):
    """Stamp every log record with the current request's ID so one
    request's entries can be traced end-to-end. '-' outside a request context."""

    def filter(self, record):
        try:
            from flask import g as _g
            from flask import has_request_context

            record.request_id = getattr(_g, "request_id", "-") if has_request_context() else "-"
        except Exception:
            record.request_id = "-"
        return True


class JsonLogFormatter(logging.Formatter):
    """Machine-parseable one-JSON-object-per-line log format.
    Enabled with LOG_FORMAT=json in .env; default stays human-readable."""

    def format(self, record):
        payload = {
            "ts": self.formatTime(record),
            "level": record.levelname,
            "logger": record.name,
            "request_id": getattr(record, "request_id", "-"),
            "message": record.getMessage(),
        }
        if record.exc_info:
            payload["exc"] = self.formatException(record.exc_info)
        return json.dumps(payload)


if os.environ.get("LOG_FORMAT", "").lower() == "json":
    formatter = JsonLogFormatter()
else:
    formatter = logging.Formatter("%(asctime)s - %(name)s - %(levelname)s - [%(request_id)s] %(message)s")
handler.addFilter(RequestIdFilter())
handler.setFormatter(formatter)
logger.addHandler(handler)

# Dedicated alert log for 500s — a small, high-signal file an admin
# (or a watcher script / mail hook) can monitor instead of grepping app.log.
try:
    alert_handler = RotatingFileHandler(os.path.join(log_dir, "alerts.log"), maxBytes=1000000, backupCount=2)
except (TypeError, OSError):
    alert_handler = logging.StreamHandler()
alert_handler.setFormatter(logging.Formatter("%(asctime)s - %(message)s"))
alert_logger = logging.getLogger("fundtrail.alerts")
alert_logger.addHandler(alert_handler)
alert_logger.setLevel(logging.ERROR)


# ---------------------------------------------------------------------------
# 2. Application setup, middleware & security headers
# ---------------------------------------------------------------------------


def resource_path(relative_path):
    if hasattr(sys, "_MEIPASS"):
        base_path = sys._MEIPASS  # PyInstaller
    else:
        base_path = os.path.dirname(os.path.abspath(__file__))  # normal run
    return os.path.join(base_path, relative_path)


def _env_bool(name, default=False):
    value = os.environ.get(name)
    if value is None:
        return default
    return value.strip().lower() in {"1", "true", "yes", "on"}


def _env_float(name, default):
    value = os.environ.get(name)
    if value is None or value == "":
        return default
    try:
        return float(value)
    except ValueError:
        logger.warning("Invalid %s=%r; using %s", name, value, default)
        return default


# ---------------------------------------------------------------------------
# Sentry PII scrubbing (financial-crime case data must not leave the process).
# Field/param/variable-name substrings whose VALUES may carry case PII. Derived
# from the actual column names in models.py (Transaction / KYCDetails / POH):
# from_account, to_account, account_number, ack_no, ifsc_code, cheque_ifsc,
# txn_id, put_on_hold_txn_id, kyc_name/aadhar/mobile/address, aadhar, mobile,
# address, name. Matched case-insensitively as substrings, so request params,
# headers, "extra", contexts, breadcrumbs and exception-frame local variables
# are all covered.
# ---------------------------------------------------------------------------
SENTRY_SENSITIVE_KEY_PARTS = (
    "account", "ack_no", "aadhar", "aadhaar", "kyc", "mobile",
    "phone", "ifsc", "txn_id", "pan_number", "pancard", "address",
)
# Name-like keys deliberately KEPT (not scrubbed) for debugging value — they are
# not case PII even though they contain "name".
SENTRY_KEEP_KEYS = frozenset({
    "username", "filename", "file_name", "bank_name", "hostname",
    "classname", "funcname", "modulename", "pathname", "step_label",
})
_SENTRY_FILTERED = "[Filtered]"


def _sentry_key_is_sensitive(key):
    k = str(key).lower()
    if k in SENTRY_KEEP_KEYS:
        return False
    if k == "name":  # KYCDetails.name / kyc_name — suspect/victim names
        return True
    return any(part in k for part in SENTRY_SENSITIVE_KEY_PARTS)


def _sentry_scrub(obj, _depth=0):
    """Recursively replace values under sensitive keys with [Filtered]. Walks
    dicts/lists in place so exception-frame vars, request data, extra, contexts
    and breadcrumbs are all covered. Depth-bounded to avoid pathological cycles."""
    if _depth > 12:
        return obj
    if isinstance(obj, dict):
        for key, value in list(obj.items()):
            if _sentry_key_is_sensitive(key):
                obj[key] = _SENTRY_FILTERED
            else:
                obj[key] = _sentry_scrub(value, _depth + 1)
    elif isinstance(obj, list):
        return [_sentry_scrub(item, _depth + 1) for item in obj]
    return obj


def _sentry_before_send(event, _hint):
    """Strip case PII before any event leaves the process. Returning the event
    (not None) keeps error reporting working; only sensitive fields are redacted."""
    # URL query strings can carry ack_no / account identifiers — drop them whole.
    request = event.get("request")
    if isinstance(request, dict):
        request.pop("query_string", None)
        url = request.get("url")
        if isinstance(url, str) and "?" in url:
            request["url"] = url.split("?", 1)[0]
    return _sentry_scrub(event)


try:
    import sentry_sdk
    from sentry_sdk.integrations.flask import FlaskIntegration

    if os.environ.get("SENTRY_DSN"):
        sentry_sdk.init(
            dsn=os.environ["SENTRY_DSN"],
            integrations=[FlaskIntegration()],
            environment=os.environ.get("SENTRY_ENVIRONMENT", "development"),
            release=os.environ.get("SENTRY_RELEASE", globals().get("APP_VERSION")),
            send_default_pii=_env_bool("SENTRY_SEND_DEFAULT_PII", False),
            # Off by default: log lines can contain case data; opt in explicitly.
            enable_logs=_env_bool("SENTRY_ENABLE_LOGS", False),
            traces_sample_rate=_env_float("SENTRY_TRACES_SAMPLE_RATE", 0.1),
            profile_session_sample_rate=_env_float("SENTRY_PROFILE_SESSION_SAMPLE_RATE", 0.0),
            profile_lifecycle=os.environ.get("SENTRY_PROFILE_LIFECYCLE", "trace"),
            before_send=_sentry_before_send,
        )
except ImportError:  # pragma: no cover
    pass
except TypeError as exc:  # pragma: no cover - older SDK compatibility
    logger.warning("Sentry SDK option not supported by installed version: %s", exc)


app = Flask(__name__, template_folder=resource_path("templates"), static_folder=resource_path("static"))

try:
    from werkzeug.serving import WSGIRequestHandler

    WSGIRequestHandler.server_version = ""
    WSGIRequestHandler.sys_version = ""
except Exception:
    pass


def _default_data_dir():
    """A stable, per-user data directory independent of where the repo/EXE lives, so
    the database (and therefore the login) survives re-clones, folder moves and EXE
    rebuilds. Overridden by FUNDTRAIL_DATA_DIR (Docker sets it to /data)."""
    if os.name == "nt":
        base = os.environ.get("LOCALAPPDATA") or os.path.expanduser("~")
    else:
        base = os.environ.get("XDG_DATA_HOME") or os.path.join(os.path.expanduser("~"), ".local", "share")
    return os.path.join(base, "FundTrail")


secure_base = os.environ.get("FUNDTRAIL_DATA_DIR") or _default_data_dir()
os.makedirs(secure_base, exist_ok=True)

# One-time migration: copy an existing legacy repo-local data/ into the stable dir so
# pre-existing installs keep their database and login (only when not overridden).
if not os.environ.get("FUNDTRAIL_DATA_DIR"):
    _legacy_base = os.path.abspath(os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "data"))
    if os.path.isdir(_legacy_base) and _legacy_base != secure_base:
        import shutil

        for _f in (
            "fundtrail.db", INITIAL_CREDENTIALS_FILENAME,
            "poh_refund_details.db", "kyc_details.db", "ifsc_state_cache.json",
        ):
            _src, _dst = os.path.join(_legacy_base, _f), os.path.join(secure_base, _f)
            if os.path.exists(_src) and not os.path.exists(_dst):
                try:
                    shutil.copy2(_src, _dst)
                except OSError:
                    pass


class StripServerHeaderMiddleware:
    def __init__(self, app):
        self.app = app

    def __call__(self, environ, start_response):
        def custom_start_response(status, headers, exc_info=None):
            filtered = [
                (k, v)
                for (k, v) in headers
                if k.lower() not in ("server", "x-powered-by", "x-aspnet-version", "x-aspnetmvc-version")
            ]
            return start_response(status, filtered, exc_info)

        return self.app(environ, custom_start_response)


app.wsgi_app = StripServerHeaderMiddleware(app.wsgi_app)

# Initialize rate limiter
limiter = Limiter(
    app=app,
    key_func=get_remote_address,
    # Generous global cap for an interactive internal tool (graphs/analytics fire many
    # AJAX calls). 50/hour used to make normal use randomly hit "Too Many Requests".
    default_limits=["5000 per day", "1000 per hour"],
    storage_uri="memory://",
)


@app.before_request
def generate_csp_nonce():
    """Per-request setup: CSP nonce + request ID."""
    g.csp_nonce = secrets.token_hex(16)
    g.request_id = secrets.token_hex(8)


@app.context_processor
def inject_csp_nonce():
    """Inject CSP nonce into templates"""
    return {"csp_nonce": g.csp_nonce}


@app.after_request
def set_security_headers(response):
    """Add security headers to all responses"""
    response.headers["X-Content-Type-Options"] = "nosniff"
    response.headers["X-Frame-Options"] = "DENY"
    # '0' disables the legacy XSS auditor (deprecated; it introduced its own XSS
    # vectors). CSP below is the real control.
    response.headers["X-XSS-Protection"] = "0"
    # Only send HSTS when running over HTTPS (LAN plain-HTTP deployments skip this).
    if not app.config.get("SESSION_COOKIE_INSECURE", False):
        response.headers["Strict-Transport-Security"] = "max-age=31536000; includeSubDomains"

    # Use the nonce generated in before_request
    nonce = getattr(g, "csp_nonce", "")

    csp_policy = (
        "default-src 'self'; "
        f"script-src 'self' 'nonce-{nonce}'; "
        "style-src 'self' 'unsafe-inline'; "
        "img-src 'self' data:; "
        "font-src 'self'; "
        "connect-src 'self'; "
        "frame-ancestors 'none'; "
        "base-uri 'self'; "
        "form-action 'self';"
    )

    response.headers["Content-Security-Policy"] = csp_policy
    response.headers["Referrer-Policy"] = "strict-origin-when-cross-origin"
    response.headers["Permissions-Policy"] = "geolocation=(), microphone=(), camera=()"

    # Remove Server header to prevent version disclosure
    response.headers.pop("Server", None)

    if request.path.startswith("/static"):
        # Static assets are immutable between deploys — let browsers cache them
        # for a day instead of re-downloading the ~2MB vendored JS every page.
        response.headers.setdefault("Cache-Control", "public, max-age=86400")
    else:
        # Prevent caching of sensitive pages
        if "Cache-Control" not in response.headers:
            response.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, private, max-age=0"
        response.headers["Pragma"] = "no-cache"
        response.headers["Expires"] = "0"

    return response


csrf = CSRFProtect(app)

# Transparently gzip HTML/JSON/JS responses. Optional dependency —
# the app runs uncompressed if flask_compress is missing.
try:
    from flask_compress import Compress

    Compress(app)
except ImportError:  # pragma: no cover
    pass


migrate = Migrate(app, db, directory=resource_path("migrations"))

# Cache for IFSC to state mappings to avoid repeated API calls
ifsc_cache = {}

# IFSC district code to state mapping
DISTRICT_TO_STATE = {
    "01": "Jammu and Kashmir",
    "02": "Himachal Pradesh",
    "03": "Punjab",
    "04": "Chandigarh",
    "05": "Uttarakhand",
    "06": "Haryana",
    "07": "Delhi",
    "08": "Rajasthan",
    "09": "Uttar Pradesh",
    "10": "Bihar",
    "11": "Sikkim",
    "12": "Arunachal Pradesh",
    "13": "Nagaland",
    "14": "Manipur",
    "15": "Mizoram",
    "16": "Tripura",
    "17": "Meghalaya",
    "18": "Assam",
    "19": "West Bengal",
    "20": "Jharkhand",
    "21": "Odisha",
    "22": "Chhattisgarh",
    "23": "Madhya Pradesh",
    "24": "Gujarat",
    "25": "Daman and Diu",
    "26": "Dadra and Nagar Haveli",
    "27": "Maharashtra",
    "28": STATE_ANDHRA_PRADESH,
    "29": "Karnataka",
    "30": "Goa",
    "31": "Lakshadweep",
    "32": "Kerala",
    "33": STATE_TAMIL_NADU,
    "34": "Puducherry",
    "35": STATE_ANDAMAN_NICOBAR,
    "36": "Telangana",
    "37": STATE_ANDHRA_PRADESH,
    "38": "Ladakh",
}


# ---------------------------------------------------------------------------
# 3. IFSC bank/state lookup & case access
# ---------------------------------------------------------------------------


def get_state(ifsc):
    if not ifsc or len(ifsc) < 6:
        return "Unknown"
    if ifsc in ifsc_cache:
        return ifsc_cache[ifsc]
    district_code = ifsc[4:6]
    state = DISTRICT_TO_STATE.get(district_code)
    if state:
        ifsc_cache[ifsc] = state
        return state
    # Fallback to local IFSC_CODES.xlsx lookup
    try:
        from ifsc_utils import get_state as excel_get_state

        state = excel_get_state(ifsc)
        if state and state != "Unknown":
            ifsc_cache[ifsc] = state
            return state
    except Exception as e:
        logger.exception(f"Error fetching state for IFSC from excel {ifsc}: {e}")
    ifsc_cache[ifsc] = "Unknown"
    return "Unknown"


def _safe_log(value):
    """Strip CR/LF from user-controlled values before logging them, to prevent
    log forging / injection (SonarCloud pythonsecurity:S5145)."""
    return str(value).replace("\r", " ").replace("\n", " ")


def check_case_access(ack_no):
    """
    Verifies if the current user has access to the case identified by ack_no.
    Aborts with 403 if unauthorized.
    """
    if not current_user.is_authenticated:
        abort(403)

    ack_no = str(ack_no).strip()
    complaint = Complaint.query.filter_by(ack_no=ack_no).first()

    # SuperAdmin can access every case.
    if is_superadmin():
        return

    # Regular Admin: can access cases in their own group, plus legacy unowned cases.
    if is_admin():
        if complaint is None or complaint.owner_admin_id in (None, current_user.id):
            return
        logger.warning(
            "Unauthorized cross-group access attempt to %s by admin %s",
            _safe_log(ack_no), current_user.username,
        )
        abort(403)

    # Investigative Officer: only cases they uploaded or were explicitly assigned.
    # Every upload now creates a Complaint row (upload_excel + backfill_complaints),
    # so a missing or unmatched row always means "not your case".
    if complaint and (complaint.uploaded_by == current_user.id or complaint.assigned_to == current_user.id):
        return

    logger.warning(
        "Unauthorized access attempt to %s by officer %s", _safe_log(ack_no), current_user.username
    )
    abort(403)


# Password complexity is enforced in one place: models.validate_password (called by
# User.set_password). A duplicate validate_password() previously lived here but was never
# called — removed during de-duplication so there's a single source of truth for the policy.

ALLOWED_IFSC_DOMAIN = "ifsc.razorpay.com"
MAX_WORKERS = 10

# Login brute-force policy (Fix: magic numbers -> named constants).
MAX_LOGIN_ATTEMPTS = 5
LOCKOUT_MINUTES = 15
# Force a password change after N days (0 disables expiry). Users with
# no recorded password age (pre-feature accounts) are not retroactively forced.
PASSWORD_MAX_AGE_DAYS = int(os.environ.get("PASSWORD_MAX_AGE_DAYS", "90"))
IFSC_CACHE_FILE = os.path.join(secure_base, "ifsc_state_cache.json")
ifsc_api_cache = {}

# 2FA replay prevention: tracks (user_id, code) pairs that have already been accepted
# within their 30-second window. Each entry expires after 90 s (3 windows) then is pruned.
import threading as _threading

_USED_TOTP: dict = {}
_USED_TOTP_LOCK = _threading.Lock()
MAX_2FA_ATTEMPTS = 3          # failed TOTP tries before aborting the pending session
PENDING_2FA_TTL = 300         # seconds a pending-2FA session stays valid (5 min)


def _totp_is_fresh(user_id: int, code: str) -> bool:
    """Return True and mark used if this (user_id, code) hasn't been seen yet.
    Returns False if the code is a replay — prevents the same OTP being accepted twice."""
    now = time.time()
    key = (user_id, code)
    with _USED_TOTP_LOCK:
        expired = [k for k, exp in _USED_TOTP.items() if now > exp]
        for k in expired:
            del _USED_TOTP[k]
        if key in _USED_TOTP:
            return False
        _USED_TOTP[key] = now + 90
        return True


def load_ifsc_cache():
    global ifsc_api_cache
    if os.path.exists(IFSC_CACHE_FILE):
        try:
            with open(IFSC_CACHE_FILE, "r") as f:
                ifsc_api_cache = json.load(f)
            logger.info(f"Loaded {len(ifsc_api_cache)} entries from IFSC cache.")
        except Exception as e:
            logger.exception(f"Error loading IFSC cache: {e}")
            ifsc_api_cache = {}
    else:
        ifsc_api_cache = {}


def save_ifsc_cache():
    try:
        with open(IFSC_CACHE_FILE, "w") as f:
            json.dump(ifsc_api_cache, f)
        logger.info(f"Saved {len(ifsc_api_cache)} entries to IFSC cache.")
    except Exception as e:
        logger.exception(f"Error saving IFSC cache: {e}")


# Initialize cache
load_ifsc_cache()


def get_state_from_api(ifsc_code):
    if not ifsc_code or ifsc_code == "N/A":
        return "Unknown"

    # Validate IFSC format (11 alphanumeric characters)
    if not re.match(r"^[A-Z]{4}0[A-Z0-9]{6}$", ifsc_code):
        return "Invalid IFSC"

    if ifsc_code in ifsc_api_cache:
        return ifsc_api_cache[ifsc_code]

    # Try local Excel/Pickle first
    try:
        from ifsc_utils import get_state as get_state_local

        local_state = get_state_local(ifsc_code)
        if local_state and local_state != "Unknown":
            ifsc_api_cache[ifsc_code] = local_state
            return local_state
    except Exception as e:
        logger.exception(f"Error fetching state from local utils for {ifsc_code}: {e}")

    try:
        # Use only trusted API endpoint
        url = f"https://{ALLOWED_IFSC_DOMAIN}/{ifsc_code}"
        response = requests.get(url, timeout=5, allow_redirects=False)

        if response.status_code == 200:
            data = response.json()
            state = data.get("STATE", "Unknown")
            ifsc_api_cache[ifsc_code] = state
            return state
        else:
            return "Unknown"
    except Exception as e:
        logger.exception(f"Error fetching state for {ifsc_code}: {e}")
        return "Unknown"


@app.route("/ifsc_info/<ifsc>", methods=["GET"])
@login_required
def ifsc_info(ifsc):
    try:
        from ifsc_utils import get_ifsc_info

        info = get_ifsc_info(ifsc) or {}
        return jsonify(info)
    except Exception as e:
        logger.exception(f"Error returning IFSC info for {ifsc}: {e}")
        return jsonify({}), 500


login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = "login"  # redirect to this route if not logged in


# Tell Flask-Login how to load a user
# ---------------------------------------------------------------------------
# 4. Auth, session & validation helpers
# ---------------------------------------------------------------------------


@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))


def admin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not current_user.is_authenticated or current_user.role != "Admin":
            abort(403)
        return f(*args, **kwargs)

    return decorated_function


def is_admin():
    """Authoritative admin check — uses current_user (DB-backed), not the possibly-stale
    session['role'] copy. Safe to call for anonymous users (returns False, no error)."""
    return current_user.is_authenticated and getattr(current_user, "role", None) == "Admin"


def is_superadmin():
    """True when the current user is both an Admin AND carries the SuperAdmin flag.
    SuperAdmins bypass all group-isolation filters and can see every admin's data."""
    return is_admin() and getattr(current_user, "is_superadmin", False)


def _cases_q():
    """Return a scoped Complaint base-query for the current user.

    - SuperAdmin  → all cases across every admin group
    - Admin       → only cases belonging to their admin group (owner_admin_id == self.id)
                    plus legacy cases with no owner set (owner_admin_id IS NULL)
    - Officer     → only cases they uploaded or were assigned to them
    """
    if is_superadmin():
        return Complaint.query
    if is_admin():
        return Complaint.query.filter(
            or_(Complaint.owner_admin_id == current_user.id, Complaint.owner_admin_id.is_(None))
        )
    return Complaint.query.filter(
        or_(Complaint.uploaded_by == current_user.id, Complaint.assigned_to == current_user.id)
    )


def _officers_q():
    """Return a scoped User base-query for Investigative Officers visible to the current admin.

    - SuperAdmin  → all officers across every admin group
    - Admin       → only officers whose admin_id points to them
    """
    base = User.query.filter_by(role=ROLE_INVESTIGATIVE_OFFICER)
    if is_superadmin():
        return base
    return base.filter_by(admin_id=current_user.id)


# Control characters that must never be stored in text fields (keep tab/newline/return).
_CONTROL_RE = re.compile("[" + re.escape("".join(chr(c) for c in range(32) if c not in (9, 10, 13)) + chr(127)) + "]")


def sanitize_text(value, max_len=255):
    """Defence-in-depth for user-supplied text that gets stored/displayed: strip null
    bytes & control characters, trim whitespace, cap length. (Output is still
    Jinja-escaped; this hardens what we persist and pass around.)"""
    if value is None:
        return ""
    return _CONTROL_RE.sub("", str(value)).strip()[:max_len]


def _is_posix():
    return os.name != "nt"


def _ensure_secure_dir(path):
    os.makedirs(path, exist_ok=True)
    if _is_posix():
        try:
            os.chmod(path, 0o700)
        except Exception:
            pass


def _ensure_secure_file(path):
    if not os.path.exists(path):
        with open(path, "a"):
            # touch — create the empty file if it does not yet exist
            pass
    if _is_posix():
        try:
            os.chmod(path, 0o600)
        except Exception:
            pass


# The SECRET_KEY MUST come from the environment (no weak/hardcoded fallback).
# There is deliberately NO hardcoded fallback — the app refuses to start without
# a real secret. This restores the behaviour verified in the Feb-2026 security
# retest. Do NOT re-add a fallback string here (that was the regression we fixed).
SECRET_KEY = os.environ.get("SECRET_KEY")
if not SECRET_KEY:
    raise RuntimeError(
        "SECRET_KEY environment variable not set! Generate one with: "
        'python3 -c "import secrets; print(secrets.token_hex(32))" '
        "and add it to your .env file."
    )
app.config["SECRET_KEY"] = SECRET_KEY
_ensure_secure_dir(secure_base)
primary_db_path = os.path.join(secure_base, "fundtrail.db")
# Legacy per-feature SQLite files. POH/KYC data now lives in the main DB; these paths
# are kept ONLY so migrate_split_dbs_into_main() can import any pre-existing rows out
# of them. We no longer create them (don't resurrect empty split DBs).
poh_db_path = os.path.join(secure_base, "poh_refund_details.db")
kyc_db_path = os.path.join(secure_base, "kyc_details.db")
_ensure_secure_file(primary_db_path)
db_url = os.environ.get("DATABASE_URL")
if not db_url:
    db_url = f"sqlite:///{primary_db_path}"
app.config["SQLALCHEMY_DATABASE_URI"] = db_url
app.config["SQLALCHEMY_ENGINE_OPTIONS"] = {"pool_pre_ping": True, "pool_recycle": 1800}

# Concurrency: enable SQLite WAL + a per-connection busy timeout so several officers
# can read/write at once without "database is locked". (No-op for a MySQL DATABASE_URL.)
if db_url.startswith("sqlite"):
    from sqlalchemy import event
    from sqlalchemy.engine import Engine

    @event.listens_for(Engine, "connect")
    def _set_sqlite_pragma(dbapi_connection, _record):
        try:
            cur = dbapi_connection.cursor()
            cur.execute("PRAGMA journal_mode=WAL")
            cur.execute("PRAGMA busy_timeout=5000")
            cur.close()
        except Exception:
            pass
# Single database: POHRefundDetails and KYCDetails are plain tables in the main DB now
# (no SQLALCHEMY_BINDS, so no cross-DB syncing and one consistent source of truth).

# Secure session cookie configuration.
# The Secure flag MUST be True in production (HTTPS). But over plain http
# (local dev / offline http://127.0.0.1) browsers never send a Secure cookie, which
# silently breaks login. So we default to True (secure) and allow an explicit,
# documented opt-out for local development only.
#   -> set SESSION_COOKIE_INSECURE=true in .env ONLY on a local dev machine.
_cookie_insecure = os.environ.get("SESSION_COOKIE_INSECURE", "false").lower() == "true"
app.config["SESSION_COOKIE_SECURE"] = not _cookie_insecure  # True unless explicitly opted out
if _cookie_insecure:
    logger.warning(
        "SESSION_COOKIE_INSECURE=true -> session cookies are NOT marked Secure. "
        "This is for LOCAL HTTP DEV ONLY; never run a real deployment this way (use HTTPS)."
    )

app.config["SESSION_COOKIE_HTTPONLY"] = True  # No JavaScript access
app.config["SESSION_COOKIE_SAMESITE"] = "Lax"  # CSRF protection
# Removed __Host- prefix as it requires Secure=True (HTTPS)
app.config["SESSION_COOKIE_NAME"] = "session"
app.config["PERMANENT_SESSION_LIFETIME"] = timedelta(minutes=30)
app.config["SESSION_REFRESH_EACH_REQUEST"] = False
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
app.config["UPLOAD_FOLDER"] = os.path.join(secure_base, "uploads")
os.makedirs(app.config["UPLOAD_FOLDER"], exist_ok=True)

# Hard ceiling on any incoming request body — stops an oversized/malicious POST from
# exhausting memory before per-file checks run. (Excel uploads are capped at 10MB
# separately; this bounds the whole request.)
app.config["MAX_CONTENT_LENGTH"] = 25 * 1024 * 1024  # 25 MB

db.init_app(app)

PASSWORD_RESET_MAX_AGE = 1800


def _password_reset_serializer():
    return URLSafeTimedSerializer(app.config["SECRET_KEY"], salt="fundtrail-password-reset")


def _password_reset_version(user):
    if not user or not user.password_changed_at:
        return ""
    changed = user.password_changed_at.replace(tzinfo=timezone.utc) if user.password_changed_at.tzinfo is None else user.password_changed_at
    return changed.isoformat()


def _make_password_reset_token(user):
    return _password_reset_serializer().dumps({"uid": user.id, "pw": _password_reset_version(user)})


def _load_password_reset_user(token):
    data = _password_reset_serializer().loads(token, max_age=PASSWORD_RESET_MAX_AGE)
    user = User.query.get(data.get("uid"))
    if not user or data.get("pw") != _password_reset_version(user):
        raise BadSignature("stale password reset token")
    return user


def _send_password_reset_if_email_matches(user, submitted_email):
    submitted = (submitted_email or "").strip().lower()
    stored = (user.email or "").strip().lower() if user else ""
    if user and stored and submitted and submitted == stored:
        token = _make_password_reset_token(user)
        link = url_for("reset_password", token=token, _external=True)
        try:
            send_password_reset(user.email, link)
            app.logger.info("Password reset email requested for user id %s", user.id)
        except Exception as exc:
            app.logger.warning("Password reset email send failed for user id %s: %s", user.id, exc)


def ensure_usage_log_table():
    inspector = inspect(db.engine)
    if "usage_log" not in inspector.get_table_names():
        db.create_all()


def log_usage(action, filename=None, ack_no=None):
    try:
        # Prefer the authenticated identity (DB-backed) over the session copy, which
        # is blank for any route that never sets session['username'/'role'].
        uname = session.get("username")
        urole = session.get("role")
        try:
            if current_user.is_authenticated:
                uname = current_user.username
                urole = current_user.role
        except Exception:
            pass
        entry = UsageLog(username=uname, role=urole, action=action, filename=filename, ack_no=ack_no)
        db.session.add(entry)
        db.session.commit()
    except Exception as e:
        try:
            db.session.rollback()
        except Exception:
            pass
        logger.exception(f"UsageLog error: {e}")


def validate_account_number(account):
    """Validate account number format"""
    if not account or account == "N/A":
        return True
    return bool(re.match(r"^\d{9,18}$", str(account)))


def validate_amount(amount):
    """Validate transaction amount"""
    try:
        amt = Decimal(str(amount))
        return amt > 0 and amt < Decimal("999999999.99")
    except (InvalidOperation, ValueError):
        return False


def validate_aadhar(value):
    """Aadhaar = 12 digits (spaces/dashes ignored). Blank allowed (optional field)."""
    if not value:
        return True
    return bool(re.fullmatch(r"\d{12}", re.sub(r"[\s-]", "", str(value))))


def validate_mobile(value):
    """Indian mobile = 10 digits starting 6-9, optional +91/0 prefix. Blank allowed."""
    if not value:
        return True
    return bool(re.fullmatch(r"(?:\+91|0)?[6-9]\d{9}", re.sub(r"[\s-]", "", str(value))))


def validate_court_order_date(value):
    """Blank, an HTML date value (yyyy-mm-dd), or dd-mm-yyyy / dd/mm/yyyy."""
    if not value:
        return True
    s = str(value).strip()
    return bool(re.fullmatch(r"\d{4}-\d{2}-\d{2}", s) or re.fullmatch(r"\d{2}[-/]\d{2}[-/]\d{4}", s))


# Mirrors the fixed <select> options in static/graph.js (blank = clear the field).
ALLOWED_REFUND_STATUSES = {"", "Refunded", STATUS_PARTIALLY_REFUNDED, "Not Refunded"}

# Money Restoration Module (MRM) — 7-stage sequential workflow. Order is significant:
# a stage can only be dated once the previous one is dated. Stage 6 (1-based),
# "Amount Refunded to Victim", carries refund_type (FULL/PARTIAL) + amount.
# Labels surface in the graph "Status of MRM" control and the audit timeline.
MRM_STEPS = [
    "MRM Bank Entry for Hold",
    "MRM Requested by Victim",
    "Verification of Suspect Account Holder by IO",
    "Approval for Refund by DC/SP",
    "Refund Requested to Bank",
    "Amount Refunded to Victim",
    "Refund Completion Intimated to Court",
]
MRM_REFUND_STEP = 6  # 1-based index of the "Amount Refunded to Victim" stage
ALLOWED_REFUND_TYPES = {"FULL", "PARTIAL"}


IST = timezone(timedelta(hours=5, minutes=30))


def to_ist(dt):
    """Convert a stored timestamp to IST. Stored datetimes are UTC (utc_now); SQLite
    hands them back naive, so assume UTC when tzinfo is missing. Returns None on None."""
    if dt is None:
        return None
    if dt.tzinfo is None:
        dt = dt.replace(tzinfo=timezone.utc)
    return dt.astimezone(IST)


@app.template_filter("ist")
def _ist_filter(dt, fmt=DATETIME_DISPLAY_FORMAT):
    """Jinja filter: render a UTC datetime in IST, e.g. {{ log.timestamp | ist }}."""
    d = to_ist(dt)
    return d.strftime(fmt) if d else ""


def ordinal(n):
    try:
        n = int(n)
    except Exception:
        return str(n)
    if 10 <= n % 100 <= 20:
        suf = "th"
    else:
        suf = {1: "st", 2: "nd", 3: "rd"}.get(n % 10, "th")
    return f"{n}{suf}"


# ---------------------------------------------------------------------------
# 5. Schema bootstrap & legacy data migration
# ---------------------------------------------------------------------------


def _ensure_columns(table_name, required_columns):
    """Add any missing columns to an EXISTING table, on both MySQL and SQLite.

    db.create_all() creates new *tables* but never alters existing ones, so when a
    new column is introduced an already-created DB (esp. the SQLite dev DB) silently
    lacks it -> runtime 'no such column' errors. This backfills the columns. Column
    names come only from our own hard-coded dicts (never user input)."""
    try:
        driver = db.engine.url.drivername or ""
    except Exception:
        driver = ""
    try:
        inspector = inspect(db.engine)
        if not inspector.has_table(table_name):
            return
        existing = {c["name"] for c in inspector.get_columns(table_name)}
        missing = {c: d for c, d in required_columns.items() if c not in existing}
        if not missing:
            return
        # Quote identifiers: backticks for MySQL, double-quotes elsewhere (SQLite).
        q = "`" if driver.startswith("mysql") else '"'
        with db.engine.begin() as conn:
            for col, ddl in missing.items():
                try:
                    conn.execute(text(f"ALTER TABLE {q}{table_name}{q} ADD COLUMN {q}{col}{q} {ddl}"))
                    logger.info(f"Added missing column {table_name}.{col}")
                except Exception as exc:
                    logger.warning(f"Skipping column {table_name}.{col}: {exc}")
    except Exception as exc:
        logger.warning(f"_ensure_columns({table_name}) failed: {exc}")


def ensure_transaction_columns():
    """Add any newly introduced columns to the transaction table if missing."""
    required_columns = {
        "layer": "INT",
        "from_account": DDL_VARCHAR_100,
        "to_account": DDL_VARCHAR_100,
        "ack_no": DDL_VARCHAR_100,
        "bank_name": DDL_VARCHAR_100,
        "ifsc_code": DDL_VARCHAR_50,
        "txn_date": DDL_VARCHAR_100,
        "txn_id": DDL_VARCHAR_100,
        "amount": "FLOAT",
        "disputed_amount": "FLOAT",
        "action_taken": "VARCHAR(255)",
        "account_number": DDL_VARCHAR_50,
        "state": DDL_VARCHAR_50,
        "atm_id": DDL_VARCHAR_100,
        "atm_withdraw_amount": "FLOAT",
        "atm_withdraw_date": DDL_VARCHAR_100,
        "atm_location": "VARCHAR(200)",
        "cheque_no": DDL_VARCHAR_100,
        "cheque_withdraw_amount": "FLOAT",
        "cheque_withdraw_date": DDL_VARCHAR_100,
        "cheque_ifsc": DDL_VARCHAR_50,
        "put_on_hold_txn_id": DDL_VARCHAR_100,
        "put_on_hold_date": DDL_VARCHAR_100,
        "put_on_hold_amount": "FLOAT",
        "court_order_date": DDL_VARCHAR_20,
        "refund_status": DDL_VARCHAR_50,
        "refund_amount": "FLOAT",
        "refund_type": "VARCHAR(10)",
        "kyc_name": "VARCHAR(120)",
        "kyc_aadhar": DDL_VARCHAR_20,
        "kyc_mobile": DDL_VARCHAR_20,
        "kyc_address": "VARCHAR(200)",
        "upload_id": "INT",
    }

    _ensure_columns("transaction", required_columns)


def ensure_user_columns():
    """Add newly introduced columns to the user table if they are missing."""
    required_columns = {
        "name": DDL_VARCHAR_100,
        "rank": DDL_VARCHAR_100,
        "email": "VARCHAR(120)",
        "manual_upload_count": DDL_INT_NULL,
        "must_change_password": "BOOLEAN DEFAULT 0",
        "last_login_at": "DATETIME NULL",
        "password_changed_at": "DATETIME NULL",
        "totp_secret": "VARCHAR(64) NULL",
        # Multi-admin isolation (added in v2.1)
        "admin_id": DDL_INT_NULL,
        "is_superadmin": "BOOLEAN DEFAULT 0",
    }
    _ensure_columns("user", required_columns)


def ensure_complaint_columns():
    """Add newly introduced columns to the complaint table if missing."""
    _ensure_columns("complaint", {
        "status": "VARCHAR(30) DEFAULT 'Open'",
        # Multi-admin isolation (added in v2.1): which admin group owns this case.
        "owner_admin_id": DDL_INT_NULL,
    })


def _mrm_row_from_hold_txn(t, poh, refund_date_fn):
    """Build an MRMTracking row seeded from one hold transaction, or None if no
    stage could be dated. Mirrors the legacy refund/court -> MRM step mapping."""
    ref_status = (poh.refund_status if poh else None) or t.refund_status
    ref_amount = poh.refund_amount if (poh and poh.refund_amount is not None) else t.refund_amount
    row = MRMTracking(ack_no=t.ack_no, txn_id=t.put_on_hold_txn_id)
    if t.put_on_hold_date:
        row.step1 = str(t.put_on_hold_date)
    if ref_status == "Refunded":
        setattr(row, f"step{MRM_REFUND_STEP}", refund_date_fn(poh))
        row.refund_type = "FULL"
        row.refund_amount = ref_amount
        t.refund_type = "FULL"
    elif ref_status == STATUS_PARTIALLY_REFUNDED:
        setattr(row, f"step{MRM_REFUND_STEP}", refund_date_fn(poh))
        row.refund_type = "PARTIAL"
        row.refund_amount = ref_amount
        t.refund_type = "PARTIAL"
    if any(getattr(row, f"step{i}") for i in range(1, len(MRM_STEPS) + 1)):
        return row
    return None


def ensure_mrm_backfill():
    """Seed MRMTracking from pre-existing refund/court data. Idempotent: only
    creates a row for a (ack_no, txn_id) that has none yet, so officer edits made
    through /save_mrm_status are never overwritten on a later startup.

    Mapping (per the migration spec):
      put_on_hold_date                 -> step1 (MRM Bank Entry for Hold)
      refund_status Refunded           -> step6 dated, refund_type FULL    + amount
      refund_status Partially Refunded -> step6 dated, refund_type PARTIAL + amount
    """
    try:
        inspector = inspect(db.engine)
        if "mrm_tracking" not in inspector.get_table_names():
            db.create_all()
    except Exception as exc:
        logger.warning(f"ensure_mrm_backfill: table check failed: {exc}")
        return

    def _refund_date(poh):
        if poh and poh.updated_at:
            return poh.updated_at.strftime("%Y-%m-%d")
        return datetime.now().strftime("%Y-%m-%d")

    try:
        existing = {
            (m.ack_no, m.txn_id)
            for m in MRMTracking.query.with_entities(MRMTracking.ack_no, MRMTracking.txn_id).all()
        }
        poh_map = {(p.ack_no, p.txn_id): p for p in POHRefundDetails.query.all()}
        hold_txns = Transaction.query.filter(Transaction.put_on_hold_txn_id.isnot(None)).all()
        created = 0
        for t in hold_txns:
            key = (t.ack_no, t.put_on_hold_txn_id)
            if None in key or key in existing:
                continue
            row = _mrm_row_from_hold_txn(t, poh_map.get(key), _refund_date)
            if row is not None:
                db.session.add(row)
                existing.add(key)
                created += 1
        if created:
            db.session.commit()
            logger.info(f"MRM backfill: created {created} tracking row(s) from existing refund/court data")
    except Exception as exc:
        try:
            db.session.rollback()
        except Exception:
            pass
        logger.warning(f"ensure_mrm_backfill failed: {exc}")


def generate_secure_password(length=16):
    """Generate a secure random first-login password.

    Keep the symbol set copy-safe for users reading it from a dialog or console.
    The password policy accepts more symbols, but quotes/brackets are easy to
    mistype and were causing first-login confusion in the packaged EXE.
    """
    safe_symbols = "!@#$%&*?"
    alphabet = string.ascii_letters + string.digits + safe_symbols
    while True:
        password = "".join(secrets.choice(alphabet) for _ in range(length))
        if (
            any(c.islower() for c in password)
            and any(c.isupper() for c in password)
            and any(c.isdigit() for c in password)
            and any(c in safe_symbols for c in password)
        ):
            return password


def migrate_split_dbs_into_main():
    """One-time, idempotent consolidation: import POH/KYC rows from the old per-feature
    SQLite files (poh_refund_details.db / kyc_details.db) into the main DB now that those
    models live there. Each old file is renamed to *.migrated after a successful import
    so it is never re-read. No-op on a fresh install or once already migrated."""
    import sqlite3

    jobs = [
        (poh_db_path, "poh_refund_details", POHRefundDetails, ("ack_no", "txn_id")),
        (kyc_db_path, "kyc_details", KYCDetails, ("txn_id",)),
    ]
    for old_path, table, model, key_cols in jobs:
        if not old_path or not os.path.exists(old_path):
            continue
        valid_cols = {c.name for c in model.__table__.columns}
        try:
            conn = sqlite3.connect(old_path)
            conn.row_factory = sqlite3.Row
            try:
                # `table` is selected from a fixed allowlist above, not user input.
                rows = conn.execute(f"SELECT * FROM {table}").fetchall()  # nosec B608
            except sqlite3.OperationalError:
                rows = []  # table missing in the old file
            finally:
                conn.close()

            imported = 0
            for r in rows:
                d = {k: r[k] for k in r.keys() if k in valid_cols and k != "id"}
                key = {k: d.get(k) for k in key_cols}
                if model.query.filter_by(**key).first():
                    continue  # idempotent: already present in the main DB
                db.session.add(model(**d))
                imported += 1
            db.session.commit()
            os.rename(old_path, old_path + ".migrated")
            if imported:
                logger.info(f"Consolidated {imported} row(s) from {os.path.basename(old_path)} into main DB.")
        except Exception as exc:
            db.session.rollback()
            logger.exception(f"DB consolidation for {old_path} failed: {exc}")


def backfill_complaints():
    """One-time/idempotent: ensure every uploaded case has a Complaint row tied to
    its uploader, so per-officer isolation also works for data uploaded before this
    feature existed. Maps UploadedFile.uploader (username) -> User.id; if the
    uploader is unknown the case is left admin-only (uploaded_by/assigned_to=None)."""
    try:
        existing = {row[0] for row in db.session.query(Complaint.ack_no).all()}
        umap = {u.username: u.id for u in User.query.with_entities(User.id, User.username).all()}
        rows = (
            db.session.query(Transaction.ack_no, UploadedFile.uploader, UploadedFile.filename, UploadedFile.upload_time)
            .join(UploadedFile, Transaction.upload_id == UploadedFile.id)
            .filter(Transaction.ack_no.isnot(None))
            .distinct()
            .all()
        )
        created, seen = 0, set()
        for ack, uploader, fname, utime in rows:
            ack_s = str(ack).strip()
            if not ack_s or ack_s in existing or ack_s in seen:
                continue
            seen.add(ack_s)
            uid = umap.get(uploader)
            db.session.add(
                Complaint(
                    ack_no=ack_s,
                    file_name=fname,
                    uploaded_by=uid,
                    assigned_to=uid,
                    upload_time=utime,
                )
            )
            created += 1
        if created:
            db.session.commit()
            logger.info(f"Backfilled {created} complaint row(s) for pre-existing uploads.")
    except Exception as e:
        db.session.rollback()
        logger.exception(f"backfill_complaints failed: {e}")


with app.app_context():
    ensure_transaction_columns()
    ensure_user_columns()
    ensure_complaint_columns()
    ensure_usage_log_table()
    ensure_mrm_backfill()


VIEW_ONLY_ROLES = set()  # Viewer (read-only) role removed — only Admin & Investigative Officer


# ---------------------------------------------------------------------------
# 6. Routes: authentication & session
# ---------------------------------------------------------------------------


@app.route("/", methods=["GET"])
def home():
    return redirect(ROUTE_LOGIN)


@app.route("/home", methods=["GET"])
def role_home():
    """Single 'Home' entry point. Always sends a user to THEIR OWN dashboard, so an
    admin never lands on the officer dashboard (and vice-versa). Generic 'Return to
    Home' links point here instead of hard-coding one dashboard."""
    if "username" not in session:
        return redirect(url_for("login"))
    if session.get("role") == "Admin":
        return redirect(url_for("admin_dashboard"))
    return redirect(url_for("index"))


ALLOWED_HOSTS = {"localhost", "127.0.0.1"}


def is_safe_url(target):
    if not target:
        return False
    # Only allow relative URLs starting with /
    if target.startswith("/") and not target.startswith("//"):
        return True
    parsed = urlparse(target)
    return parsed.netloc in ALLOWED_HOSTS and parsed.scheme in ("http", "https")


@app.route("/change_password", methods=["GET", "POST"])
@login_required
def change_password():
    # Accessible voluntarily by any logged-in user, OR forced when
    # must_change_password is True (e.g. first login after admin creates account).
    forced = getattr(current_user, "must_change_password", False)
    error = None

    if request.method == "POST":
        current_pw = request.form.get("current_password", "")
        new_pw = request.form.get("new_password", "")
        confirm_pw = request.form.get("confirm_password", "")

        # 1. Verify the user knows their present password
        if not current_user.check_password(current_pw):
            error = "Current password is incorrect."
        elif new_pw != confirm_pw:
            error = "New passwords do not match."
        elif new_pw == current_pw:
            error = "New password must be different from your current password."
        else:
            try:
                current_user.set_password(new_pw)  # also validates complexity
                current_user.must_change_password = False
                db.session.commit()
                log_usage("change_password")
                try:
                    os.remove(os.path.join(secure_base, INITIAL_CREDENTIALS_FILENAME))
                except OSError:
                    pass
                flash("Password changed successfully.", "success")
                # Forced change → log out so they re-authenticate with new password
                if forced:
                    logout_user()
                    return redirect(url_for("login"))
                # Voluntary change → stay logged in, return to their dashboard
                role = session.get("role", "")
                return redirect(url_for("admin_dashboard") if role == "Admin" else url_for("index"))
            except ValueError as e:
                error = str(e)

    return render_template("change_password.html", error=error, forced=forced)


@app.route(ROUTE_LOGIN, methods=["GET", "POST"])
# Brute-force throttle: only count actual POST login attempts, NOT page views/refreshes.
# The per-account lockout (MAX_LOGIN_ATTEMPTS fails -> LOCKOUT_MINUTES, below) is the
# primary defence; this per-IP cap is a secondary guard.
@limiter.limit("30 per minute", exempt_when=lambda: request.method != "POST")
def login():
    """Authenticate a user.

    Every failure mode (unknown user, wrong password, locked account) shows the SAME
    generic message so usernames and lockout state cannot be enumerated. The lockout
    itself is still enforced. On success the session is rotated (fixation hygiene),
    last-login is recorded, password expiry is checked, and 2FA users are sent to the
    TOTP verification step before the session is established.
    """
    error = None
    if request.method == "POST":
        role = request.form.get("role", "").strip()
        username = request.form.get("username", "").strip()
        password = request.form.get("password", "")

        user = User.query.filter_by(username=username, role=role).first()

        locked = False
        if user and user.account_locked_until:
            lock_time = (
                user.account_locked_until.replace(tzinfo=timezone.utc)
                if user.account_locked_until.tzinfo is None
                else user.account_locked_until
            )
            locked = lock_time > datetime.now(timezone.utc)

        password_valid = False
        if user:
            password_valid = user.check_password(password)
        else:
            # Dummy hash op to keep response timing consistent when the user doesn't exist.
            # Hash a random throwaway value (no hardcoded credential — SonarCloud S6437).
            check_password_hash(generate_password_hash(secrets.token_urlsafe(16)), password)

        if user and password_valid and not locked:
            user.failed_login_attempts = 0
            user.account_locked_until = None
            prev_login = user.last_login_at
            user.last_login_at = datetime.now(timezone.utc)
            # Password expiry — force a change when the password is too old.
            if PASSWORD_MAX_AGE_DAYS and user.password_changed_at:
                changed = (
                    user.password_changed_at.replace(tzinfo=timezone.utc)
                    if user.password_changed_at.tzinfo is None
                    else user.password_changed_at
                )
                if datetime.now(timezone.utc) - changed > timedelta(days=PASSWORD_MAX_AGE_DAYS):
                    user.must_change_password = True
            db.session.commit()

            # Session-fixation hygiene: never carry pre-auth session state past login.
            session.clear()

            # Users with 2FA enabled must enter a TOTP code first.
            if user.totp_secret and pyotp is not None:
                session["pending_2fa_user_id"] = user.id
                session["pending_2fa_role"] = role
                session["pending_2fa_prev"] = prev_login.isoformat() if prev_login else ""
                session["pending_2fa_ts"] = time.time()
                session["pending_2fa_fails"] = 0
                return redirect(url_for("verify_2fa"))

            return _establish_session(user, role, prev_login)

        # Count the failure (unless already locked) and lock when over the limit.
        if user and not locked:
            user.failed_login_attempts = (user.failed_login_attempts or 0) + 1
            if user.failed_login_attempts >= MAX_LOGIN_ATTEMPTS:
                user.account_locked_until = datetime.now(timezone.utc) + timedelta(minutes=LOCKOUT_MINUTES)
            db.session.commit()

        # Constant-bounds random delay to mask timing differences across all failure modes.
        time.sleep(secrets.SystemRandom().uniform(0.05, 0.15))
        # One generic message for ALL failure modes (anti-enumeration).
        error = "Invalid credentials or role"
    return render_template("login.html", error=error)


def _establish_session(user, role, prev_login):
    """Finish a successful authentication: Flask-Login session + role routing."""
    login_user(user)
    session["username"] = user.username
    session["role"] = role
    # Surface the PREVIOUS login time so users can spot misuse.
    if prev_login:
        session["prev_login"] = to_ist(prev_login).strftime("%d %b %Y, %H:%M IST")
    log_usage("login")

    if getattr(user, "must_change_password", False):
        return redirect(url_for("change_password"))

    next_page = request.args.get("next")
    if next_page and is_safe_url(next_page):
        return redirect(next_page)
    return redirect(url_for("admin_dashboard") if role == "Admin" else url_for("index"))


@app.route("/verify_2fa", methods=["GET", "POST"])
@limiter.limit("10 per minute", exempt_when=lambda: request.method != "POST")
def verify_2fa():
    """Second login step for users with TOTP enabled."""
    uid = session.get("pending_2fa_user_id")
    if not uid or pyotp is None:
        return redirect(url_for("login"))
    user = User.query.get(uid)
    if not user or not user.totp_secret:
        session.pop("pending_2fa_user_id", None)
        return redirect(url_for("login"))

    # Abort if the pending session is stale (> 5 minutes since password was accepted).
    if time.time() - session.get("pending_2fa_ts", 0) > PENDING_2FA_TTL:
        session.clear()
        return redirect(url_for("login"))

    error = None
    if request.method == "POST":
        code = request.form.get("code", "").strip().replace(" ", "")
        # valid_window=0 means only the current 30-s window; still accept ±1 for clock skew.
        if pyotp.TOTP(user.totp_secret).verify(code, valid_window=1) and _totp_is_fresh(uid, code):
            role = session.pop("pending_2fa_role", user.role)
            prev_raw = session.pop("pending_2fa_prev", "")
            session.pop("pending_2fa_user_id", None)
            prev_login = datetime.fromisoformat(prev_raw) if prev_raw else None
            return _establish_session(user, role, prev_login)
        # Count failures; abort back to login after MAX_2FA_ATTEMPTS wrong codes.
        fails = session.get("pending_2fa_fails", 0) + 1
        if fails >= MAX_2FA_ATTEMPTS:
            session.clear()
            return redirect(url_for("login"))
        session["pending_2fa_fails"] = fails
        error = f"Invalid code. {MAX_2FA_ATTEMPTS - fails} attempt(s) remaining."
    return render_template("verify_2fa.html", error=error)


@app.route("/setup_2fa", methods=["GET", "POST"])
@login_required
def setup_2fa():
    """Enable TOTP 2FA for the logged-in user (authenticator app)."""
    if pyotp is None:
        flash("Two-factor support is not installed on this machine (pip install pyotp).", "warning")
        return redirect(url_for("role_home"))

    if request.method == "POST":
        secret = session.get("totp_setup_secret")
        code = request.form.get("code", "").strip().replace(" ", "")
        if secret and pyotp.TOTP(secret).verify(code, valid_window=1):
            current_user.totp_secret = secret
            db.session.commit()
            session.pop("totp_setup_secret", None)
            log_usage("enable_2fa")
            flash("Two-factor authentication enabled.", "success")
            return redirect(url_for("role_home"))
        flash("Code did not match — scan/enter the secret again and retry.", "danger")

    secret = session.get("totp_setup_secret") or pyotp.random_base32()
    session["totp_setup_secret"] = secret
    otpauth_uri = pyotp.TOTP(secret).provisioning_uri(name=current_user.username, issuer_name="FundTrail")
    qr_svg = None
    if qrcode:
        try:
            import qrcode.image.svg as _qr_svg_mod
            img = qrcode.make(otpauth_uri, image_factory=_qr_svg_mod.SvgPathImage)
            buf = io.BytesIO()
            img.save(buf)
            raw = buf.getvalue().decode("utf-8", errors="replace")
            start = raw.find("<svg")
            if start != -1:
                svg = raw[start:]
                # Replace mm-based width/height with explicit px values so the SVG
                # renders at a consistent size across all browsers and platforms.
                svg = re.sub(r'(<svg[^>]*)\bwidth="[^"]*"', r'\1width="220"', svg, count=1)
                svg = re.sub(r'(<svg[^>]*)\bheight="[^"]*"', r'\1height="220"', svg, count=1)
                # Ensure a white background (SvgPathImage has no background rect).
                svg = svg.replace(
                    'xmlns="http://www.w3.org/2000/svg">',
                    'xmlns="http://www.w3.org/2000/svg"><rect width="100%" height="100%" fill="white"/>',
                    1,
                )
                qr_svg = svg
        except Exception as exc:
            app.logger.warning("QR SVG generation failed: %s", exc)
    return render_template(
        "setup_2fa.html", secret=secret, otpauth_uri=otpauth_uri,
        enabled=bool(current_user.totp_secret), qr_svg=qr_svg
    )


@app.route("/disable_2fa", methods=["POST"])
@login_required
def disable_2fa():
    """Turn TOTP off for the logged-in user (requires current password)."""
    if not current_user.check_password(request.form.get("password", "")):
        flash("Password incorrect — 2FA unchanged.", "danger")
        return redirect(url_for("setup_2fa"))
    current_user.totp_secret = None
    db.session.commit()
    log_usage("disable_2fa")
    flash("Two-factor authentication disabled.", "success")
    return redirect(url_for("role_home"))


@app.route("/forgot_password", methods=["GET", "POST"])
def forgot_password():
    """Self-service password reset using TOTP or a validated account email."""
    if current_user.is_authenticated:
        return redirect(url_for("role_home"))

    if request.method == "GET":
        session.pop("fp_step", None)
        session.pop("fp_uid", None)
        session.pop("fp_verified", None)
        return render_template(TEMPLATE_FORGOT_PASSWORD, step=1)

    step = session.get("fp_step", 1)

    # Step 1: username lookup
    if step == 1:
        username = request.form.get("username", "").strip()
        user = User.query.filter_by(username=username).first()
        if not user:
            return render_template(TEMPLATE_FORGOT_PASSWORD, step=1, no_2fa=True)
        session["fp_uid"] = user.id
        session["fp_step"] = 2
        if user.totp_secret and pyotp is not None:
            return render_template(TEMPLATE_FORGOT_PASSWORD, step=2, email_fallback=True)
        return render_template(TEMPLATE_FORGOT_PASSWORD, step=2, email_required=True)

    # Step 2: verify TOTP
    if step == 2:
        uid = session.get("fp_uid")
        user = User.query.get(uid) if uid else None
        if not user:
            return redirect(url_for("forgot_password"))
        if request.form.get("reset_method") == "email":
            _send_password_reset_if_email_matches(user, request.form.get("email", ""))
            return render_template(TEMPLATE_FORGOT_PASSWORD, step=2, email_required=not bool(user.totp_secret and pyotp), email_sent=True)
        code = request.form.get("code", "").strip().replace(" ", "")
        if pyotp and user.totp_secret and pyotp.TOTP(user.totp_secret).verify(code, valid_window=1):
            session["fp_step"] = 3
            session["fp_verified"] = True
            return render_template(TEMPLATE_FORGOT_PASSWORD, step=3)
        return render_template(
            TEMPLATE_FORGOT_PASSWORD,
            step=2,
            error="Invalid code — check your authenticator app and try again.",
            email_fallback=True,
        )

    # Step 3: set new password
    if step == 3:
        if not session.get("fp_verified"):
            return redirect(url_for("forgot_password"))
        uid = session.get("fp_uid")
        user = User.query.get(uid) if uid else None
        if not user:
            return redirect(url_for("forgot_password"))
        new_pw = request.form.get("password", "")
        confirm_pw = request.form.get("confirm_password", "")
        if new_pw != confirm_pw:
            return render_template(TEMPLATE_FORGOT_PASSWORD, step=3, error="Passwords do not match.")
        try:
            user.set_password(new_pw)
            db.session.commit()
            session.pop("fp_step", None)
            session.pop("fp_uid", None)
            session.pop("fp_verified", None)
            log_usage("self_service_pw_reset")
            flash("Password reset successfully. Please sign in with your new password.", "success")
            return redirect(url_for("login"))
        except ValueError as exc:
            return render_template(TEMPLATE_FORGOT_PASSWORD, step=3, error=str(exc))

    return redirect(url_for("forgot_password"))


@app.route("/reset_password/<token>", methods=["GET", "POST"])
def reset_password(token):
    if current_user.is_authenticated:
        return redirect(url_for("role_home"))
    try:
        user = _load_password_reset_user(token)
    except SignatureExpired:
        flash("Password reset link expired. Request a new one.", "danger")
        return redirect(url_for("forgot_password"))
    except BadSignature:
        flash("Password reset link is invalid. Request a new one.", "danger")
        return redirect(url_for("forgot_password"))

    if request.method == "GET":
        return render_template(TEMPLATE_FORGOT_PASSWORD, step=3, reset_token=token)

    new_pw = request.form.get("password", "")
    confirm_pw = request.form.get("confirm_password", "")
    if new_pw != confirm_pw:
        return render_template(TEMPLATE_FORGOT_PASSWORD, step=3, reset_token=token, error="Passwords do not match.")
    try:
        user.set_password(new_pw)
        user.must_change_password = False
        db.session.commit()
        log_usage("self_service_pw_reset_email")
        flash("Password reset successfully. Please sign in with your new password.", "success")
        return redirect(url_for("login"))
    except ValueError as exc:
        return render_template(TEMPLATE_FORGOT_PASSWORD, step=3, reset_token=token, error=str(exc))


@app.route("/logout", methods=["GET", "POST"])
def logout():
    """Log out. POST (CSRF-protected) is preferred; GET remains supported only for
    the legacy link inside the graph page, which must stay untouched."""
    log_usage("logout")
    logout_user()
    session.clear()
    return redirect(ROUTE_LOGIN)


# ---------------------------------------------------------------------------
# 7. Routes: dashboards & Excel upload
# ---------------------------------------------------------------------------


@app.route("/admin_dashboard", methods=["GET"])
def admin_dashboard():
    """Admin home. Redirects (not 403) for non-admins — preserved historical
    behaviour that navigation and the smoke tests rely on."""
    if not is_admin():
        return redirect(url_for("login"))
    my_cases = _cases_q()
    total_cases = my_cases.count()
    active_cases = my_cases.filter(
        or_(Complaint.status.is_(None), Complaint.status != "Closed")
    ).count()
    stats = {
        "total_cases": total_cases,
        "active_cases": active_cases,
        "total_txns": db.session.query(func.count(Transaction.id)).scalar() or 0,
        "held_amount": db.session.query(func.sum(Transaction.put_on_hold_amount)).scalar() or 0.0,
        "recovered_amount": db.session.query(func.sum(Transaction.refund_amount)).scalar() or 0.0,
    }
    recent_cases = _cases_q().order_by(Complaint.upload_time.desc()).limit(10).all()
    id_to_username = {u.id: u.username for u in User.query.with_entities(User.id, User.username).all()}
    recent_activity = UsageLog.query.order_by(UsageLog.timestamp.desc()).limit(5).all()
    return render_template(
        "admin_dashboard.html",
        username=current_user.username,
        stats=stats,
        recent_cases=recent_cases,
        id_to_username=id_to_username,
        recent_activity=recent_activity,
    )


@app.route(ROUTE_INDEX, methods=["GET"])
def index():
    """Officer dashboard: KPI stat cards, recent cases, activity feed, upload."""
    if "username" not in session:
        return redirect(ROUTE_LOGIN)
    # Strict role separation: the admin has their own dashboard. If an admin reaches
    # the officer dashboard, bounce them back to the admin dashboard.
    if session.get("role") == "Admin":
        return redirect(url_for("admin_dashboard"))

    me = session.get("username")
    my_cases = (
        Complaint.query.filter(or_(Complaint.uploaded_by == current_user.id, Complaint.assigned_to == current_user.id))
        .order_by(Complaint.upload_time.desc())
        .all()
    )
    my_file_ids = [f.id for f in UploadedFile.query.with_entities(UploadedFile.id).filter_by(uploader=me)]
    base = Transaction.query.filter(Transaction.upload_id.in_(my_file_ids or [-1]))
    stats = {
        "total_cases": len(my_cases),
        "active_cases": sum(1 for c in my_cases if (c.status or "Open") != "Closed"),
        "total_txns": base.count(),
        "held_amount": (
            db.session.query(func.sum(Transaction.put_on_hold_amount))
            .filter(Transaction.upload_id.in_(my_file_ids or [-1]))
            .scalar()
            or 0.0
        ),
        "recovered_amount": (
            db.session.query(func.sum(Transaction.refund_amount))
            .filter(Transaction.upload_id.in_(my_file_ids or [-1]))
            .scalar()
            or 0.0
        ),
    }
    recent_cases = my_cases[:10]
    recent_activity = UsageLog.query.filter_by(username=me).order_by(UsageLog.timestamp.desc()).limit(5).all()
    return render_template("index.html", stats=stats, recent_cases=recent_cases, recent_activity=recent_activity)


MAX_ROWS = 100000  # Hard row cap per imported Excel file


@contextlib.contextmanager
def file_lock(filepath):
    """Context manager for file locking (Cross-platform)"""
    lock_file = f"{filepath}.lock"
    # Create lock file if it doesn't exist
    lock_fd = os.open(lock_file, os.O_RDWR | os.O_CREAT | os.O_TRUNC)
    try:
        if os.name == "nt":
            # Windows: Blocking lock for 10 attempts (approx 10s)
            msvcrt.locking(lock_fd, msvcrt.LK_RLCK, 1)
        else:
            # Unix: Exclusive blocking lock
            fcntl.flock(lock_fd, fcntl.LOCK_EX)

        yield
    finally:
        # Unlock and clean up
        if os.name == "nt":
            msvcrt.locking(lock_fd, msvcrt.LK_UNLCK, 1)
        else:
            fcntl.flock(lock_fd, fcntl.LOCK_UN)

        os.close(lock_fd)
        try:
            os.remove(lock_file)
        except OSError:
            pass


@app.route("/upload_excel", methods=["POST"])
@login_required
@limiter.limit("20 per hour")
def upload_excel():
    """Import a bank-transaction Excel file and rebuild the case's fund trail."""
    file = request.files.get("excel_file")
    if file is None or not file.filename:
        flash("No file selected. Please choose a .xlsx file.", "warning")
        return redirect(ROUTE_INDEX)

    # secure_filename strips any directory components / path-traversal sequences;
    # enforce the .xlsx extension case-insensitively on the *sanitised* name.
    filename = secure_filename(file.filename)
    if not filename or not filename.lower().endswith(".xlsx"):
        flash("Invalid file type. Only .xlsx workbooks are accepted.", "warning")
        return redirect(ROUTE_INDEX)
    file_path = os.path.join(app.config["UPLOAD_FOLDER"], filename)

    # Use file lock to prevent race conditions
    with file_lock(file_path):
        # Check for duplicate filename
        existing_file = UploadedFile.query.filter_by(filename=filename).first()
        if existing_file:
            old_path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
            try:
                # Remove old file if it exists
                if os.path.exists(old_path):
                    os.remove(old_path)
            except OSError as e:
                logger.warning(f"Failed to delete old file {old_path}: {e}")

            # Delete associated transactions first
            try:
                Transaction.query.filter_by(upload_id=existing_file.id).delete()
            except Exception as e:
                logger.exception(f"Error deleting associated transactions for file {filename}: {e}")

            # Delete database record
            db.session.delete(existing_file)
            db.session.commit()
            logger.info(f"Existing file and record deleted: {filename}")

        file.save(file_path)

    # Defence-in-depth: a genuine .xlsx is an OOXML ZIP container and must begin with
    # the ZIP magic bytes "PK". Reject (and delete) anything else BEFORE the parser
    # touches it — this stops a renamed executable / script / macro-laden blob that
    # slipped past the extension check from ever being parsed or stored in the DB.
    try:
        with open(file_path, "rb") as _sig_fh:
            _magic = _sig_fh.read(4)
    except OSError:
        _magic = b""
    if _magic[:2] != b"PK":
        try:
            os.remove(file_path)
        except OSError:
            pass
        logger.warning("Rejected non-xlsx upload (bad magic bytes): %s", filename)
        flash("That file is not a valid .xlsx workbook and was rejected.", "danger")
        return redirect(ROUTE_INDEX)

    # Any failure while parsing below would otherwise leave the saved working copy
    # orphaned in UPLOAD_FOLDER (an untracked, attacker-supplied blob). Track whether
    # the upload was actually persisted so the finally-block can delete it otherwise.
    upload_persisted = False
    xls = None
    try:
        # ✅ Step 2: Read Excel and extract acknowledgment numbers

        # Helper for sanitizing cell values to prevent CSV/Excel injection
        def sanitize_cell(value):
            """Neutralise spreadsheet formula injection (CWE-1236) WITHOUT corrupting
            legitimate numbers: only quote-prefix a value that starts with a formula
            trigger AND is not a plain number (so '-500' / '+91...' stay intact)."""
            if isinstance(value, str):
                s = value.strip()
                if s[:1] in ("=", "+", "-", "@", "\t", "\r"):
                    try:
                        float(s.replace(",", ""))
                    except ValueError:
                        return "'" + value
            return value

        try:
            xls = pd.ExcelFile(file_path, engine="openpyxl")
        except Exception as e:
            logger.exception(f"Error reading Excel file: {str(e)}", exc_info=True)
            flash("Error reading Excel file. Please ensure it is a valid .xlsx file.", "error")
            return redirect(ROUTE_INDEX)

        # Robustly find the sheet name (handles case/spacing/extra text)
        def find_sheet_name(xls_obj, target):
            names = xls_obj.sheet_names or []
            # Exact match first
            if target in names:
                return target
            # Trim & case-insensitive
            for n in names:
                if n and n.strip().lower() == target.strip().lower():
                    return n
            # Containment (target inside sheet name)
            for n in names:
                if n and target.strip().lower() in n.strip().lower():
                    return n

            # Normalize non-alphanum and compare
            def simple(s):
                return re.sub(r"[^0-9a-z]+", "", str(s).lower())

            t_simple = simple(target)
            for n in names:
                if simple(n) == t_simple:
                    return n
            return None

        sheet_name = find_sheet_name(xls, "Money Transfer to")
        if not sheet_name:
            available = ", ".join([f"'{s}'" for s in (xls.sheet_names or [])])
            raise ValueError(f"Worksheet named 'Money Transfer to' not found. Available sheets: {available}")

        tx_df = pd.read_excel(xls, sheet_name=sheet_name)

        # Check row count
        if len(tx_df) > MAX_ROWS:
            xls.close()
            if os.path.exists(file_path):
                os.remove(file_path)
            flash(f"File contains too many rows. Maximum {MAX_ROWS} rows allowed.", "error")
            return redirect(ROUTE_INDEX)

        # Sanitize all cells in main dataframe
        for col in tx_df.columns:
            tx_df[col] = tx_df[col].apply(sanitize_cell)

        acknos_in_excel = tx_df["Acknowledgement No."].dropna().unique()

        if len(acknos_in_excel) == 0:
            flash("⚠️ No Acknowledgement numbers found in the Excel file!", "warning")
            return redirect(ROUTE_INDEX)

        # ✅ Step 3: Purge ALL existing transactions for these ACK numbers.
        # This is belt-and-suspenders: the filename-dedup above handles same-name
        # re-uploads, but this catches any lingering rows from a differently-named
        # copy of the same data (which would otherwise double every amount shown).
        existing_acks = Transaction.query.filter(Transaction.ack_no.in_(acknos_in_excel)).all()
        if existing_acks:
            # Bulk-delete is faster and avoids per-object overhead
            Transaction.query.filter(Transaction.ack_no.in_(acknos_in_excel)).delete(synchronize_session="fetch")
            db.session.commit()
            flash("ℹ️ Existing records for these ACK numbers replaced.", "info")

        # Re-read file to get binary data for storage
        with open(file_path, "rb") as f:
            file_data = f.read()

        uploaded_file = UploadedFile(
            filename=filename, data=file_data, uploader=session.get("username"), mimetype=file.mimetype
        )
        db.session.add(uploaded_file)
        db.session.commit()
        upload_persisted = True  # from here the disk copy is the legit download source
        log_usage("upload_excel", filename=filename)

        # ✅ Step 5: Process and insert transactions
        atm_df = (
            pd.read_excel(xls, sheet_name=TXN_WITHDRAWAL_ATM)
            if TXN_WITHDRAWAL_ATM in xls.sheet_names
            else pd.DataFrame()
        )
        chq_df = (
            pd.read_excel(xls, sheet_name="Cash Withdrawal through Cheque")
            if "Cash Withdrawal through Cheque" in xls.sheet_names
            else pd.DataFrame()
        )
        hold_df = (
            pd.read_excel(xls, sheet_name="Transaction put on hold")
            if "Transaction put on hold" in xls.sheet_names
            else pd.DataFrame()
        )

        # Sanitize additional dataframes
        if not atm_df.empty:
            for col in atm_df.columns:
                atm_df[col] = atm_df[col].apply(sanitize_cell)

        if not chq_df.empty:
            for col in chq_df.columns:
                chq_df[col] = chq_df[col].apply(sanitize_cell)

        if not hold_df.empty:
            for col in hold_df.columns:
                hold_df[col] = hold_df[col].apply(sanitize_cell)

        def normalize_columns(df):
            return [
                str(c).encode("ascii", "ignore").decode().strip().replace("\u00a0", " ").replace("\xa0", " ")
                for c in df.columns
            ]

        # Normalize all dataframe columns for consistent matching
        tx_df.columns = normalize_columns(tx_df)
        if not atm_df.empty:
            atm_df.columns = normalize_columns(atm_df)
        if not chq_df.empty:
            chq_df.columns = normalize_columns(chq_df)
        if not hold_df.empty:
            hold_df.columns = normalize_columns(hold_df)

        # DEBUG: Log all columns in Excel
        logger.debug("[UPLOAD DEBUG] All columns in 'Money Transfer to' sheet:")
        for i, col in enumerate(tx_df.columns):
            logger.debug(f"  {i}: '{col}'")

        # DEBUG: Log presence of data (without sensitive details)
        if not tx_df.empty:
            logger.debug(f"[UPLOAD DEBUG] First row data present. Columns: {list(tx_df.columns)}")

        def clean_amount(value):
            if pd.isna(value):
                return 0.0
            try:
                return float(str(value).replace(",", "").strip())
            except Exception:
                return 0.0

        def clean_location(value):
            """Normalize ATM location strings and strip common prefixes."""
            if pd.isna(value) or value is None:
                return None
            s = str(value).strip()
            prefixes = [
                "Place of ATM :-",
                "Place of ATM :",
                "Place/Location of ATM :",
                "Place/Location of ATM",
                "Place / Location of ATM",
                "Place of ATM",
                "Place/Location",
            ]
            for p in prefixes:
                if s.lower().startswith(p.lower()):
                    s = s[len(p) :].strip(" :-")
                    break
            return s or None

        def clean_bank_name(value):
            """Clean bank name by removing HTML tags and truncating to 100 characters."""
            if pd.isna(value) or value is None:
                return ""
            # Convert to string and strip whitespace
            s = str(value).strip()
            # Remove HTML tags (e.g., <hr/>, <br/>, etc.)
            s = re.sub(r"<[^>]+>", "", s)
            # Replace multiple spaces with single space
            s = re.sub(r"\s+", " ", s)
            # Strip again after HTML removal
            s = s.strip()
            # Truncate to 100 characters (column size limit)
            if len(s) > 100:
                s = s[:100]
            return s

        def safe_txn_id(val):
            if pd.isna(val) or val is None:
                return ""
            s = str(val).strip()
            if re.fullmatch(r"\d+\.0", s):
                s = s[:-2]
            if re.fullmatch(r"\d+(\.\d+)?e\+\d+", s, flags=re.IGNORECASE):
                try:
                    num = pd.to_numeric(val, errors="coerce")
                    if pd.notna(num):
                        s = f"{int(num):d}"
                except Exception:
                    pass
            return s

        def get_first_value(df, columns):
            """Return first non-empty value for given column names from df."""
            if df.empty:
                return None
            for col in columns:
                if col in df.columns:
                    val = df.iloc[0].get(col)
                    if pd.notna(val) and str(val).strip():
                        return str(val).strip()
            return None

        def norm(s):
            s = str(s).replace("\u00a0", " ")
            s = re.sub(COLUMN_NAME_NORM_RE, " ", s).lower().strip()
            return s

        def get_txn_id_from_row(row, utr2_col=None):
            """Extract Transaction ID / UTR Number2 from a row."""
            # Try the detected column first (if provided)
            if utr2_col and utr2_col in row.index:
                val = row.get(utr2_col, "")
                s = safe_txn_id(val)
                if s:
                    return s

            # Try common column name variants (exact matches) - EXPANDED LIST
            variants = UTR_TXN_ID_VARIANTS
            for col in variants:
                if col in row.index:
                    val = row.get(col, "")
                    s = safe_txn_id(val)
                    if s:
                        return s

            # Fuzzy matching: Look for columns containing UTR and Number 2 in normalized form
            def norm(s):
                s = str(s).replace("\u00a0", " ")
                s = re.sub(COLUMN_NAME_NORM_RE, " ", s).lower().strip()
                return s

            for col in row.index:
                nc = norm(col)
                # Match columns that have UTR and (NUMBER 2 or ends with 2) and (TRANSACTION or TXN or ID)
                has_utr = "utr" in nc
                has_number2 = (UTR_FUZZY_NUMBER2 in nc) or ("number2" in nc) or nc.endswith(" 2")
                has_txn_id = any(x in nc for x in ["transaction", "txn", "id"])

                if has_utr and has_number2 and has_txn_id:
                    val = row.get(col, "")
                    s = safe_txn_id(val)
                    if s:
                        return s

            # Last resort: Try any column with just "UTR" and "NUMBER"
            for col in row.index:
                nc = norm(col)
                if "utr" in nc and "number" in nc:
                    val = row.get(col, "")
                    s = safe_txn_id(val)
                    if s:
                        return s

            return ""

        def find_utr2_column(columns):
            """Find the 'Transaction ID / UTR Number2' column among various possible names."""
            # First, try exact matches
            exact_matches = UTR_COLUMN_VARIANTS
            for col in columns:
                if col in exact_matches:
                    logger.debug(f"[FIND_UTR2] Exact match found: {col}")
                    return col

            # Then try fuzzy matching with normalized names
            def norm(s):
                s = str(s).replace("\u00a0", " ")
                s = re.sub(COLUMN_NAME_NORM_RE, " ", s).lower().strip()
                return s

            known_normalized = [
                "transaction id utr number2",
                "transaction id utr number 2",
                "transaction id  utr number2",
                "transaction id  utr number 2",
                "transaction id number2",
                "transaction id utr",
                "txn id utr number2",
                "txn id utr number 2",
                "txn id utr",
                "utr number2",
                "utr number 2",
            ]

            normalized_map = {col: norm(col) for col in columns}
            for col, nc in normalized_map.items():
                if nc in known_normalized:
                    logger.debug(f"[FIND_UTR2] Normalized match found: '{col}' -> '{nc}'")
                    return col

            # Final fallback: Look for columns with all required components
            for col, nc in normalized_map.items():
                has_utr = "utr" in nc
                has_number2 = (UTR_FUZZY_NUMBER2 in nc) or ("number2" in nc) or nc.endswith(" 2")
                has_txn_id = any(x in nc for x in ["transaction", "txn", "id"])

                if has_utr and has_number2 and has_txn_id:
                    logger.debug(f"[FIND_UTR2] Fuzzy pattern match found: '{col}' -> '{nc}'")
                    return col

            # Try ANY column with UTR in name
            for col, nc in normalized_map.items():
                if "utr" in nc and "number" in nc:
                    logger.debug(f"[FIND_UTR2] Fallback match (UTR + Number): '{col}' -> '{nc}'")
                    return col

            logger.debug(f"[FIND_UTR2] No match found. Available normalized columns: {list(normalized_map.values())}")
            return None

        utr2_col = find_utr2_column(tx_df.columns)
        logger.debug(f"[UPLOAD] Excel columns in 'Money Transfer to': {list(tx_df.columns)}")
        logger.debug(f"[UPLOAD] Detected UTR2 column: {utr2_col}")

        # Performance: pre-index the ATM / cheque / hold sheets ONCE, keyed by the
        # account number (and, for holds, the txn id). The previous code re-filtered
        # the entire sheet inside the per-row loop (`df[df[col].astype(str)...] == x`),
        # which is O(rows × sheet) and re-stringifies the whole column every iteration —
        # the reason a "small" upload felt slow. Now each row is an O(1) dict lookup.
        _ACC_COL = "Account No./ (Wallet /PG/PA) Id"

        def _index_by_account(df):
            if df.empty or _ACC_COL not in df.columns:
                return {}
            tmp = df.copy()
            tmp["_acc_key"] = tmp[_ACC_COL].astype(str).str.strip()
            return dict(tmp.groupby("_acc_key"))

        atm_by_acc = _index_by_account(atm_df)
        chq_by_acc = _index_by_account(chq_df)

        hold_by_key = {}
        if not hold_df.empty and _ACC_COL in hold_df.columns and "Transaction Id / UTR Number" in hold_df.columns:
            _h = hold_df.copy()
            _h["_acc_key"] = _h[_ACC_COL].astype(str).str.strip()
            _h["_txn_key"] = _h["Transaction Id / UTR Number"].astype(str).str.strip()
            hold_by_key = dict(_h.groupby(["_acc_key", "_txn_key"]))

        _EMPTY_DF = pd.DataFrame()
        transactions = []

        txn_id_counts = {"found": 0, "missing": 0}
        for idx, (_, row) in enumerate(tx_df.iterrows()):
            ack_no = str(row.get("Acknowledgement No.", "")).strip()
            if not ack_no:
                # ✅ Skip rows with missing ACK
                continue

            acc_to = str(row.get("Account No", "")).strip()
            extracted_txn_id = get_txn_id_from_row(row, utr2_col)
            atm_info = atm_by_acc.get(acc_to, _EMPTY_DF)
            chq_info = chq_by_acc.get(acc_to, _EMPTY_DF)
            hold_info = (
                hold_by_key.get((acc_to, extracted_txn_id), _EMPTY_DF) if extracted_txn_id else _EMPTY_DF
            )
            if extracted_txn_id:
                txn_id_counts["found"] += 1
                if idx < 2:  # Log first 2 rows with txn_id
                    logger.debug(f"[UPLOAD] Row {idx}: ACK={ack_no}, TXN_ID={extracted_txn_id}")
            else:
                txn_id_counts["missing"] += 1
                if idx < 2:  # Log first 2 rows without txn_id
                    logger.debug(f"[UPLOAD] Row {idx}: ACK={ack_no}, TXN_ID=EMPTY")
                    # Debug: log what's in the row related to UTR/Transaction
                    for col in row.index:
                        if "utr" in col.lower() or "transaction" in col.lower() or "txn" in col.lower():
                            logger.debug(f"  -> {col}: [REDACTED]")

            # Validate before creating transaction
            from_account = str(row.get("Account No./ (Wallet /PG/PA) Id", "")).strip()
            to_account = acc_to  # Already stripped and extracted above
            raw_amount = row.get("Transaction Amount")
            cleaned_amount = clean_amount(raw_amount)

            if not validate_account_number(from_account):
                logger.warning(f"Invalid from_account: {from_account}")
                continue

            if not validate_account_number(to_account):
                logger.warning(f"Invalid to_account: {to_account}")
                continue

            if not validate_amount(cleaned_amount):
                logger.warning(f"Invalid amount: {cleaned_amount}")
                continue

            transaction = Transaction(
                layer=int(row.get("Layer", 0)),
                from_account=from_account,
                to_account=to_account,
                account_number=to_account,
                ack_no=ack_no,
                bank_name=clean_bank_name(row.get("Bank/FIs")),
                ifsc_code=str(row.get("Ifsc Code", "")).strip(),
                txn_date=str(row.get("Transaction Date", "")).strip(),
                txn_id=extracted_txn_id,
                amount=cleaned_amount,
                disputed_amount=clean_amount(row.get("Disputed Amount")),
                action_taken=str(row.get("Action Taken By bank", "")).strip(),
                atm_id=str(atm_info.iloc[0]["ATM ID"]) if not atm_info.empty else None,
                atm_withdraw_amount=clean_amount(atm_info.iloc[0]["Withdrawal Amount"]) if not atm_info.empty else None,
                atm_withdraw_date=str(atm_info.iloc[0]["Withdrawal Date & Time"]) if not atm_info.empty else None,
                atm_location=clean_location(
                    get_first_value(
                        atm_info,
                        [
                            "ATM Location",
                            "Location",
                            "ATM Location / City",
                            "ATM Address",
                            "ATM Location/City",
                            "Place/Location of ATM",
                            "Place / Location of ATM",
                        ],
                    )
                ),
                cheque_no=str(chq_info.iloc[0]["Cheque No"]) if not chq_info.empty else None,
                cheque_withdraw_amount=clean_amount(chq_info.iloc[0]["Withdrawal Amount"])
                if not chq_info.empty
                else None,
                cheque_withdraw_date=str(chq_info.iloc[0]["Withdrawal Date & Time"]) if not chq_info.empty else None,
                cheque_ifsc=str(chq_info.iloc[0]["Ifsc Code"]) if not chq_info.empty else None,
                put_on_hold_txn_id=str(hold_info.iloc[0]["Transaction Id / UTR Number"])
                if not hold_info.empty
                else None,
                put_on_hold_date=str(hold_info.iloc[0]["Put on hold Date"]) if not hold_info.empty else None,
                put_on_hold_amount=clean_amount(hold_info.iloc[0]["Put on hold Amount"])
                if not hold_info.empty
                else None,
                upload_id=uploaded_file.id,
            )

            # Restore saved refund details if any (Persistent POH Data)
            if transaction.put_on_hold_txn_id:
                poh_details = POHRefundDetails.query.filter_by(
                    ack_no=transaction.ack_no, txn_id=transaction.put_on_hold_txn_id
                ).first()
                if poh_details:
                    transaction.court_order_date = poh_details.court_order_date
                    transaction.refund_status = poh_details.refund_status
                    transaction.refund_amount = poh_details.refund_amount

            # Collect transaction in list instead of adding directly
            transactions.append(transaction)

        # ✅ Sort transactions by layer in ascending order before inserting
        transactions.sort(key=lambda t: t.layer)

        # Add sorted transactions to the session
        for t in transactions:
            db.session.add(t)

        db.session.commit()

        # ── Per-officer isolation: register a Complaint per ACK so the uploader
        #    (and any admin-assigned officer) can see/access this case. Without it,
        #    check_case_access() and available_ack_nos() treat the case as nobody's.
        #    Get-or-create; never clobber an admin's manual reassignment.
        try:
            for _ack in acknos_in_excel:
                _ack_s = str(_ack).strip()
                if not _ack_s:
                    continue
                _c = Complaint.query.filter_by(ack_no=_ack_s).first()
                # Determine the owning admin: admins own their own cases directly;
                # officers' cases belong to their managing admin.
                _owner_admin = (
                    current_user.id
                    if is_admin()
                    else getattr(current_user, "admin_id", None)
                )
                if not _c:
                    db.session.add(
                        Complaint(
                            ack_no=_ack_s,
                            file_name=filename,
                            uploaded_by=current_user.id,
                            assigned_to=current_user.id,
                            upload_time=datetime.now(timezone.utc),
                            owner_admin_id=_owner_admin,
                        )
                    )
                else:
                    _c.file_name = filename
                    # Same ACK re-uploaded (e.g. the bank sent an updated Excel): refresh
                    # the upload date so the case reflects when it was last updated. The
                    # refund status lives in POHRefundDetails and is left untouched.
                    _c.upload_time = datetime.now(timezone.utc)
                    if _c.uploaded_by is None:
                        _c.uploaded_by = current_user.id
                    if _c.assigned_to is None:
                        _c.assigned_to = current_user.id
                    if _c.owner_admin_id is None:
                        _c.owner_admin_id = _owner_admin
            db.session.commit()
        except Exception as _ce:
            db.session.rollback()
            logger.exception(f"Failed to register complaint rows for upload {filename}: {_ce}")

        logger.info(
            f"[UPLOAD] Txn ID extraction summary: Found={txn_id_counts['found']}, Missing={txn_id_counts['missing']}"
        )
        flash("✅ Excel uploaded and data saved successfully.", "success")

    except Exception as e:
        db.session.rollback()
        logger.exception(f"Failed to process Excel: {e}", exc_info=True)
        flash("Failed to process Excel file. Please check the file format and try again.", "danger")
    finally:
        # Close the workbook handle (Windows won't delete an open file) and purge the
        # orphaned working copy whenever the upload did not complete — so an
        # unsupported / unparseable file never lingers in UPLOAD_FOLDER.
        if xls is not None:
            try:
                xls.close()
            except Exception:
                pass
        if not upload_persisted:
            try:
                if os.path.exists(file_path):
                    os.remove(file_path)
                    logger.info("Removed orphaned upload (not persisted): %s", filename)
            except OSError as exc:
                logger.warning("Could not remove orphaned upload %s: %s", filename, exc)

    return redirect(ROUTE_INDEX)


@app.route("/download/<filename>", methods=["GET"])
@login_required
def download_file(filename):
    # Sanitize filename
    safe_filename = secure_filename(filename)

    # Verify file exists in database (authorization check)
    uploaded_file = UploadedFile.query.filter_by(filename=safe_filename).first()
    if not uploaded_file:
        abort(404)

    # Construct safe path
    file_path = os.path.join(app.config["UPLOAD_FOLDER"], safe_filename)

    # Verify path is within upload directory
    upload_dir = os.path.abspath(app.config["UPLOAD_FOLDER"])
    requested_path = os.path.abspath(file_path)

    if not requested_path.startswith(upload_dir):
        abort(403)

    # Verify file exists
    if not os.path.isfile(requested_path):
        abort(404)

    log_usage("download_file", filename=safe_filename)
    return send_from_directory(app.config["UPLOAD_FOLDER"], safe_filename)


# ---------------------------------------------------------------------------
# 8. Routes: fund-trail graph & transaction views
# ---------------------------------------------------------------------------


@app.route("/view_graph", methods=["GET"])
@login_required
def view_graph():
    ack_no = request.args.get("ack_no")
    check_case_access(ack_no)
    try:
        fname_row = (
            db.session.query(UploadedFile.filename)
            .join(Transaction, Transaction.upload_id == UploadedFile.id)
            .filter(Transaction.ack_no == ack_no)
            .order_by(UploadedFile.upload_time.desc())
            .first()
        )
        fname = fname_row[0] if fname_row else None
        log_usage("view_graph", filename=fname, ack_no=ack_no)
    except Exception as e:
        logger.exception(f"UsageLog view_graph error: {e}")
    return redirect(url_for("graph_tree1", ack_no=ack_no))


@app.route("/graph/<ack_no>", methods=["GET"])
@login_required
def graph_tree1(ack_no):
    check_case_access(ack_no)
    try:
        fname_row = (
            db.session.query(UploadedFile.filename)
            .join(Transaction, Transaction.upload_id == UploadedFile.id)
            .filter(Transaction.ack_no == ack_no)
            .order_by(UploadedFile.upload_time.desc())
            .first()
        )
        fname = fname_row[0] if fname_row else None
        log_usage("graph_page", filename=fname, ack_no=ack_no)
    except Exception as e:
        logger.exception(f"UsageLog graph_page error: {e}")

    # Calculate Layer 1 Total
    layer_1_total = 0.0
    formatted_l1_total = "0.00"
    try:
        l1_txns = Transaction.query.filter_by(ack_no=ack_no, layer=1).all()
        if l1_txns:
            layer_1_total = sum(t.amount for t in l1_txns if t.amount)

        def format_indian_currency(amount):
            try:
                amount = float(amount)
            except (ValueError, TypeError):
                return "0.00"
            s = "{:.2f}".format(amount)
            parts = s.split(".")
            integer_part = parts[0]
            if len(integer_part) > 3:
                last_three = integer_part[-3:]
                rest = integer_part[:-3]
                rest_formatted = ""
                while len(rest) > 2:
                    rest_formatted = "," + rest[-2:] + rest_formatted
                    rest = rest[:-2]
                rest_formatted = rest + rest_formatted
                integer_part = rest_formatted + "," + last_three
            return f"₹{integer_part}"

        formatted_l1_total = format_indian_currency(layer_1_total)
    except Exception as e:
        logger.exception(f"Error calculating layer 1 total: {e}")

    return render_template(
        "graph_tree1.html", ack_no=ack_no, role=session.get("role"), layer_1_total=formatted_l1_total
    )


@app.route("/complaints", methods=["GET"])
@login_required
def complaints():
    """'My Cases' — the real cases the current user may access (per-officer
    isolation): admin sees every registered case; an officer sees only cases they
    uploaded or that an admin assigned to them."""
    try:
        log_usage("view_complaints")
    except Exception:
        pass

    rows = _cases_q().order_by(Complaint.upload_time.desc()).all()

    # Only mark cases that actually have transaction data as viewable.
    acks_with_txn = {
        row[0] for row in db.session.query(Transaction.ack_no).filter(Transaction.ack_no.isnot(None)).distinct().all()
    }
    id_to_username = {u.id: u.username for u in User.query.with_entities(User.id, User.username).all()}

    cases = [
        {
            "ack_no": c.ack_no,
            "file_name": c.file_name or "—",
            "assigned_to": id_to_username.get(c.assigned_to) or "Unassigned",
            "upload_time": c.upload_time,
            "has_data": c.ack_no in acks_with_txn,
            "status": c.status or "Open",
        }
        for c in rows
    ]

    return render_template("complaint.html", cases=cases, statuses=CASE_STATUSES)


@app.route("/graph_data/<ack_no>", methods=["GET"])
@login_required
def graph_data(ack_no):
    # Authorization check
    check_case_access(ack_no)

    try:
        ack_no = ack_no.strip()
        logger.info(f"Fetching graph data for ACK: {ack_no}")
        transactions = Transaction.query.filter_by(ack_no=ack_no).all()
        logger.info(f"Found {len(transactions)} transactions for ACK {ack_no}")

        # ✅ Restore POH Refund Details from persistent store (POHRefundDetails)
        # This ensures that even if the Transaction table is stale or reset, the refund details are preserved and displayed.
        try:
            poh_details_list = POHRefundDetails.query.filter_by(ack_no=ack_no).all()
            if poh_details_list:
                poh_map = {p.txn_id: p for p in poh_details_list}
                updated_count = 0
                for t in transactions:
                    if t.put_on_hold_txn_id and t.put_on_hold_txn_id in poh_map:
                        pdata = poh_map[t.put_on_hold_txn_id]
                        # Update in-memory object attributes for graph generation
                        t.court_order_date = pdata.court_order_date
                        t.refund_status = pdata.refund_status
                        t.refund_amount = pdata.refund_amount
                        updated_count += 1
                logger.info(f"Restored persistent POH details for {updated_count} transactions")
        except Exception as e:
            logger.exception(f"Error restoring POH details in graph_data: {e}")

        # ✅ Restore KYC Details from persistent store (KYCDetails)
        try:
            txn_ids = [t.txn_id for t in transactions if t.txn_id]
            if txn_ids:
                # SQLite has limit on variables in IN clause, but usually 999. If txn count is high, might need chunking.
                # For now assuming manageable chunk or SQLAlchemy handles it (it often doesn't automatically chunk large lists).
                # To be safe, let's just query what we can. If it fails, we catch exception.
                kyc_list = KYCDetails.query.filter(KYCDetails.txn_id.in_(txn_ids)).all()
                kyc_map = {k.txn_id: k for k in kyc_list}

                kyc_updated_count = 0
                for t in transactions:
                    if t.txn_id and t.txn_id in kyc_map:
                        kdata = kyc_map[t.txn_id]
                        t.kyc_name = kdata.name
                        t.kyc_aadhar = kdata.aadhar
                        t.kyc_mobile = kdata.mobile
                        t.kyc_address = kdata.address
                        kyc_updated_count += 1
                logger.info(f"Restored persistent KYC details for {kyc_updated_count} transactions")
        except Exception as e:
            logger.exception(f"Error restoring KYC details in graph_data: {e}")

        if not transactions:
            logger.warning(f"No transactions found for ACK {ack_no}")
            return jsonify({"error": "No transactions found for this Acknowledgement No."})

        from_to_map = defaultdict(lambda: defaultdict(list))
        incoming_map = defaultdict(list)

        # Dedup on the FULL transaction identity, not just txn_id. The old check
        # treated every blank txn_id as "already seen", so multiple distinct transfers
        # between the same two accounts that lacked a UTR collapsed into one (a cause of
        # the "missing transactions / partial data" symptom). Keep rows differing in
        # amount/date/txn_id; drop only exact duplicates.
        seen_edge_keys = set()
        for t in transactions:
            key = (t.from_account, t.to_account, t.amount, t.txn_date, t.txn_id)
            if key in seen_edge_keys:
                continue
            seen_edge_keys.add(key)
            from_to_map[t.from_account][t.to_account].append(
                {"txn_id": t.txn_id, "amount": t.amount, "date": t.txn_date, "ack_no": t.ack_no}
            )
            incoming_map[t.to_account].append(
                {
                    "from_account": t.from_account,
                    "amount": t.amount,
                    "date": t.txn_date,
                    "ack_no": t.ack_no,
                    "txn_id": t.txn_id,
                }
            )

        from_layer_map = {t.from_account: t.layer for t in transactions if t.from_account}

        def build_hierarchy(rows):
            root = {"name": "Flow", "children": []}

            def find_node(n, target):
                if n["name"] == target:
                    return n
                for c in n.get("children", []):
                    found = find_node(c, target)
                    if found:
                        return found
                return None

            for t in rows:
                if t.layer == 1:
                    parent = next((c for c in root["children"] if c["name"] == t.from_account), None)
                    if not parent:
                        parent = {
                            "name": t.from_account,
                            "children": [],
                            "kyc_name": t.kyc_name,
                            "kyc_aadhar": t.kyc_aadhar,
                            "kyc_mobile": t.kyc_mobile,
                            "kyc_address": t.kyc_address,
                            "action": t.action_taken,
                        }
                        root["children"].append(parent)
                else:
                    parent = find_node(root, t.from_account)

                if parent:
                    existing = next((c for c in parent["children"] if c["name"] == t.to_account), None)
                    if not existing:
                        # ✅ calculate total amount of all transactions between parent → child
                        child = {
                            "name": t.to_account,
                            "children": [],
                            "layer": from_layer_map.get(t.to_account, t.layer),
                            "ack": t.ack_no,
                            "bank": t.bank_name,
                            "ifsc": t.ifsc_code,
                            "date": t.txn_date,
                            "txid": (
                                t.txn_id
                                or next(
                                    (
                                        tx.get("txn_id")
                                        for tx in from_to_map[t.from_account][t.to_account]
                                        if tx.get("txn_id")
                                    ),
                                    None,
                                )
                            ),
                            "amt": str(t.amount),
                            "disputed": str(t.disputed_amount),
                            "action": t.action_taken,
                            "state": t.state if t.state and t.state != "Unknown" else (t.ifsc_code or "Unknown State"),
                            "atm_info": {
                                "atm_id": t.atm_id,
                                "amount": t.atm_withdraw_amount,
                                "date": t.atm_withdraw_date,
                                "location": t.atm_location,
                            }
                            if t.atm_id
                            else None,
                            "cheque_info": {
                                "cheque_no": t.cheque_no,
                                "amount": t.cheque_withdraw_amount,
                                "date": t.cheque_withdraw_date,
                                "ifsc": t.cheque_ifsc,
                            }
                            if t.cheque_no
                            else None,
                            "hold_info": {
                                "txn_id": t.put_on_hold_txn_id,
                                "amount": t.put_on_hold_amount,
                                "date": t.put_on_hold_date,
                                "court_order_date": t.court_order_date,
                                "refund_status": t.refund_status,
                                "refund_amount": t.refund_amount,
                            }
                            if t.put_on_hold_txn_id
                            else None,
                            "kyc_name": t.kyc_name,
                            "kyc_aadhar": t.kyc_aadhar,
                            "kyc_mobile": t.kyc_mobile,
                            "kyc_address": t.kyc_address,
                            # ✅ Keep all individual transactions for popup
                            "transactions_from_parent": from_to_map[t.from_account][t.to_account],
                        }
                        parent["children"].append(child)

            return root

        result = build_hierarchy(transactions)
        logger.info(f"Built hierarchy with {len(result['children'])} root children")
        return jsonify(result)
    except Exception as e:
        logger.exception(f"Error processing graph data for ACK {ack_no}: {e}", exc_info=True)
        # Return a safe error message to the client while logging full traceback to server logs
        return jsonify({"error": "Internal server error while processing graph data."}), 500


@app.route("/available_ack_nos", methods=["GET"])
@login_required
def available_ack_nos():
    """List all available ACK numbers accessible to the current user"""
    # Use the same scoped query as the rest of the app so group isolation is consistent.
    allowed_ack_nos = [
        row.ack_no for row in _cases_q().with_entities(Complaint.ack_no).all() if row.ack_no
    ]
    if not allowed_ack_nos:
        return jsonify({"available_ack_nos": []})
    ack_nos = (
        db.session.query(Transaction.ack_no)
        .filter(Transaction.ack_no.in_(allowed_ack_nos))
        .distinct()
        .all()
    )
    ack_list = [ack[0] for ack in ack_nos if ack[0]]
    return jsonify({"available_ack_nos": sorted(ack_list)})


# Parsed ATM-sheet results cached by (upload_id, ack_no). The stored Excel blob is
# immutable per upload, so re-parsing it on every click was pure waste (the slow part).
_ATM_RESULT_CACHE = {}
_ATM_CACHE_MAX = 64  # bound the cache so it can't grow without limit (was a slow leak)


@app.route("/atm_data/<ack_no>", methods=["GET"])
@login_required
def atm_data(ack_no):
    """Return rows from the ATM-related sheet for the uploaded Excel associated with this ack_no."""
    check_case_access(ack_no)
    try:
        # Find the most recent uploaded file for this ack_no
        up_row = (
            db.session.query(UploadedFile)
            .join(Transaction, Transaction.upload_id == UploadedFile.id)
            .filter(Transaction.ack_no == ack_no)
            .order_by(UploadedFile.upload_time.desc())
            .first()
        )

        if not up_row:
            return jsonify({"atm": []})

        up = up_row
        _ck = (up.id, ack_no)
        if _ck in _ATM_RESULT_CACHE:  # fast path: already parsed this file
            return jsonify(_ATM_RESULT_CACHE[_ck])
        file_bytes = None
        if up.data:
            file_bytes = up.data
        else:
            # Fallback: try to read the saved file from uploads folder
            # Ensure absolute path usage
            upload_folder = app.config.get("UPLOAD_FOLDER", "uploads")
            if not os.path.isabs(upload_folder):
                upload_folder = os.path.join(app.root_path, upload_folder)

            possible_path = os.path.join(upload_folder, up.filename) if up.filename else None

            if possible_path and os.path.exists(possible_path):
                try:
                    with open(possible_path, "rb") as fh:
                        file_bytes = fh.read()
                except Exception as ex:
                    logger.warning(f"[atm_data] Failed to read file from disk: {ex}")
            else:
                logger.warning(f"[atm_data] File not found at {possible_path}")

        if not file_bytes:
            return jsonify({"atm": []})

        xls = pd.ExcelFile(io.BytesIO(file_bytes))

        # Robust sheet finding
        sheet_name = None
        # 1. Exact match candidates
        candidates = [TXN_WITHDRAWAL_ATM, "Withdrawal through ATM ", "Withdrawal through ATM\n"]
        for c in candidates:
            if c in xls.sheet_names:
                sheet_name = c
                break

        # 2. Case-insensitive match
        if not sheet_name:
            lower_sheets = {s.lower().strip(): s for s in xls.sheet_names}
            if "withdrawal through atm" in lower_sheets:
                sheet_name = lower_sheets["withdrawal through atm"]

        # 3. Partial match
        if not sheet_name:
            for s in xls.sheet_names:
                if "withdrawal" in s.lower() and "atm" in s.lower():
                    sheet_name = s
                    break

        if not sheet_name:
            return jsonify({"atm": []})

        df = pd.read_excel(xls, sheet_name=sheet_name)

        # Filter by ACK Number if the column exists
        ack_col = None
        for col in df.columns:
            if "acknowledgement" in str(col).lower() and "no" in str(col).lower():
                ack_col = col
                break

        if ack_col:
            # Normalize ACK numbers for comparison
            df["ack_str"] = df[ack_col].astype(str).str.strip()
            # Filter logic can be added here if needed, currently returns all for the file

        df = df.fillna("")

        # Ensure columns are strings and replace non-ascii spacing
        def normalize_columns(cols):
            return [
                str(c).encode("ascii", "ignore").decode().strip().replace("\u00a0", " ").replace("\xa0", " ")
                for c in cols
            ]

        df.columns = normalize_columns(df.columns)
        logger.debug(f"[atm_data] Normalized columns: {list(df.columns)}")

        # Fetch layer mapping from Transactions for this ACK No
        layer_map = {}
        try:
            # Get distinct to_account -> layer mapping
            txns = (
                db.session.query(Transaction.to_account, Transaction.layer)
                .filter(Transaction.ack_no == ack_no)
                .filter(Transaction.to_account.isnot(None))
                .distinct()
                .all()
            )

            def normalize_db_acc(val):
                s = str(val).strip()
                # If it looks like a float ending in .0, strip it
                if s.endswith(".0"):
                    return s[:-2]
                return s

            for acc, lay in txns:
                if acc:
                    norm_acc = normalize_db_acc(acc)
                    layer_map[norm_acc] = lay
                    # Also store the original just in case
                    layer_map[str(acc).strip()] = lay

            logger.info(f"[atm_data] Built layer map for {len(layer_map)} accounts")
        except Exception as e:
            logger.exception(f"[atm_data] Error building layer map: {e}")

        # Find Account Number column in the dataframe
        acc_col = None
        for col in df.columns:
            c_lower = col.lower()
            # Common patterns for account number column
            if "account no" in c_lower or "account number" in c_lower:
                acc_col = col
                break

        if acc_col:
            logger.info(f"[atm_data] Found account column: {acc_col}")

            def get_layer(row):
                val = row.get(acc_col)
                if pd.isna(val) or val == "":
                    return ""

                # Normalize the ATM account value
                s_val = str(val).strip()

                # 1. Try exact match
                if s_val in layer_map:
                    return layer_map[s_val]

                # 2. Try removing .0 suffix (common in pandas float->str)
                if s_val.endswith(".0") and s_val[:-2] in layer_map:
                    return layer_map[s_val[:-2]]

                # 3. Try converting scientific notation or float to int string
                try:
                    f_val = float(s_val)
                    int_val = str(int(f_val))
                    if int_val in layer_map:
                        return layer_map[int_val]
                except (ValueError, TypeError):
                    pass

                return ""

            df["Layer"] = df.apply(get_layer, axis=1)
        else:
            logger.warning("[atm_data] Could not find Account Number column in ATM sheet")
            df["Layer"] = ""

        rows = df.fillna("").to_dict(orient="records")
        logger.info(f"[atm_data] Rows extracted: {len(rows)}")
        _result = {"atm": rows, "columns": list(df.columns)}
        if len(_ATM_RESULT_CACHE) >= _ATM_CACHE_MAX:  # bound the cache (simple evict-all)
            _ATM_RESULT_CACHE.clear()
        _ATM_RESULT_CACHE[(up.id, ack_no)] = _result  # cache: blob is immutable per upload
        return jsonify(_result)
    except Exception as e:
        logger.exception(f"[atm_data] Error: {e}", exc_info=True)
        return jsonify({"atm": [], "error": "Failed to load ATM data"})


@app.route("/statewise_summary/<ack_no>", methods=["GET"])
@login_required
def statewise_summary(ack_no):
    check_case_access(ack_no)
    try:
        # Check if we have any transactions with known states
        known_states_count = (
            db.session.query(Transaction)
            .filter(Transaction.ack_no == ack_no, Transaction.state.isnot(None), Transaction.state != "Unknown")
            .count()
        )

        if known_states_count == 0:
            # If no known states, fetch synchronously for immediate response
            transactions_unknown = Transaction.query.filter(
                Transaction.ack_no == ack_no,
                (Transaction.state.is_(None) | (Transaction.state == "Unknown")),
                Transaction.ifsc_code.isnot(None),
            ).all()

            if transactions_unknown:
                ifscs_to_fetch = {t.ifsc_code for t in transactions_unknown}
                ifsc_to_state = {}
                with concurrent.futures.ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
                    future_to_ifsc = {executor.submit(get_state_from_api, ifsc): ifsc for ifsc in ifscs_to_fetch}
                    for future in concurrent.futures.as_completed(future_to_ifsc):
                        ifsc = future_to_ifsc[future]
                        try:
                            state = future.result()
                        except Exception as exc:
                            logger.exception(f"IFSC {ifsc} generated an exception: {exc}")
                            state = "Unknown"
                        ifsc_to_state[ifsc] = state

                # Save persistent cache
                save_ifsc_cache()

                # Update the database with the fetched states (normalized to title case)
                for t in transactions_unknown:
                    if t.ifsc_code in ifsc_to_state:
                        t.state = ifsc_to_state[t.ifsc_code].title()
                db.session.commit()

        # Now use database aggregation for efficiency
        from sqlalchemy import distinct, func

        state_summaries = (
            db.session.query(
                Transaction.state,
                func.count(Transaction.id).label("total_transactions"),
                func.sum(Transaction.amount).label("total_amount"),
                func.group_concat(distinct(Transaction.ifsc_code)).label("ifsc_codes"),
            )
            .filter(Transaction.ack_no == ack_no, Transaction.state.isnot(None), Transaction.state != "Unknown")
            .group_by(Transaction.state)
            .all()
        )

        # Define regions
        regions = {
            "Southern": [
                STATE_TAMIL_NADU,
                "Kerala",
                "Karnataka",
                STATE_ANDHRA_PRADESH,
                "Telangana",
                "Puducherry",
                "Lakshadweep",
                STATE_ANDAMAN_NICOBAR,
            ],
            "Western": ["Maharashtra", "Gujarat", "Rajasthan", "Goa", "Daman and Diu", "Dadra and Nagar Haveli"],
            "Eastern": [
                "West Bengal",
                "Odisha",
                "Bihar",
                "Jharkhand",
                "Assam",
                "Arunachal Pradesh",
                "Nagaland",
                "Manipur",
                "Mizoram",
                "Tripura",
                "Meghalaya",
                "Sikkim",
            ],
            "Northern": [
                "Jammu and Kashmir",
                "Himachal Pradesh",
                "Punjab",
                "Chandigarh",
                "Uttarakhand",
                "Haryana",
                "Delhi",
                "Uttar Pradesh",
                "Madhya Pradesh",
                "Chhattisgarh",
                "Ladakh",
            ],
        }

        # Convert to list
        summaries = []
        for summary in state_summaries:
            state, total_transactions, total_amount, ifsc_codes_str = summary
            ifsc_codes = sorted(ifsc_codes_str.split(",")) if ifsc_codes_str else []
            formatted_state = state.title() if state and state != "Unknown" else state
            summaries.append(
                {
                    "state": formatted_state,
                    "total_transactions": total_transactions,
                    "total_amount": float(total_amount) if total_amount else 0.0,
                    "ifsc_codes": ifsc_codes,
                }
            )

        # Group by regions
        regional_summaries = {region: [] for region in regions}
        other_summaries = []
        for summary in summaries:
            state = summary["state"]
            placed = False
            for region, states in regions.items():
                if state in states:
                    regional_summaries[region].append(summary)
                    placed = True
                    break
            if not placed:
                other_summaries.append(summary)

        # Sort within each region by total_amount descending, except Southern which uses custom order
        southern_order = [
            STATE_TAMIL_NADU,
            "Kerala",
            "Karnataka",
            STATE_ANDHRA_PRADESH,
            "Telangana",
            "Puducherry",
            "Lakshadweep",
            STATE_ANDAMAN_NICOBAR,
        ]
        for region in regional_summaries:
            if region == "Southern":
                # Sort Southern by custom order, then by amount descending for ties
                regional_summaries[region].sort(
                    key=lambda x: (
                        southern_order.index(x["state"]) if x["state"] in southern_order else len(southern_order),
                        -x["total_amount"],
                    )
                )
            else:
                regional_summaries[region].sort(key=lambda x: x["total_amount"], reverse=True)

        # Concatenate in order: Southern, Western, Eastern, Northern, then others
        result = (
            regional_summaries["Southern"]
            + regional_summaries["Western"]
            + regional_summaries["Eastern"]
            + regional_summaries["Northern"]
            + other_summaries
        )

        return jsonify(result)
    except Exception as e:
        logger.exception(f"Error in statewise_summary for {ack_no}: {e}")
        return jsonify({"error": MSG_INTERNAL_SERVER_ERROR}), 500


@app.route("/put_on_hold_transactions/<ack_no>", methods=["GET"])
@login_required
def put_on_hold_transactions(ack_no):
    """Return all put-on-hold transactions for a complaint."""
    check_case_access(ack_no)
    try:
        from ifsc_utils import get_ifsc_info

        hold_txns = Transaction.query.filter(
            Transaction.ack_no == ack_no.strip(), Transaction.put_on_hold_txn_id.isnot(None)
        ).all()

        # ✅ Sync with persistent POH details
        poh_map = {}
        try:
            poh_details_list = POHRefundDetails.query.filter_by(ack_no=ack_no.strip()).all()
            poh_map = {p.txn_id: p for p in poh_details_list}
        except Exception as e:
            logger.exception(f"Error fetching POH details in put_on_hold_transactions: {e}")

        # MRM 7-step progress, keyed by put_on_hold_txn_id (additive to legacy fields).
        mrm_map = {}
        try:
            mrm_map = {m.txn_id: m for m in MRMTracking.query.filter_by(ack_no=ack_no.strip()).all()}
        except Exception as e:
            logger.exception(f"Error fetching MRM details in put_on_hold_transactions: {e}")

        response = []
        for t in hold_txns:
            # Fetch branch info on backend to avoid N+1 frontend calls
            ifsc_data = get_ifsc_info(t.ifsc_code) or {}
            branch_name = ifsc_data.get("BRANCH") or ifsc_data.get("Branch") or "Unknown"

            # Get persistent details if available
            pdata = poh_map.get(t.put_on_hold_txn_id)
            court_date = pdata.court_order_date if pdata else t.court_order_date
            ref_status = pdata.refund_status if pdata else t.refund_status
            ref_amount = pdata.refund_amount if pdata else t.refund_amount

            response.append(
                {
                    "account_number": t.account_number or t.to_account,
                    "bank_name": t.bank_name,
                    "branch_name": branch_name,
                    "ifsc_code": t.ifsc_code,
                    "amount": t.put_on_hold_amount,
                    "layer": t.layer,
                    "court_order_date": court_date,
                    "refund_status": ref_status,
                    "refund_amount": ref_amount,
                    # txn_id lets the frontend address this hold row for MRM saves.
                    "hold_txn_id": t.put_on_hold_txn_id,
                    "mrm": _mrm_to_dict(mrm_map.get(t.put_on_hold_txn_id)),
                }
            )

        return jsonify(response)
    except Exception as e:
        logger.exception(f"Error fetching hold transactions for {_safe_log(ack_no)}: {e}")
        return jsonify({"error": MSG_INTERNAL_SERVER_ERROR}), 500


@app.route("/state_transactions/<ack_no>/<state>", methods=["GET"])
@login_required
def state_transactions(ack_no, state):
    check_case_access(ack_no)

    def _safe_int(v, default, lo, hi):
        try:
            n = int(v)
        except (TypeError, ValueError):
            return default
        return max(lo, min(n, hi))

    page = _safe_int(request.args.get("page"), 1, 1, 10_000_000)
    per_page = _safe_int(request.args.get("per_page"), 50, 1, 500)
    offset = (page - 1) * per_page

    # Use case-insensitive and trimmed comparison for state matching
    state_lower = state.strip().lower()

    # Get total count for this specific state
    total_count = Transaction.query.filter(
        Transaction.ack_no == ack_no, db.func.lower(db.func.trim(Transaction.state)) == state_lower
    ).count()

    # Query transactions for this state with pagination, sorting ATM transactions first
    transactions = (
        Transaction.query.filter(
            Transaction.ack_no == ack_no, db.func.lower(db.func.trim(Transaction.state)) == state_lower
        )
        .order_by(Transaction.atm_id.isnot(None).desc())
        .offset(offset)
        .limit(per_page)
        .all()
    )

    state_transactions = []
    for t in transactions:
        # Case-insensitive comparison for state matching
        if t.state and t.state.lower() == state.lower():
            # Determine transaction type
            if t.atm_id is not None:
                txn_type = "ATM Withdrawal"
                txn_amt = t.atm_withdraw_amount
                txn_id = "N/A"
            elif t.cheque_no is not None:
                txn_type = "Cheque Withdrawal"
                txn_amt = t.cheque_withdraw_amount
                txn_id = "N/A"
            elif t.put_on_hold_txn_id is not None:
                txn_type = "Put on Hold"
                txn_amt = t.put_on_hold_amount
                txn_id = "N/A"
            else:
                txn_type = "Account Transfer"
                txn_amt = str(t.amount)
                txn_id = t.txn_id

            state_transactions.append(
                {
                    "ack_no": t.ack_no,
                    "account_name": t.account_number,
                    "bank_name": t.bank_name,
                    "amount": txn_amt,
                    "ifsc_code": t.ifsc_code,
                    "date": t.txn_date or "N/A",
                    "transaction_type": txn_type,
                    #'status': t.action_taken or 'N/A',
                    "transaction_id": txn_id,
                    "layer": t.layer or "N/A",  # added by D
                    "is_atm": t.atm_id is not None,
                }
            )

    return jsonify(
        {
            "transactions": state_transactions,
            "total_count": total_count,
            "page": page,
            "per_page": per_page,
            "total_pages": (total_count + per_page - 1) // per_page,
        }
    )


# ---------------------------------------------------------------------------
# 9. Routes: KYC, hold/refund & letter generation
# ---------------------------------------------------------------------------


@app.route("/save_kyc", methods=["POST"])
@login_required
def save_kyc():
    if session.get("role") in VIEW_ONLY_ROLES:
        return jsonify({"status": "error", "message": "View-only users cannot edit KYC"}), 403
    data = request.get_json(silent=True) or {}  # malformed JSON -> {} instead of a 500
    txn_id = data.get("txn_id")

    if not txn_id:
        return jsonify({"status": "error", "message": "Transaction ID missing"}), 400

    # AuthZ (Bug #7 / IDOR): resolve the case for this txn_id and enforce per-officer
    # access before writing KYC. Checked before the try so abort(403) propagates.
    _txn_for_access = Transaction.query.filter_by(txn_id=txn_id).first()
    check_case_access(_txn_for_access.ack_no if _txn_for_access else None)

    try:
        # Save to persistent KYC store
        kyc_entry = KYCDetails.query.filter_by(txn_id=txn_id).first()
        if not kyc_entry:
            kyc_entry = KYCDetails(txn_id=txn_id)
            db.session.add(kyc_entry)

        # Sanitise once (strip control chars + cap to column lengths), apply to both stores
        name = sanitize_text(data.get("name"), 120)
        aadhar = sanitize_text(data.get("aadhar"), 20)
        mobile = sanitize_text(data.get("mobile"), 20)
        address = sanitize_text(data.get("address"), 200)

        # Format checks (blank allowed; reject clearly-malformed identifiers)
        if not validate_aadhar(aadhar):
            return jsonify({"status": "error", "message": "Aadhaar must be 12 digits."}), 400
        if not validate_mobile(mobile):
            return jsonify({"status": "error", "message": "Mobile must be a valid 10-digit number."}), 400

        kyc_entry.name = name
        kyc_entry.aadhar = aadhar
        kyc_entry.mobile = mobile
        kyc_entry.address = address

        # Also update the main Transaction table for backward compatibility
        txn = Transaction.query.filter_by(txn_id=txn_id).first()
        if txn:
            txn.kyc_name = name
            txn.kyc_aadhar = aadhar
            txn.kyc_mobile = mobile
            txn.kyc_address = address

        db.session.commit()
        log_usage("save_kyc", ack_no=(txn.ack_no if txn else None))
        logger.info(f"KYC Save Request processed for txn_id: {_safe_log(txn_id)}")
        return jsonify({"status": "success"})
    except Exception as e:
        logger.exception(f"Error saving KYC: {e}")
        db.session.rollback()
        return jsonify({"status": "error", "message": "Internal error while saving. See server log."}), 500


@app.route("/save_hold_refund", methods=["POST"])
@login_required
def save_hold_refund():
    """Save court order / refund details for a put-on-hold transaction."""
    if session.get("role") in VIEW_ONLY_ROLES:
        return jsonify({"status": "error", "message": "View-only users cannot edit refund details"}), 403

    data = request.get_json(silent=True) or {}
    ack_no = (data.get("ack_no") or "").strip()
    hold_txn_id = (data.get("hold_txn_id") or "").strip()

    if not ack_no or not hold_txn_id:
        logger.error("Missing required identifiers in save_hold_refund")
        return jsonify({"status": "error", "message": "Missing required identifiers"}), 400

    # AuthZ (Bug #7 / IDOR): only let users edit refund details for a case they may
    # access. Checked before the try so abort(403) isn't swallowed into a 500.
    check_case_access(ack_no)

    try:
        # (Trimmed log: do not dump the full payload — avoids logging sensitive data.)
        logger.info(f"save_hold_refund: ack_no={ack_no}, hold_txn_id={hold_txn_id}")
        txn = Transaction.query.filter_by(ack_no=ack_no, put_on_hold_txn_id=hold_txn_id).first()

        if not txn:
            logger.error(f"Put-on-hold transaction not found for ack={ack_no}, id={hold_txn_id}")
            return jsonify({"status": "error", "message": "Put-on-hold transaction not found"}), 404

        court_order_date = sanitize_text(data.get("court_order_date"), 20)
        refund_status = sanitize_text(data.get("refund_status"), 50)
        if not validate_court_order_date(court_order_date):
            return jsonify({"status": "error", "message": "Invalid court order date."}), 400
        if refund_status not in ALLOWED_REFUND_STATUSES:
            return jsonify({"status": "error", "message": "Invalid refund status."}), 400
        txn.court_order_date = court_order_date or None
        txn.refund_status = refund_status or None

        refund_amount_raw = data.get("refund_amount")
        try:
            txn.refund_amount = float(refund_amount_raw) if refund_amount_raw not in (None, "") else None
        except (TypeError, ValueError):
            logger.exception(f"Invalid refund amount: {_safe_log(refund_amount_raw)}")
            return jsonify({"status": "error", "message": "Invalid refund amount"}), 400

        # Update or Create POHRefundDetails for persistence
        poh_details = POHRefundDetails.query.filter_by(ack_no=ack_no, txn_id=hold_txn_id).first()
        if not poh_details:
            poh_details = POHRefundDetails(ack_no=ack_no, txn_id=hold_txn_id)
            db.session.add(poh_details)

        poh_details.court_order_date = txn.court_order_date
        poh_details.refund_status = txn.refund_status
        poh_details.refund_amount = txn.refund_amount

        db.session.commit()
        log_usage(
            "save_hold_refund",
            filename=f"status={txn.refund_status or 'cleared'}, amount={txn.refund_amount}",
            ack_no=ack_no,
        )
        logger.info("Refund details saved successfully (persisted to POHRefundDetails)")

        return jsonify({"status": "success"})
    except Exception as e:
        logger.exception(f"Error in save_hold_refund: {e}", exc_info=True)
        return jsonify({"status": "error", "message": "Internal error while saving. See server log."}), 500


def _mrm_to_dict(row, ack_no=None, txn_id=None):
    """Serialise an MRMTracking row (or None) into the shape the graph "Status of MRM"
    control and the read-only timeline consume: per-stage label/date/done, the refund
    info, the latest completed stage and the next actionable stage. When ack_no/txn_id
    are supplied, also include the who/what/when audit trail from MRMStatusLog."""
    steps = []
    latest = 0
    for i, label in enumerate(MRM_STEPS, start=1):
        date = getattr(row, f"step{i}") if row else None
        done = bool(date)
        if done:
            latest = i
        steps.append({"index": i, "label": label, "date": date or "", "done": done})
    out = {
        "steps": steps,
        "latest_step": latest,
        "latest_label": MRM_STEPS[latest - 1] if latest else "",
        "next_step": latest + 1 if latest < len(MRM_STEPS) else 0,
        "next_label": MRM_STEPS[latest] if latest < len(MRM_STEPS) else "",
        "refund_step": MRM_REFUND_STEP,
        "refund_type": (row.refund_type if row else None) or "",
        "refund_amount": (row.refund_amount if row else None),
    }
    if ack_no is not None and txn_id is not None:
        logs = (
            MRMStatusLog.query.filter_by(ack_no=str(ack_no).strip(), txn_id=str(txn_id).strip())
            .order_by(MRMStatusLog.step.asc(), MRMStatusLog.created_at.asc())
            .all()
        )
        out["audit"] = [
            {
                "step": g.step,
                "label": g.step_label,
                "date_completed": g.date_completed or "",
                "refund_type": g.refund_type or "",
                "refund_amount": g.refund_amount,
                "performed_by": g.performed_by or "",
                "performed_role": g.performed_role or "",
                "recorded_at": (to_ist(g.created_at).strftime("%Y-%m-%d %H:%M") + " IST") if g.created_at else "",
            }
            for g in logs
        ]
    return out


@app.route("/save_mrm_status", methods=["POST"])
@login_required
def save_mrm_status():
    """Record one step of the 7-step MRM workflow for a put-on-hold transaction.

    Enforces: sequential progress (step N needs step N-1 dated; step1 always allowed),
    set-once (a dated step cannot be changed), and that the refund step carries a
    refund_type (FULL/PARTIAL) and a positive amount. The refund step is mirrored
    into Transaction + POHRefundDetails so the refund dashboard and legacy graph
    fields reflect the same figure.
    """
    if session.get("role") in VIEW_ONLY_ROLES:
        return jsonify({"status": "error", "message": "View-only users cannot edit MRM status"}), 403

    data = request.get_json(silent=True) or {}
    ack_no = (data.get("ack_no") or "").strip()
    hold_txn_id = (data.get("hold_txn_id") or "").strip()
    if not ack_no or not hold_txn_id:
        return jsonify({"status": "error", "message": "Missing required identifiers"}), 400

    check_case_access(ack_no)

    try:
        step = int(data.get("step"))
    except (TypeError, ValueError):
        return jsonify({"status": "error", "message": "Invalid step"}), 400
    if step < 1 or step > len(MRM_STEPS):
        return jsonify({"status": "error", "message": "Step out of range"}), 400

    date = sanitize_text(data.get("date"), 20)
    if not date or not validate_court_order_date(date):
        return jsonify({"status": "error", "message": "A valid date is required for this step."}), 400

    refund_type = (data.get("refund_type") or "").strip().upper()
    refund_amount = None
    if step == MRM_REFUND_STEP:
        if refund_type not in ALLOWED_REFUND_TYPES:
            return jsonify({"status": "error", "message": "Refund type (FULL/PARTIAL) is required."}), 400
        try:
            refund_amount = float(data.get("refund_amount"))
        except (TypeError, ValueError):
            return jsonify({"status": "error", "message": "A valid refund amount is required."}), 400
        if refund_amount <= 0:
            return jsonify({"status": "error", "message": "Refund amount must be greater than zero."}), 400

    try:
        txn = Transaction.query.filter_by(ack_no=ack_no, put_on_hold_txn_id=hold_txn_id).first()
        if not txn:
            return jsonify({"status": "error", "message": "Put-on-hold transaction not found"}), 404

        row = MRMTracking.query.filter_by(ack_no=ack_no, txn_id=hold_txn_id).first()
        if not row:
            row = MRMTracking(ack_no=ack_no, txn_id=hold_txn_id)
            db.session.add(row)

        # Set-once: never overwrite a step that already has a date.
        if getattr(row, f"step{step}"):
            return jsonify({"status": "error", "message": f"Step {step} is already recorded."}), 409
        # Sequential: every step except the first needs its predecessor dated.
        if step > 1 and not getattr(row, f"step{step - 1}"):
            return jsonify({"status": "error", "message": f"Complete step {step - 1} before step {step}."}), 409

        setattr(row, f"step{step}", date)

        if step == MRM_REFUND_STEP:
            row.refund_type = refund_type
            row.refund_amount = refund_amount
            # Mirror into the durable refund stores so existing views stay correct.
            txn.refund_type = refund_type
            txn.refund_amount = refund_amount
            txn.refund_status = "Refunded" if refund_type == "FULL" else STATUS_PARTIALLY_REFUNDED
            poh = POHRefundDetails.query.filter_by(ack_no=ack_no, txn_id=hold_txn_id).first()
            if not poh:
                poh = POHRefundDetails(ack_no=ack_no, txn_id=hold_txn_id)
                db.session.add(poh)
            poh.court_order_date = txn.court_order_date
            poh.refund_status = txn.refund_status
            poh.refund_amount = refund_amount

        # Append the immutable audit entry: what / when / who.
        try:
            actor = current_user.username
            actor_role = current_user.role
        except Exception:
            actor = session.get("username")
            actor_role = session.get("role")
        db.session.add(
            MRMStatusLog(
                ack_no=ack_no,
                txn_id=hold_txn_id,
                step=step,
                step_label=MRM_STEPS[step - 1],
                date_completed=date,
                refund_type=refund_type if step == MRM_REFUND_STEP else None,
                refund_amount=refund_amount if step == MRM_REFUND_STEP else None,
                performed_by=actor,
                performed_role=actor_role,
            )
        )

        db.session.commit()
        log_usage("save_mrm_status", filename=f"step={step} ({MRM_STEPS[step - 1]})", ack_no=ack_no)
        return jsonify({"status": "success", "mrm": _mrm_to_dict(row, ack_no=ack_no, txn_id=hold_txn_id)})
    except Exception as e:
        try:
            db.session.rollback()
        except Exception:
            pass
        logger.exception(f"Error in save_mrm_status: {e}", exc_info=True)
        return jsonify({"status": "error", "message": "Internal error while saving MRM status."}), 500


@app.route("/mrm_timeline/<ack_no>/<txn_id>", methods=["GET"])
@login_required
def mrm_timeline(ack_no, txn_id):
    """Read-only MRM timeline for one put-on-hold transaction (Admin + IO).
    Access is gated to the case owner/assignee via check_case_access."""
    check_case_access(ack_no)
    row = MRMTracking.query.filter_by(ack_no=ack_no.strip(), txn_id=txn_id.strip()).first()
    return jsonify(_mrm_to_dict(row, ack_no=ack_no, txn_id=txn_id))


def format_indian_currency(amount):
    try:
        amount = float(amount)
    except (ValueError, TypeError):
        return "0.00"

    s = "{:.2f}".format(amount)
    parts = s.split(".")
    integer_part = parts[0]

    if len(integer_part) > 3:
        last_three = integer_part[-3:]
        rest = integer_part[:-3]
        rest_formatted = ""
        while len(rest) > 2:
            rest_formatted = "," + rest[-2:] + rest_formatted
            rest = rest[:-2]
        rest_formatted = rest + rest_formatted
        integer_part = rest_formatted + "," + last_three

    return f"₹{integer_part}"


def get_layer_1_total(ack_no):
    try:
        # Assuming Transaction model and db session are available globally
        l1_txns = Transaction.query.filter_by(ack_no=ack_no, layer=1).all()
        return sum(t.amount for t in l1_txns if t.amount) if l1_txns else 0.0
    except Exception as e:
        logger.exception(f"Error calculating layer 1 total: {e}")
        return 0.0


@app.route("/generate_letter", methods=["POST"])
@login_required
def generate_letter():
    """Generate a letter (HTML) from a template and provided details.
    Expects JSON with keys: ack_no, account_number, letter_type ('suspect'|'victim'),
    and officer details: officer_name, officer_designation, officer_phone, officer_email,
    letter_date, crime_no, ncrp_ack_no
    Returns rendered HTML for client to open/print.
    """
    # AuthZ (Bug #2 / IDOR): require login AND that the user may access this case.
    # check_case_access() runs BEFORE the try so its abort(403) is not swallowed by
    # the broad except below and turned into a 500.
    data = request.get_json(silent=True) or {}
    ack_no = data.get("ack_no")
    check_case_access(ack_no)
    try:
        account_number = data.get("account_number")
        letter_type = data.get("letter_type")

        # Collect officer/letter details with sensible defaults
        context = {
            "ack_no": ack_no,
            "account_number": account_number,
            "officer_name": data.get("officer_name", ""),
            "officer_designation": data.get("officer_designation", ""),
            "officer_phone": data.get("officer_phone", ""),
            "officer_email": data.get("officer_email", ""),
            "letter_date": data.get("letter_date", datetime.now().strftime("%d-%m-%Y")),
            "crime_no": data.get("crime_no", ""),
            "ncrp_ack_no": data.get("ncrp_ack_no", ack_no or ""),
            "layer_label": "Unknown layer",
            "layer_1_total": format_indian_currency(get_layer_1_total(ack_no)),
        }

        # Choose template
        if letter_type == "victim":
            template_name = "letter_victim.html"
        else:
            template_name = "letter_suspect.html"

        # Log usage
        try:
            log_usage("generate_letter", filename=template_name, ack_no=ack_no)
        except Exception:
            pass

        return render_template(template_name, **context)
    except Exception as e:
        logger.exception(f"Error generating letter: {e}")
        return jsonify({"error": MSG_DOC_GENERATION_ERROR}), 500


@app.route("/generate_letter_pdf", methods=["POST"])
@login_required
def generate_letter_pdf():
    """Generate a letter (PDF) from a template and provided details.
    Saves the PDF in a folder named after the ACK no.
    """
    # AuthZ (Bug #2 / IDOR): require login AND case access; checked before the try
    # so abort(403) isn't swallowed by the broad except.
    data = request.get_json(silent=True) or {}
    ack_no = data.get("ack_no")
    check_case_access(ack_no)
    try:
        account_number = data.get("account_number")
        letter_type = data.get("letter_type")

        # Collect officer/letter details with sensible defaults
        context = {
            "ack_no": ack_no,
            "account_number": account_number,
            "officer_name": data.get("officer_name", ""),
            "officer_designation": data.get("officer_designation", ""),
            "officer_phone": data.get("officer_phone", ""),
            "officer_email": data.get("officer_email", ""),
            "letter_date": data.get("letter_date", datetime.now().strftime("%d-%m-%Y")),
            "crime_no": data.get("crime_no", ""),
            "ncrp_ack_no": data.get("ncrp_ack_no", ack_no or ""),
            "layer_1_total": format_indian_currency(get_layer_1_total(ack_no)),
        }

        # Choose template
        if letter_type == "victim":
            template_name = "letter_victim.html"
        else:
            template_name = "letter_suspect.html"

        # Derive layer
        try:
            txns_for_layer = Transaction.query.filter(
                Transaction.ack_no == ack_no,
                or_(Transaction.to_account == account_number, Transaction.account_number == account_number),
            ).all()
            for t in txns_for_layer:
                if getattr(t, "layer", None):
                    context["layer_label"] = f"{ordinal(t.layer)} layer"
                    break
        except Exception:
            pass

        # Render HTML
        html_content = render_template(template_name, **context)

        # Create PDF
        pdf_buffer = io.BytesIO()
        pisa_status = pisa.CreatePDF(io.BytesIO(html_content.encode("utf-8")), dest=pdf_buffer)

        if pisa_status.err:
            return jsonify({"error": "PDF generation failed"}), 500

        # Create folder and save file
        folder_name = context["ncrp_ack_no"] or "Unknown_ACK"
        # Clean folder name to be safe
        folder_name = secure_filename(folder_name)

        # (Formerly built a generated_letters/<folder> path here — the PDF is
        #  streamed back to the browser below, with no server-side persistence.)
        # (No server-side persistence: the PDF is streamed back to the browser
        #  below. It used to be written under generated_letters/ on every request.)
        filename = f"{letter_type}_Letter_{account_number}.pdf"
        safe_filename_val = secure_filename(filename)

        # Reset buffer position for download (no disk write)
        pdf_buffer.seek(0)

        log_usage("download_letter_pdf", filename=safe_filename_val, ack_no=ack_no)
        return send_file(pdf_buffer, mimetype=MIME_PDF, as_attachment=True, download_name=safe_filename_val)

    except Exception as e:
        logger.exception(f"Error generating PDF letter: {e}")
        return jsonify({"error": MSG_DOC_GENERATION_ERROR}), 500


@app.route("/generate_letter_docx", methods=["POST"])
@login_required
def generate_letter_docx():
    """Generate a letter (DOCX) from a template and provided details.
    Saves the DOCX in a folder named after the ACK no.
    """
    # AuthZ (Bug #2 / IDOR): require login AND case access; checked before the try
    # so abort(403) isn't swallowed by the broad except.
    data = request.get_json(silent=True) or {}
    ack_no = data.get("ack_no")
    check_case_access(ack_no)
    try:
        # Support both single and multiple accounts
        account_numbers = data.get("account_numbers", [])
        l1_txns = Transaction.query.filter_by(ack_no=ack_no, layer=1).all()
        if not account_numbers and data.get("account_number"):
            account_numbers = [data.get("account_number")]

        if not account_numbers:
            return jsonify({"error": "No account numbers provided"}), 400

        letter_type = data.get("letter_type")
        is_poh = data.get("is_poh", False)

        # Validate the officer's contact details before they are written into an
        # official letter to a bank. A malformed email or phone in a legal notice is
        # not acceptable, so reject bad input (when provided) and normalise the phone.
        _officer_email = (data.get("officer_email") or "").strip()
        _officer_phone_raw = (data.get("officer_phone") or "").strip()
        if _officer_email and not re.match(r"^[^@\s]+@[^@\s]+\.[^@\s]+$", _officer_email):
            return jsonify({"error": "Please enter a valid email address for the officer."}), 400
        if _officer_phone_raw and _officer_phone_raw != "<Phone>":
            _phone_digits = re.sub(r"[\s\-()]", "", _officer_phone_raw)
            # Accept a bare 10-digit number by auto-prefixing +91, so a valid number
            # typed without the country code (or with spaces/dashes) isn't rejected.
            if re.match(r"^[6-9]\d{9}$", _phone_digits):
                _phone_digits = "+91" + _phone_digits
            if not re.match(r"^\+91[6-9]\d{9}$", _phone_digits):
                return jsonify(
                    {"error": "Please enter a valid 10-digit mobile number (optionally prefixed with +91)."}
                ), 400
            data["officer_phone"] = _phone_digits  # clean, normalised number flows into the letter
        is_l1_grouped = data.get("is_l1_grouped", False)

        # Collect details
        context = {
            "ack_no": ack_no,
            "officer_name": data.get("officer_name", "<Name>"),
            "officer_designation": data.get("officer_designation", "<Rank>"),
            "officer_phone": data.get("officer_phone", "<Phone>"),
            "officer_email": data.get("officer_email", "<email id>"),
            "letter_date": data.get("letter_date", datetime.now().strftime("%d-%m-%Y")),
            "crime_no": data.get("crime_no", "<Cr.No>"),
            "ncrp_ack_no": data.get("ncrp_ack_no", ack_no or "<Acknowledgement no>"),
        }

        # Calculate Layer 1 Total (Common for all letters in this ACK)
        layer_1_total = 0.0
        formatted_l1_total = "0.00"
        try:
            if l1_txns:
                layer_1_total = sum(t.amount for t in l1_txns if t.amount)

            # Helper for currency
            def format_indian_currency(amount):
                try:
                    amount = float(amount)
                except (ValueError, TypeError):
                    return "0.00"
                s = "{:.2f}".format(amount)
                parts = s.split(".")
                integer_part = parts[0]
                if len(integer_part) > 3:
                    last_three = integer_part[-3:]
                    rest = integer_part[:-3]
                    rest_formatted = ""
                    while len(rest) > 2:
                        rest_formatted = "," + rest[-2:] + rest_formatted
                        rest = rest[:-2]
                    rest_formatted = rest + rest_formatted
                    integer_part = rest_formatted + "," + last_three
                return f"₹{integer_part}"

            formatted_l1_total = format_indian_currency(layer_1_total)
        except Exception as e:
            logger.exception(f"Error calculating layer 1 total: {e}")

        # Prepare ZIP
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
            # If grouped, first fetch all Layer 1 txns and group them
            txns_groups = []
            all_txns = Transaction.query.filter_by(ack_no=ack_no).all()
            if is_l1_grouped:
                bmap = {}
                for t in l1_txns:
                    # Robust bank name normalization
                    bname = (t.bank_name or t.ifsc_code or BANK_UNKNOWN).strip().upper()
                    if bname not in bmap:
                        bmap[bname] = []
                    bmap[bname].append(t)
                txns_groups = list(bmap.values())
                logger.info(f"Grouped {len(l1_txns)} Layer 1 transactions into {len(txns_groups)} bank groups.")
            elif is_poh:
                txn_map = {str(t.txn_id): t for t in all_txns if t.txn_id}

                for acc in account_numbers:
                    poh_txns = [
                        t
                        for t in all_txns
                        if (str(t.account_number or "") == str(acc) or str(t.to_account or "") == str(acc))
                        and t.put_on_hold_txn_id
                    ]

                    seed_ids = {str(t.put_on_hold_txn_id) for t in poh_txns if t.put_on_hold_txn_id}

                    if not seed_ids:
                        continue

                    expanded_ids = set(seed_ids)
                    queue = list(seed_ids)

                    while queue:
                        cur_id = queue.pop(0)
                        cur_txn = txn_map.get(str(cur_id))
                        if not cur_txn:
                            continue

                        # upstream
                        parent_acc = cur_txn.from_account
                        if parent_acc:
                            parents = [
                                t
                                for t in all_txns
                                if (t.to_account == parent_acc or t.account_number == parent_acc) and t.txn_id
                            ]
                            for p in parents:
                                pid = str(p.txn_id)
                                if pid not in expanded_ids:
                                    expanded_ids.add(pid)
                                    queue.append(pid)

                    expanded_txns = [t for t in all_txns if str(t.txn_id) in expanded_ids]

                    if expanded_txns:
                        grouped = {}
                        for t in expanded_txns:
                            key = t.to_account or t.account_number or t.from_account or "UNKNOWN"
                            grouped.setdefault(key, []).append(t)

                        for _key, tx_list in grouped.items():
                            txns_groups.append({"transactions": tx_list, "folder": acc})

            else:
                # Path-to-Root / explicit account selection (e.g. the green
                # "Generate Letters (Path to Root)" button). The officer has picked
                # specific accounts, so generate a suspect-account letter for EACH
                # selected account FROM THAT ACCOUNT'S OWN real transactions. No
                # fabrication: an account with no transactions is skipped, and if none
                # qualify we return a clear message below instead of an empty ZIP.
                for acc in account_numbers:
                    acc_txns = [
                        t
                        for t in all_txns
                        if str(t.account_number or "") == str(acc)
                        or str(t.to_account or "") == str(acc)
                        or str(t.from_account or "") == str(acc)
                    ]
                    if acc_txns:
                        txns_groups.append({"transactions": acc_txns, "folder": str(acc)})

            # If nothing matched (e.g. selected accounts have no transactions, or a
            # Put-on-Hold request hit no held accounts) return a clear message rather
            # than streaming an empty ZIP — and never invent letters from unrelated data.
            if not txns_groups:
                return jsonify({
                    "error": "No accounts with transactions to generate letters for in this selection."
                }), 400

            # Helper to create and save a DOCX for a transactions list. Optionally place under a seed subfolder.
            def _save_doc_for_transactions(transactions, seed_folder=None):
                if not transactions:
                    return

                account_number_local = (
                    transactions[0].to_account
                    or transactions[0].account_number
                    or transactions[0].from_account
                    or "UNKNOWN"
                )

                bank_name = BANK_UNKNOWN
                amount_lost = 0.0
                from_date_local = context["letter_date"]

                for t in transactions:
                    if t.bank_name:
                        bank_name = t.bank_name
                        break
                    if t.ifsc_code:
                        bank_name = t.ifsc_code

                dates = []
                if is_poh:
                    amount_lost = sum(t.put_on_hold_amount for t in transactions if t.put_on_hold_amount)
                    dates = [t.put_on_hold_date for t in transactions if t.put_on_hold_date]
                else:
                    amount_lost = sum(t.amount for t in transactions if t.amount)
                    dates = [t.txn_date for t in transactions if t.txn_date]

                if dates:
                    try:
                        from_date_local = min(dates, key=lambda x: pd.to_datetime(x, dayfirst=True))
                    except Exception:
                        from_date_local = min(dates)

                layer_label = "Unknown layer"
                date_minus_6_str = from_date_local

                for t in transactions:
                    if getattr(t, "layer", None):
                        layer_label = f"{ordinal(t.layer)} layer"
                        break
                try:
                    dt_obj = pd.to_datetime(from_date_local, dayfirst=True)
                    new_month = dt_obj.month - 6
                    new_year = dt_obj.year
                    if new_month <= 0:
                        new_month += 12
                        new_year -= 1
                    dt_minus_6 = dt_obj.replace(year=new_year, month=new_month, day=1)
                    date_minus_6_str = dt_minus_6.strftime("%d-%m-%Y")
                except Exception:
                    date_minus_6_str = from_date_local

                if letter_type == "victim":
                    template_name_local = "letter_template_victim_account.docx"
                else:
                    template_name_local = "letter_template_suspect_accounts.docx"

                template_path_local = resource_path(template_name_local)
                if os.path.exists(template_path_local):
                    doc = Document(template_path_local)
                else:
                    doc = Document()
                    doc.add_paragraph("Template not found.")

                replacements = {
                    "<Bank Name>": bank_name,
                    "<amount lost>": f"{amount_lost:,.2f}",
                    "<total amount>": f"{amount_lost:,.2f}",
                    "(corresponding layer)": layer_label,
                    "<layer>": layer_label,
                    "<add all layer 1 amount>": formatted_l1_total,
                    "over all 1st layer total amount": formatted_l1_total,
                    "overall 1st layer total amount": formatted_l1_total,
                    "< Layer >": layer_label,
                    "<Layer>": layer_label,
                    "< layer >": layer_label,
                    "<transaction date - 6 months>": date_minus_6_str,
                    "<from date>": from_date_local,
                    "<Rank>": context["officer_designation"],
                    "<Name>": context["officer_name"],
                    "<Phno>": context["officer_phone"],
                    "<email id>": context["officer_email"],
                    "<Cr.No>": context["crime_no"],
                    "<Acknowledgement no>": context["ncrp_ack_no"],
                    "<Account No>": account_number_local,
                    "<To Date>": context["letter_date"],
                    "<current date>": context["letter_date"],
                    "<Suspect Account Number>": str(account_number_local),
                    "<Suspect account number>": str(account_number_local),
                    "<suspect account number>": str(account_number_local),
                    "<Suspect Account>": str(account_number_local),
                }

                def replace_in_paragraph(paragraph):
                    for key, val in replacements.items():
                        if key in paragraph.text:
                            for run in paragraph.runs:
                                if key in run.text:
                                    run.text = run.text.replace(key, str(val))
                    for key, val in replacements.items():
                        if key in paragraph.text:
                            paragraph.text = paragraph.text.replace(key, str(val))

                for paragraph in doc.paragraphs:
                    replace_in_paragraph(paragraph)
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                replace_in_paragraph(paragraph)

                target_text_start = "I am the investigating officer"
                new_paragraph_text = f"I am the investigating officer of the case mentioned in the subject. In this case, the victim has lost {formatted_l1_total}. Of the amount lost, partial amounts has been sent to the account mentioned below. Request you to provide the required details to proceed with the further investigation."
                for p in doc.paragraphs:
                    if p.text.strip().startswith(target_text_start):
                        p.clear()
                        run = p.add_run(new_paragraph_text)
                        run.font.name = "Bookman Old Style"
                        run.font.size = Pt(12)
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        break

                try:
                    heading_text = "Suspect Account Txn Details"
                    for p in doc.paragraphs:
                        t = p.text or ""
                        if heading_text in t:
                            p.text = f"({layer_label}) {heading_text}"
                            break
                except Exception:
                    pass

                target_table = None
                for table in doc.tables:
                    if len(table.rows) > 0:
                        header_cells = table.rows[0].cells
                        headers_text = [c.text.strip() for c in header_cells]
                        if (
                            LABEL_SUSPECT_ACCOUNT_NUMBER in headers_text
                            or "Transaction Id / UTR Number" in headers_text
                            or "Victim Account Number" in headers_text
                        ):
                            target_table = table
                            break

                if target_table:
                    tr = target_table.rows[0]._tr
                    tr_pr = tr.get_or_add_trPr()
                    tbl_header = OxmlElement("w:tblHeader")
                    tr_pr.append(tbl_header)
                    for i in range(len(target_table.rows) - 1, 0, -1):
                        row = target_table.rows[i]
                        tbl = row._element
                        tbl.getparent().remove(tbl)
                else:
                    logger.warning("Transaction table not found in template, searching for insertion point.")
                    target_idx = -1
                    headings_to_find = [
                        "Suspect Account Txn Details",
                        LABEL_SUSPECT_ACCOUNT_NUMBER,
                        "Suspect Account Details",
                    ]
                    for i, p in enumerate(doc.paragraphs):
                        p_text = p.text.strip()
                        if any(h in p_text for h in headings_to_find):
                            target_idx = i
                            break
                    if target_idx != -1:
                        target_table = doc.add_table(rows=1, cols=6)
                        target_table.style = "Table Grid"
                        tr = target_table.rows[0]._tr
                        tr_pr = tr.get_or_add_trPr()
                        tbl_header = OxmlElement("w:tblHeader")
                        tr_pr.append(tbl_header)
                        hdr_cells = target_table.rows[0].cells
                        headers = [
                            "S. No.",
                            LABEL_SUSPECT_ACCOUNT_NUMBER,
                            "Transaction Date",
                            "Transaction Amount",
                            "Transaction Id / UTR Number",
                            "IFSC Code",
                        ]
                        for i, h in enumerate(headers):
                            hdr_cells[i].text = h
                            for p in hdr_cells[i].paragraphs:
                                for r in p.runs:
                                    r.bold = True
                        tbl, p = target_table._tbl, doc.paragraphs[target_idx]._p
                        p.addnext(tbl)

                if target_table:
                    # "N/A" for genuinely-missing values so no cell is left blank.
                    def _v(x):
                        return str(x) if (x is not None and str(x).strip() != "") else "N/A"

                    for idx, t in enumerate(transactions, 1):
                        row = target_table.add_row().cells
                        row[0].text = str(idx)
                        if is_poh:
                            # For Put-on-Hold rows, fall back to the original transaction's
                            # date / amount / id when the hold-specific value is empty —
                            # that is why some letters were missing Date/Amount/Txn ID/UTR.
                            row[1].text = _v(t.account_number or t.to_account)
                            row[2].text = _v(t.put_on_hold_date or t.txn_date)
                            row[3].text = _v(t.put_on_hold_amount if t.put_on_hold_amount is not None else t.amount)
                            row[4].text = _v(t.put_on_hold_txn_id or t.txn_id)
                            row[5].text = _v(t.ifsc_code)
                        else:
                            row[1].text = _v(t.to_account or t.account_number or t.from_account)
                            row[2].text = _v(t.txn_date)
                            row[3].text = _v(t.amount)
                            row[4].text = _v(t.txn_id)
                            row[5].text = _v(t.ifsc_code)

                    # Render the table with visible cell borders ("boxes") whether it came
                    # from the template (which may be borderless) or was created above.
                    from docx.oxml.ns import qn as _qn

                    _tbl_el = target_table._tbl
                    _tbl_pr = _tbl_el.tblPr
                    if _tbl_pr is None:
                        _tbl_pr = OxmlElement("w:tblPr")
                        _tbl_el.insert(0, _tbl_pr)
                    _borders = OxmlElement("w:tblBorders")
                    for _edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                        _e = OxmlElement(f"w:{_edge}")
                        _e.set(_qn("w:val"), "single")
                        _e.set(_qn("w:sz"), "6")
                        _e.set(_qn("w:space"), "0")
                        _e.set(_qn("w:color"), "000000")
                        _borders.append(_e)
                    _tbl_pr.append(_borders)

                if is_poh or letter_type == "suspect":
                    subfolder_local = "suspect letter"
                elif letter_type == "victim":
                    subfolder_local = "victim letter"
                else:
                    subfolder_local = "other letters"

                # If seed_folder provided, place files directly under a folder named after the seed
                # i.e. generated_letters/<ACK>/<SEED_FOLDER>/<files>
                if seed_folder:
                    arc_subfolder = secure_filename(seed_folder)
                else:
                    arc_subfolder = subfolder_local

                # (No server-side persistence: the .docx is rendered into memory and
                #  added straight to the ZIP below — it used to accumulate forever
                #  under generated_letters/.)
                prefix_local = "Suspect_Account_Letter" if is_poh else letter_type + "_Letter"
                if is_l1_grouped:
                    safe_bank = secure_filename(bank_name) or "Grouped"
                    filename_local = f"{prefix_local}_L1_{safe_bank}_{secrets.token_hex(4)}.docx"
                else:
                    filename_local = f"{prefix_local}_{account_number_local}.docx"

                safe_filename_val_local = secure_filename(filename_local)

                # Render into memory and add directly to the ZIP (no temp file on disk).
                _doc_buf = io.BytesIO()
                doc.save(_doc_buf)
                arcname_local = os.path.join(arc_subfolder, safe_filename_val_local)
                zf.writestr(arcname_local, _doc_buf.getvalue())

            for item in txns_groups:
                if isinstance(item, dict):
                    _save_doc_for_transactions(item["transactions"], seed_folder=item.get("folder"))
                else:
                    _save_doc_for_transactions(item)

        zip_buffer.seek(0)

        zip_filename = f"Letters_{ack_no}.zip"
        if len(account_numbers) == 1:
            # Try to keep single file naming convention for download if possible, or just default to zip
            pass

        log_usage(
            "download_letters",
            filename=f"{len(account_numbers)} account(s); type={letter_type or 'suspect'}",
            ack_no=ack_no,
        )
        return send_file(zip_buffer, mimetype="application/zip", as_attachment=True, download_name=zip_filename)

    except Exception as e:
        logger.exception(f"Error generating DOCX letter: {e}")
        return jsonify({"error": MSG_DOC_GENERATION_ERROR}), 500


# ---------------------------------------------------------------------------
# 10. Routes: case & officer administration
# ---------------------------------------------------------------------------


@app.route("/view_all_complaints", methods=["GET"])
@login_required
def view_all_complaints():
    try:
        log_usage("view_all_complaints")
    except Exception as e:
        logger.exception(f"UsageLog view_all_complaints error: {e}")

    # OPTIMIZATION: Use defer to skip loading the large 'data' column
    complaints = UploadedFile.query.options(defer(UploadedFile.data)).order_by(UploadedFile.upload_time.desc()).all()

    # OPTIMIZATION: Query Transaction table directly for upload_id -> ack_no mapping
    # This avoids joining the heavy UploadedFile table and uses the index on upload_id
    upload_ack_rows = (
        db.session.query(Transaction.upload_id, Transaction.ack_no)
        .filter(Transaction.ack_no.isnot(None))
        .distinct()
        .all()
    )

    upload_id_to_acks = {}
    for uid, ack in upload_ack_rows:
        if not ack or not uid:
            continue
        upload_id_to_acks.setdefault(uid, set()).add(ack)

    # Convert upload_time to IST (UTC+5:30) and attach ACK numbers by upload_id
    for c in complaints:
        if c.upload_time:
            c.upload_time = c.upload_time.replace(tzinfo=timezone.utc).astimezone(
                timezone(timedelta(hours=5, minutes=30))
            )

        ack_set = upload_id_to_acks.get(c.id, set())
        c.ack_nos = sorted(ack_set) if ack_set else []
        # logger.debug(f"File: {c.filename}, ID: {c.id}, ACK numbers: {c.ack_nos}")

    # Deduplicate by ACK number (keep the latest upload per ACK based on ordering)
    seen_acks = set()
    unique_complaints = []
    for c in complaints:
        ack_value = c.ack_nos[0] if c.ack_nos else None
        if ack_value and ack_value in seen_acks:
            continue
        if ack_value:
            seen_acks.add(ack_value)
        unique_complaints.append(c)

    # Annotate each row with its current assignee; only expose officers in this
    # admin's group so the assign-dropdown stays scoped to the right group.
    officers = _officers_q().order_by(User.username).all()
    _complaint_rows = _cases_q().all()
    assign_map = {c.ack_no: c.assigned_to for c in _complaint_rows}
    status_map = {c.ack_no: (c.status or "Open") for c in _complaint_rows}
    id_to_username = {u.id: u.username for u in User.query.with_entities(User.id, User.username).all()}
    for c in unique_complaints:
        ack0 = c.ack_nos[0] if getattr(c, "ack_nos", None) else None
        c.assigned_to_id = assign_map.get(ack0)
        c.assigned_to_name = id_to_username.get(c.assigned_to_id)
        c.case_status = status_map.get(ack0, "Open")

    return render_template(
        "view_all_complaints.html", complaint_data=unique_complaints, officers=officers, statuses=CASE_STATUSES
    )


@app.route("/assign_case", methods=["POST"])
@login_required
@admin_required
def assign_case():
    """Admin assigns (or unassigns) a case to an Investigative Officer."""
    ack_no = (request.form.get("ack_no") or "").strip()
    officer_id_raw = (request.form.get("officer_id") or "").strip()

    if not ack_no:
        flash("Missing acknowledgement number.", "danger")
        return redirect(url_for("view_all_complaints"))

    complaint = _cases_q().filter_by(ack_no=ack_no).first()
    if not complaint:
        flash(f"No case found for ACK {ack_no}.", "warning")
        return redirect(url_for("view_all_complaints"))

    # Empty / "0" means unassign.
    if officer_id_raw in ("", "0"):
        complaint.assigned_to = None
        db.session.commit()
        log_usage("assign_case", ack_no=ack_no)
        flash(f"Case {ack_no} unassigned.", "success")
        return redirect(url_for("view_all_complaints"))

    try:
        officer_id = int(officer_id_raw)
    except ValueError:
        flash("Invalid officer selection.", "danger")
        return redirect(url_for("view_all_complaints"))

    # Ensure the target officer belongs to this admin's group.
    officer = _officers_q().filter_by(id=officer_id).first()
    if not officer:
        flash("Officer not found or not in your group.", "danger")
        return redirect(url_for("view_all_complaints"))

    complaint.assigned_to = officer.id
    db.session.commit()
    log_usage("assign_case", ack_no=ack_no, filename=officer.username)
    flash(f"Case {ack_no} assigned to {officer.username}.", "success")
    return redirect(url_for("view_all_complaints"))


@app.route("/view_officers", methods=["GET"])
@login_required
@admin_required
def view_officers():
    # Scoped to this admin's group (SuperAdmin sees all officers).
    officers = _officers_q().all()

    # For each officer, count uploaded files
    officer_data = []
    for officer in officers:
        computed_upload_count = UploadedFile.query.filter_by(uploader=officer.username).count()
        upload_count = officer.manual_upload_count if officer.manual_upload_count is not None else computed_upload_count
        officer_data.append(
            {
                "username": officer.username,
                "role": officer.role,
                "upload_count": upload_count,
                "computed_upload_count": computed_upload_count,
                "manual_upload_count": officer.manual_upload_count,
                "name": officer.name,
                "rank": officer.rank,
                "email": officer.email,
            }
        )

    return render_template("view_officers.html", officers=officer_data)


@app.route("/update_officer", methods=["POST"])
@login_required
@admin_required
def update_officer():
    username = (request.form.get("username") or "").strip()
    if not username:
        flash("Invalid request.")
        return redirect(url_for("view_officers"))

    user = User.query.filter_by(username=username).first()
    if not user:
        flash("Officer not found.")
        return redirect(url_for("view_officers"))

    # Only allow editing investigative officers via this view
    if user.role != ROLE_INVESTIGATIVE_OFFICER:
        flash("Only Investigative Officers can be edited here.")
        return redirect(url_for("view_officers"))

    # Update editable fields (keep it minimal)
    user.name = sanitize_text(request.form.get("name"), 100) or None
    user.rank = sanitize_text(request.form.get("rank"), 100) or None
    user.email = sanitize_text(request.form.get("email"), 120) or None

    manual_upload_count_raw = (request.form.get("manual_upload_count") or "").strip()
    if manual_upload_count_raw == "":
        user.manual_upload_count = None
    else:
        try:
            user.manual_upload_count = int(manual_upload_count_raw)
        except ValueError:
            flash("No. of Files Uploaded must be a number (or leave blank).")
            return redirect(url_for("view_officers"))

    db.session.commit()
    log_usage("update_officer", filename=username)
    flash(f"Officer {username} updated successfully.")
    return redirect(url_for("view_officers"))


@app.route("/edit_officer/<int:officer_id>", methods=["POST"])
@login_required
def edit_officer(officer_id):
    if not is_admin():
        return jsonify({"error": "Access denied"}), 403

    officer = _officers_q().filter_by(id=officer_id).first()
    if not officer:
        return jsonify({"error": "Officer not found or not in your group"}), 404

    password = request.json.get("password")
    if password:
        try:
            officer.set_password(password)
            db.session.commit()
            log_usage("edit_officer_password", filename=officer.username)
            return jsonify({"message": "Password updated successfully"})
        except ValueError as e:
            return jsonify({"error": str(e)}), 400

    return jsonify({"error": "No password provided"}), 400


@app.route("/delete_officer", methods=["POST"])
@login_required
@admin_required
def delete_officer():
    username = request.form.get("username")
    if not username:
        flash("Invalid request.")
        return redirect(url_for("view_officers"))

    user = User.query.filter_by(username=username).first()
    if not user:
        flash("Officer not found.")
        return redirect(url_for("view_officers"))

    if user.role == "Admin":
        flash("Cannot delete admin user.")
        return redirect(url_for("view_officers"))

    db.session.delete(user)
    db.session.commit()
    log_usage("delete_officer", filename=username)
    flash(f"Officer {username} deleted successfully.")
    return redirect(url_for("view_officers"))


# ─────────────────────────────────────────────────────────────────
#  File Management — view / download / delete uploaded Excel files
# ─────────────────────────────────────────────────────────────────
@app.route("/manage_files", methods=["GET"])
@login_required
@admin_required
def manage_files():
    """List all uploaded Excel files with per-file stats."""
    try:
        log_usage("manage_files")
    except Exception:
        pass
    files = UploadedFile.query.options(defer(UploadedFile.data)).order_by(UploadedFile.upload_time.desc()).all()
    file_stats = []
    for f in files:
        txns = Transaction.query.filter_by(upload_id=f.id).all()
        ack_nos = list({t.ack_no for t in txns if t.ack_no})
        total_amt = sum(t.amount or 0 for t in txns)
        disputed_amt = sum(t.disputed_amount or 0 for t in txns)
        hold_count = sum(1 for t in txns if t.put_on_hold_txn_id)
        hold_amt = sum(t.put_on_hold_amount or 0 for t in txns if t.put_on_hold_txn_id)
        layers = sorted({t.layer for t in txns if t.layer})
        file_stats.append(
            {
                "id": f.id,
                "filename": f.filename,
                "uploader": f.uploader or "—",
                "upload_time": (to_ist(f.upload_time).strftime(DATETIME_DISPLAY_FORMAT) if f.upload_time else "—"),
                "txn_count": len(txns),
                "ack_nos": ack_nos,
                "total_amt": total_amt,
                "disputed_amt": disputed_amt,
                "hold_count": hold_count,
                "hold_amt": hold_amt,
                "layers": layers,
            }
        )
    return render_template("manage_files.html", file_stats=file_stats)


@app.route("/delete_upload/<int:file_id>", methods=["POST"])
@login_required
@admin_required
def delete_upload(file_id):
    """Delete an uploaded file and all its transactions."""
    f = UploadedFile.query.get_or_404(file_id)
    fname = f.filename
    try:
        # Delete transactions linked to this upload
        Transaction.query.filter_by(upload_id=file_id).delete(synchronize_session="fetch")
        db.session.delete(f)
        db.session.commit()
        # Remove disk copy if it exists
        disk_path = os.path.join(app.config.get("UPLOAD_FOLDER", ""), fname)
        if os.path.exists(disk_path):
            os.remove(disk_path)
        flash(f'File "{fname}" and all its transactions deleted.', "success")
        log_usage("delete_upload", filename=fname)
    except Exception as e:
        db.session.rollback()
        logger.exception(f"delete_upload error for id={file_id}: {e}")
        flash("Error deleting file — details are in the server log.", "danger")
    return redirect(url_for("manage_files"))


@app.route("/download_upload/<int:file_id>", methods=["GET"])
@login_required
@admin_required
def download_upload(file_id):
    """Stream the original Excel file back to the browser."""
    f = UploadedFile.query.get_or_404(file_id)
    if not f.data:
        flash("File data not found in database.", "danger")
        return redirect(url_for("manage_files"))
    log_usage("download_upload", filename=f.filename)
    return send_file(
        io.BytesIO(f.data),
        mimetype=f.mimetype or MIME_XLSX,
        as_attachment=True,
        download_name=f.filename,
    )


@app.route("/download_case_excel/<ack_no>", methods=["GET"])
@login_required
def download_case_excel(ack_no):
    """Download the original uploaded Excel for a case — available to the owning
    officer AND to admins (case access enforced). Keyed by acknowledgement number,
    so it works from the graph toolbar without needing the internal file id."""
    ack_no = str(ack_no).strip()
    check_case_access(ack_no)
    row = (
        db.session.query(UploadedFile)
        .join(Transaction, Transaction.upload_id == UploadedFile.id)
        .filter(Transaction.ack_no == ack_no)
        .order_by(UploadedFile.upload_time.desc())
        .first()
    )
    if not row or not row.data:
        flash("No uploaded Excel found for this case.", "warning")
        return redirect(url_for("role_home"))
    log_usage("download_case_excel", filename=row.filename, ack_no=ack_no)
    return send_file(
        io.BytesIO(row.data),
        mimetype=row.mimetype or MIME_XLSX,
        as_attachment=True,
        download_name=row.filename,
    )


@app.route("/view_analytics", methods=["GET"])
@login_required
@admin_required
def view_analytics():
    try:
        log_usage("view_analytics")
    except Exception as e:
        logger.exception(f"UsageLog view_analytics error: {e}")

    # ── Summary KPIs ──────────────────────────────────────────────────
    total_files = db.session.query(func.count(func.distinct(UploadedFile.filename))).scalar() or 0
    total_txns = db.session.query(func.count(Transaction.id)).scalar() or 0
    total_amount = float(db.session.query(func.sum(Transaction.amount)).scalar() or 0)
    disputed_amt = float(db.session.query(func.sum(Transaction.disputed_amount)).scalar() or 0)
    hold_amt = float(db.session.query(func.sum(Transaction.put_on_hold_amount)).scalar() or 0)
    hold_count = (
        db.session.query(func.count(Transaction.id)).filter(Transaction.put_on_hold_txn_id.isnot(None)).scalar() or 0
    )
    unique_accts = db.session.query(func.count(func.distinct(Transaction.to_account))).scalar() or 0
    unique_banks = db.session.query(func.count(func.distinct(Transaction.bank_name))).scalar() or 0
    unique_states = (
        db.session.query(func.count(func.distinct(Transaction.state))).filter(Transaction.state.isnot(None)).scalar()
        or 0
    )
    max_layer = db.session.query(func.max(Transaction.layer)).scalar() or 0
    active_officers = db.session.query(func.count(func.distinct(UploadedFile.uploader))).scalar() or 0

    # ── Per-file breakdown (all important numbers) ────────────────────
    files = UploadedFile.query.options(defer(UploadedFile.data)).order_by(UploadedFile.upload_time.desc()).all()
    file_rows = []
    for f in files:
        txns = Transaction.query.filter_by(upload_id=f.id).all()
        ack_nos = list({t.ack_no for t in txns if t.ack_no})
        fa = sum(t.amount or 0 for t in txns)
        fd = sum(t.disputed_amount or 0 for t in txns)
        fh = sum(t.put_on_hold_amount or 0 for t in txns)
        fhc = sum(1 for t in txns if t.put_on_hold_txn_id)
        file_rows.append(
            {
                "filename": f.filename,
                "uploader": f.uploader or "—",
                "upload_time": (to_ist(f.upload_time).strftime(DATETIME_DISPLAY_FORMAT) if f.upload_time else "—"),
                "txn_count": len(txns),
                "ack_count": len(ack_nos),
                "total_amt": fa,
                "disputed_amt": fd,
                "hold_amt": fh,
                "hold_count": fhc,
            }
        )

    # ── Layer-wise distribution ───────────────────────────────────────
    layer_dist = (
        db.session.query(Transaction.layer, func.count(Transaction.id), func.sum(Transaction.amount))
        .group_by(Transaction.layer)
        .order_by(Transaction.layer)
        .all()
    )

    # ── Top banks by transaction volume ──────────────────────────────
    top_banks = (
        db.session.query(
            Transaction.bank_name, func.count(Transaction.id).label("cnt"), func.sum(Transaction.amount).label("amt")
        )
        .filter(Transaction.bank_name.isnot(None))
        .group_by(Transaction.bank_name)
        .order_by(desc("cnt"))
        .limit(10)
        .all()
    )

    # ── Top IFSC codes ────────────────────────────────────────────────
    top_ifsc = (
        db.session.query(Transaction.ifsc_code, Transaction.bank_name, func.count(Transaction.id).label("cnt"))
        .filter(Transaction.ifsc_code.isnot(None))
        .group_by(Transaction.ifsc_code, Transaction.bank_name)
        .order_by(desc("cnt"))
        .limit(10)
        .all()
    )

    # ── State-wise distribution ───────────────────────────────────────
    # Derive state from each IFSC via the local memoised IFSC table, NOT the cached
    # Transaction.state column — that column is only filled lazily (when a case's
    # state-wise summary is opened), so it was mostly NULL and this table showed empty.
    # Aggregate by unique IFSC -> one O(1) lookup each.
    from ifsc_utils import get_ifsc_info as _gii

    _ifsc_rows = (
        db.session.query(
            Transaction.ifsc_code, func.count(Transaction.id).label("cnt"), func.sum(Transaction.amount).label("amt")
        )
        .filter(Transaction.ifsc_code.isnot(None))
        .group_by(Transaction.ifsc_code)
        .all()
    )
    _state_agg = defaultdict(lambda: {"cnt": 0, "amt": 0.0})
    for _ifsc, _cnt, _amt in _ifsc_rows:
        _info = _gii(_ifsc) or {}
        _st = (str(_info.get("STATE") or _info.get("State") or "").strip().title()) or "Unknown"
        _state_agg[_st]["cnt"] += _cnt
        _state_agg[_st]["amt"] += _amt or 0
    state_dist = sorted(_state_agg.items(), key=lambda kv: -kv[1]["cnt"])[:15]
    state_dist = [(_s, _v["cnt"], _v["amt"]) for _s, _v in state_dist]

    # ── Officer upload stats ──────────────────────────────────────────
    officer_stats = (
        db.session.query(User.username, User.name, func.count(UploadedFile.id).label("uploads"))
        .join(UploadedFile, UploadedFile.uploader == User.username)
        .filter(User.role == ROLE_INVESTIGATIVE_OFFICER)
        .group_by(User.username, User.name)
        .order_by(desc("uploads"))
        .all()
    )

    # ── ATM / Cheque withdrawal counts ───────────────────────────────
    atm_count = db.session.query(func.count(Transaction.id)).filter(Transaction.atm_id.isnot(None)).scalar() or 0
    cheque_count = db.session.query(func.count(Transaction.id)).filter(Transaction.cheque_no.isnot(None)).scalar() or 0

    return render_template(
        "view_analytics.html",
        # KPIs
        total_files=total_files,
        total_txns=total_txns,
        total_amount=total_amount,
        disputed_amt=disputed_amt,
        hold_amt=hold_amt,
        hold_count=hold_count,
        unique_accts=unique_accts,
        unique_banks=unique_banks,
        unique_states=unique_states,
        max_layer=max_layer,
        active_officers=active_officers,
        atm_count=atm_count,
        cheque_count=cheque_count,
        # Tables
        file_rows=file_rows,
        layer_dist=layer_dist,
        top_banks=top_banks,
        top_ifsc=top_ifsc,
        state_dist=state_dist,
        officer_stats=officer_stats,
    )


@app.route("/download_logs", methods=["GET"])
@login_required
@admin_required
def download_logs():
    try:
        now = datetime.now(timezone.utc)
        start_of_day = now.replace(hour=0, minute=0, second=0, microsecond=0)
        end_of_day = start_of_day + timedelta(days=1)
        logs = (
            UsageLog.query.filter(UsageLog.timestamp >= start_of_day, UsageLog.timestamp < end_of_day)
            .order_by(UsageLog.timestamp.asc())
            .all()
        )

        # PDF Generation
        buf = io.BytesIO()
        c = canvas.Canvas(buf, pagesize=landscape(A4))
        _, height = landscape(A4)
        y = height - 50

        # Title
        c.setFont("Helvetica-Bold", 16)
        c.drawString(
            30,
            y,
            f"Usage Logs - {start_of_day.astimezone(timezone(timedelta(hours=5, minutes=30))).strftime('%Y-%m-%d')}",
        )
        y -= 30

        # Headers - Adjusted for Landscape (Width ~842)
        c.setFont("Helvetica-Bold", 10)
        c.drawString(30, y, "Timestamp")  # Left aligned
        c.drawString(180, y, "User (Role)")  # ~150px space
        c.drawString(350, y, "Action")  # ~150px space
        c.drawString(500, y, "Details (File/Ack)")  # Rest of the space (~310px)
        y -= 20
        c.line(30, y + 15, 810, y + 15)  # Line spans full width

        c.setFont("Helvetica", 9)

        for l in logs:
            if y < 50:
                c.showPage()
                y = height - 50
                # Redraw headers on new page
                c.setFont("Helvetica-Bold", 10)
                c.drawString(30, y, "Timestamp")
                c.drawString(180, y, "User (Role)")
                c.drawString(350, y, "Action")
                c.drawString(500, y, "Details (File/Ack)")
                y -= 20
                c.line(30, y + 15, 810, y + 15)
                c.setFont("Helvetica", 9)

            ts_ist = l.timestamp.replace(tzinfo=timezone.utc).astimezone(timezone(timedelta(hours=5, minutes=30)))
            ts_str = ts_ist.strftime("%Y-%m-%d %H:%M:%S")
            u = l.username or ""
            r = l.role or ""
            a = l.action or ""
            f = l.filename or ""
            an = l.ack_no or ""

            details = f"{f} {an}".strip()
            # No truncation

            user_role = f"{u} ({r})"
            if len(user_role) > 35:
                user_role = user_role[:32] + "..."

            c.drawString(30, y, ts_str)
            c.drawString(180, y, user_role)
            c.drawString(350, y, a)
            c.drawString(500, y, details)

            y -= 15

        c.save()
        buf.seek(0)

        fname = (
            f"usage_logs_{start_of_day.astimezone(timezone(timedelta(hours=5, minutes=30))).strftime('%Y-%m-%d')}.pdf"
        )
        log_usage("download_logs")
        return send_file(buf, mimetype=MIME_PDF, as_attachment=True, download_name=fname)
    except Exception as e:
        logger.exception(f"Failed to generate logs PDF: {e}")
        return "Failed to generate logs.", 500


@app.route("/delete_complaint/<int:complaint_id>", methods=["DELETE"])
@login_required
def delete_complaint(complaint_id):
    complaint = Complaint.query.get_or_404(complaint_id)

    if current_user.role != "Admin":
        abort(403, description="You don't have permission to delete this complaint")
    # Group isolation: non-SuperAdmin may only delete their own group's cases.
    if not is_superadmin() and complaint.owner_admin_id not in (None, current_user.id):
        abort(403, description="This case belongs to a different admin group")

    ack_no = complaint.ack_no
    # Delete the case's transactions too (previously orphaned), matching delete_by_ack.
    Transaction.query.filter_by(ack_no=ack_no).delete(synchronize_session="fetch")
    db.session.delete(complaint)
    db.session.commit()
    logger.info(f"User {current_user.username} deleted complaint {complaint_id} (ack {ack_no})")
    log_usage("delete_complaint", ack_no=ack_no)
    return jsonify({"success": True})


@app.route("/delete_by_ack", methods=["POST"])
@login_required
def delete_by_ack():
    if current_user.role != "Admin":
        flash("Unauthorized action.", "danger")
        return redirect(url_for("admin_dashboard"))

    ack_no = request.form.get("ack_no", "").strip()
    if not ack_no:
        flash("Please provide an Acknowledgement Number.", "warning")
        return redirect(url_for("admin_dashboard"))

    try:
        # Group isolation: non-SuperAdmin admins can only delete their own cases.
        if not is_superadmin():
            _c = Complaint.query.filter_by(ack_no=ack_no).first()
            if _c and _c.owner_admin_id not in (None, current_user.id):
                flash("You can only delete cases in your own group.", "danger")
                return redirect(url_for("admin_dashboard"))

        # Delete from Transaction table
        deleted_count = Transaction.query.filter_by(ack_no=ack_no).delete()

        # Also delete from Complaint table if it exists there
        complaint_deleted = Complaint.query.filter_by(ack_no=ack_no).delete()

        db.session.commit()

        if deleted_count > 0 or complaint_deleted > 0:
            flash(f"Successfully deleted records for ACK No: {ack_no}", "success")
            logger.info(
                f"User {current_user.username} deleted ACK {ack_no}. Txns: {deleted_count}, Complaints: {complaint_deleted}"
            )
        else:
            flash(f"No records found for ACK No: {ack_no}", "warning")

    except Exception as e:
        db.session.rollback()
        logger.exception(f"Error deleting ACK {ack_no}: {e}")
        flash("An error occurred while deleting the records.", "danger")

    return redirect(url_for("admin_dashboard"))


@app.route("/view_complaint/<int:complaint_id>", methods=["GET"])
@login_required
def view_complaint(complaint_id):
    complaint = Complaint.query.get_or_404(complaint_id)

    # Authorization: admins restricted to their own group; officers to assigned/owned cases.
    if current_user.role == "Admin":
        if not is_superadmin() and complaint.owner_admin_id not in (None, current_user.id):
            abort(403, description="This case belongs to a different admin group")
    elif (
        complaint.assigned_to != current_user.id
        and complaint.uploaded_by != current_user.id
    ):
        abort(403, description="You don't have permission to view this complaint")

    # Viewing a case = viewing its fund-trail graph.
    return redirect(url_for("graph_tree1", ack_no=complaint.ack_no))


@app.route("/admin/add_officer", methods=["GET"])
@login_required
@admin_required
def add_officer():
    # This page only shows the form. Submitting it POSTs to /submit_officer, which
    # validates the input, creates the officer, and re-renders this page with an
    # inline error message if anything is wrong.
    return render_template("add_officer.html")


@app.route("/submit_officer", methods=["POST"])
@login_required
def submit_officer():
    if not is_admin():
        flash("Access Denied!", "danger")
        return redirect(url_for("login"))

    name = sanitize_text(request.form.get("name"), 100)
    rank = sanitize_text(request.form.get("rank"), 100)
    email = sanitize_text(request.form.get("email"), 120)
    username = request.form.get("username", "").strip()[:80]  # cap length; keep .strip() so it matches login
    # Do NOT strip the password: it is stored exactly as typed and must match what
    # the officer enters at login (the login route does not strip it either).
    password = request.form.get("password", "")

    # On any problem, re-render the SAME form with a clear inline error and keep
    # what was already typed. (The old code redirected to '/add_officer' — a route
    # that does not exist; the real one is '/admin/add_officer' — so a weak password
    # produced a 404 instead of a helpful message.)
    def _reject(message):
        return render_template(
            "add_officer.html",
            error=message,
            form={"name": name, "rank": rank, "email": email, "username": username},
        )

    if not (name and rank and email and username and password):
        return _reject("All fields are required.")

    if User.query.filter_by(username=username).first():
        return _reject(f"The username '{username}' is already taken. Please choose another.")

    new_officer = User(
        username=username,
        role=ROLE_INVESTIGATIVE_OFFICER,
        name=name,
        rank=rank,
        email=email,
        admin_id=current_user.id,  # officer belongs to the admin who created them
    )
    try:
        # set_password enforces the complexity policy and raises ValueError naming
        # the exact rule that failed (e.g. "Password must be at least 12 characters").
        new_officer.set_password(password)
    except ValueError as e:
        return _reject(str(e))

    # This password is TEMPORARY. Force the officer to set their own on first login:
    # must_change_password makes the login flow redirect them straight to /change_password.
    new_officer.must_change_password = True
    db.session.add(new_officer)
    db.session.commit()

    flash(f"Officer '{username}' added successfully. They must set a new password on first login.", "success")
    return redirect(url_for("view_officers"))


with app.app_context():
    db.create_all()

    # NOTE: this import-time seeding runs once per process. Multi-worker servers
    # must import the app exactly once before forking (gunicorn --preload, as the
    # Dockerfile does) or parallel workers race to insert the same first users.
    if User.query.count() == 0:
        # Generate random initial passwords that meet complexity requirements
        # uses the module-level generate_secure_password() defined earlier (de-duplicated)
        # Random per-machine first-login passwords (keeps the random-password hardening). They are
        # printed below AND saved to INITIAL_CREDENTIALS.txt so they're always easy to
        # find, and must be changed on first login.
        admin_password = generate_secure_password()
        officer_password = generate_secure_password()

        admin = User(username="admin", role="Admin", is_superadmin=True)
        admin.set_password(admin_password)
        admin.must_change_password = True

        db.session.add(admin)
        db.session.flush()  # assign admin.id before referencing it below

        officer = User(
            username="officer",
            role=ROLE_INVESTIGATIVE_OFFICER,
            admin_id=admin.id,  # officer belongs to the seed admin group
        )
        officer.set_password(officer_password)
        officer.must_change_password = True

        db.session.add(officer)
        db.session.commit()

        logger.info("Users initialized with secure random passwords.")

        # Save the generated passwords to a gitignored file as well as printing them.
        # The console message scrolls away the moment the app starts logging requests,
        # which is exactly why the admin password kept getting "lost" after a clone.
        try:
            creds_path = os.path.join(secure_base, INITIAL_CREDENTIALS_FILENAME)
            with open(creds_path, "w") as cf:
                cf.write(
                    "FundTrail initial login credentials (auto-generated)\n"
                    "Change these on first login, then delete this file.\n\n"
                    f"  admin    password: {admin_password}\n"
                    f"  officer  password: {officer_password}\n"
                )
            if _is_posix():
                os.chmod(creds_path, 0o600)
            print(f"[SETUP] Credentials also saved to: {creds_path}")
        except Exception as exc:
            logger.warning(f"Could not write INITIAL_CREDENTIALS.txt: {exc}")

        print(f"[SETUP] Initial admin password: {admin_password}")
        print(f"[SETUP] Initial officer password: {officer_password}")
        print("[SETUP] Change these passwords immediately!")

    # Consolidate the old split POH/KYC SQLite files into the main DB (one-time).
    migrate_split_dbs_into_main()
    # Ensure pre-existing uploads have Complaint rows for per-officer isolation.
    backfill_complaints()


# Configure custom error handlers
# ---------------------------------------------------------------------------
# 11. Error handlers
# ---------------------------------------------------------------------------


@app.errorhandler(CSRFError)
def handle_csrf_error(error):
    # Flask-WTF normally returns a raw "400 Bad Request: The CSRF token is missing".
    # That's what users were seeing on the Change/Reset Password form when their
    # session had expired or cookies weren't being stored. Show a friendly page that
    # explains it and lets them try again.
    return render_template("errors/csrf.html"), 400


@app.errorhandler(404)
def not_found(error):
    return render_template("errors/404.html"), 404


@app.errorhandler(403)
def forbidden(error):
    return render_template("errors/403.html"), 403


@app.errorhandler(500)
def internal_error(error):
    logger.error(MSG_INTERNAL_SERVER_ERROR, exc_info=True)
    # Mirror into alerts.log so failures are noticed, not buried.
    try:
        alert_logger.error(
            f"500 on {request.method} {request.path} "
            f"(user={session.get('username', 'anon')}, rid={getattr(g, 'request_id', '-')})"
        )
    except Exception:
        pass
    return render_template("errors/500.html"), 500


@app.errorhandler(413)
def request_too_large(error):
    # Triggered by MAX_CONTENT_LENGTH — keep it simple and clear.
    response = make_response("The request was too large. Uploads are limited to 25 MB.", 413)
    response.headers["Content-Type"] = "text/plain; charset=utf-8"
    return response


# ---------------------------------------------------------------------------
# 12. Routes: reports, analytics & exports
# ---------------------------------------------------------------------------


@app.route("/download_fundtrail_pdf", methods=["POST"])
@login_required
def download_fundtrail_pdf():
    """
    Generate a PDF visualising the fund trail using reportlab.
    Receives JSON data containing the chain of nodes.
    """
    try:
        data = request.get_json() or {}
        ack_no = data.get("ack_no", "Unknown")
        nodes = data.get("nodes", [])

        if not nodes:
            return jsonify({"error": "No data provided"}), 400

        # Verify the caller is authorised to access this case before generating PDF.
        if ack_no and ack_no != "Unknown":
            check_case_access(ack_no)

        buffer = io.BytesIO()
        c = canvas.Canvas(buffer, pagesize=A4)
        width, height = A4

        # ---------- neutral palette ----------
        NAVY = colors.HexColor("#0a2e63")
        GOLD = colors.HexColor("#f5b301")
        INK = colors.HexColor("#0f172a")
        LABELC = colors.HexColor("#475569")
        MUTED = colors.HexColor("#64748b")
        LINEC = colors.HexColor("#e2e8f0")
        SHADOW = colors.HexColor("#e6ebf2")
        CARDBG = colors.HexColor("#f8fafc")
        HEADSUB = colors.HexColor("#c7d6f0")

        # ---------- MEANING colours (each colour carries meaning) ----------
        # Victim = orange, intermediate layers = cream, last / on-hold = red.
        V_BG = colors.HexColor("#F59E0B")
        V_AC = colors.HexColor("#d97706")
        V_TI = colors.HexColor("#7c2d12")
        M_BG = colors.HexColor("#fcfaf9")
        M_AC = colors.HexColor("#94a3b8")
        M_TI = colors.HexColor("#334155")
        H_BG = colors.HexColor("#F86262")
        H_AC = colors.HexColor("#b91c1c")
        H_TI = colors.HexColor("#7f1d1d")

        MARGIN = 50
        HEAD_H = 94
        BOX_W = 360
        x = (width - BOX_W) / 2.0
        PAD = 16
        LINE_H = 15
        LABEL_W = 104

        def _to_f(v):
            try:
                return float(v)
            except Exception:
                return None

        def fmt_money(v):
            n = _to_f(v)
            if n is None:
                return str(v)
            neg = n < 0
            n = abs(n)
            ip, dec = ("%.2f" % n).split(".")
            if len(ip) > 3:
                ip = re.sub(r"(\d)(?=(\d\d)+$)", r"\1,", ip[:-3]) + "," + ip[-3:]
            return ("-" if neg else "") + "Rs. " + ip + "." + dec

        def wrap(text, font, size, maxw):
            out = []
            for raw in str(text).split("\n"):
                cur = ""
                for w in raw.split(" "):
                    t = (cur + " " + w).strip()
                    if not cur or c.stringWidth(t, font, size) <= maxw:
                        cur = t
                    else:
                        out.append(cur)
                        cur = w
                out.append(cur)
            return out or [""]

        def draw_band():
            # official government letterhead: emblem + titles on a deep-navy band
            c.setFillColor(NAVY)
            c.rect(0, height - HEAD_H, width, HEAD_H, fill=1, stroke=0)
            c.setFillColor(GOLD)
            c.rect(0, height - HEAD_H, width, 3, fill=1, stroke=0)
            logo = os.path.join(app.root_path, "static", "tn_police_logo.png")
            try:
                c.drawImage(
                    logo,
                    MARGIN,
                    height - HEAD_H + (HEAD_H - 58) / 2.0,
                    width=58,
                    height=58,
                    preserveAspectRatio=True,
                    mask="auto",
                )
            except Exception:
                pass
            tx = MARGIN + 72
            c.setFillColor(GOLD)
            c.setFont("Helvetica-Bold", 8)
            c.drawString(tx, height - 32, "GOVERNMENT OF TAMIL NADU")
            c.setFillColor(colors.white)
            c.setFont("Helvetica-Bold", 17)
            c.drawString(tx, height - 52, "TAMIL NADU CYBER CRIME POLICE")
            c.setFillColor(HEADSUB)
            c.setFont("Helvetica", 9.5)
            c.drawString(tx, height - 70, "Fund Trail Analysis  -  Put-On-Hold Money Trail")
            c.setFillColor(colors.white)
            c.setFont("Helvetica-Bold", 9)
            c.drawRightString(width - MARGIN, height - 34, "Ack No: %s" % ack_no)
            c.setFillColor(HEADSUB)
            c.setFont("Helvetica", 8.5)
            c.drawRightString(width - MARGIN, height - 50, "Generated " + datetime.now().strftime(DATETIME_DISPLAY_FORMAT))
            c.setFillColor(HEADSUB)
            c.setFont("Helvetica-Bold", 7.5)
            c.drawRightString(width - MARGIN, height - 65, "CONFIDENTIAL")

        def draw_footer():
            c.setStrokeColor(LINEC)
            c.setLineWidth(0.8)
            c.line(MARGIN, 42, width - MARGIN, 42)
            c.setFillColor(MUTED)
            c.setFont("Helvetica", 8)
            c.drawString(MARGIN, 30, "CONFIDENTIAL  -  For authorised investigation use only")
            c.drawRightString(width - MARGIN, 30, "Computer-generated  ·  Page %d" % c.getPageNumber())

        def draw_summary(y_top):
            # summary-of-findings stat row (forensic reports lead with the key numbers)
            accounts = len(nodes)
            disputed = max([_to_f(n.get("disputed_amount")) or 0 for n in nodes] or [0])
            hold = _to_f(nodes[-1].get("hold_amount")) or 0
            cards = [
                ("On-Hold / Frozen", fmt_money(hold), H_AC),
                ("Disputed Amount", fmt_money(disputed), V_AC),
                ("Accounts Traced", str(accounts), NAVY),
            ]
            gap = 12
            ch = 52
            cw = (width - 2 * MARGIN - 2 * gap) / 3.0
            cx = MARGIN
            for label, value, accent in cards:
                c.setFillColor(SHADOW)
                c.roundRect(cx + 1.5, y_top - ch - 1.5, cw, ch, 9, fill=1, stroke=0)
                c.setFillColor(CARDBG)
                c.setStrokeColor(LINEC)
                c.setLineWidth(1)
                c.roundRect(cx, y_top - ch, cw, ch, 9, fill=1, stroke=1)
                c.setFillColor(accent)
                c.roundRect(cx, y_top - ch, 4, ch, 2, fill=1, stroke=0)
                fs = 14.0
                while fs > 8 and c.stringWidth(value, "Helvetica-Bold", fs) > cw - 24:
                    fs -= 0.5
                c.setFillColor(INK)
                c.setFont("Helvetica-Bold", fs)
                c.drawString(cx + 14, y_top - 26, value)
                c.setFillColor(MUTED)
                c.setFont("Helvetica", 7.5)
                c.drawString(cx + 14, y_top - 42, label.upper())
                cx += cw + gap
            return y_top - ch

        def draw_legend(y):
            items = [(V_BG, "Victim"), (M_BG, "Intermediate"), (H_BG, "On-Hold / Frozen")]
            lx = MARGIN
            for col, lab in items:
                c.setFillColor(col)
                c.setStrokeColor(LINEC)
                c.setLineWidth(0.8)
                c.roundRect(lx, y - 8, 11, 11, 2.5, fill=1, stroke=1)
                c.setFillColor(LABELC)
                c.setFont("Helvetica", 8.5)
                c.drawString(lx + 16, y - 6, lab)
                lx += 16 + c.stringWidth(lab, "Helvetica", 8.5) + 22

        def begin_page(first):
            draw_band()
            draw_footer()
            yy = height - HEAD_H - 26
            if first:
                yy = draw_summary(yy)
                draw_legend(yy - 22)
                yy = yy - 46
            return yy

        y = begin_page(True)

        for i, node in enumerate(nodes):
            if i == 0:
                bg, accent, tcolor, badge = V_BG, V_AC, V_TI, "VICTIM ACCOUNT"
            elif i == len(nodes) - 1:
                bg, accent, tcolor, badge = H_BG, H_AC, H_TI, "ON HOLD"
            else:
                bg, accent, tcolor, badge = M_BG, M_AC, M_TI, "LAYER %d" % i

            rows = [("Account No", node.get("account_number", "N/A")), ("Bank", node.get("bank", BANK_UNKNOWN))]
            if i != 0:
                rows += [
                    ("Branch", node.get("branch", "Unknown")),
                    ("IFSC", node.get("ifsc", "N/A")),
                    ("Txn ID", node.get("txn_id", "N/A")),
                    ("Transacted", fmt_money(node.get("amount", "0"))),
                    ("Disputed", fmt_money(node.get("disputed_amount", "0"))),
                ]
            if i == len(nodes) - 1 and node.get("hold_amount"):
                rows += [("On-Hold Amt", fmt_money(node.get("hold_amount")))]

            value_maxw = BOX_W - PAD * 2 - LABEL_W
            disp = []
            for label, value in rows:
                wl = wrap(value, "Helvetica", 10, value_maxw)
                disp.append((label, wl[0]))
                for extra in wl[1:]:
                    disp.append((None, extra))

            box_h = PAD + 20 + 14 + len(disp) * LINE_H + PAD

            if y - box_h < 58:
                c.showPage()
                y = begin_page(False)

            top = y
            # soft drop shadow
            c.setFillColor(SHADOW)
            c.roundRect(x + 2.5, top - box_h - 3.5, BOX_W, box_h, 14, fill=1, stroke=0)
            # card body
            c.setFillColor(bg)
            c.setStrokeColor(LINEC)
            c.setLineWidth(1.2)
            c.roundRect(x, top - box_h, BOX_W, box_h, 14, fill=1, stroke=1)
            # left accent rail
            c.setFillColor(accent)
            c.roundRect(x, top - box_h, 6, box_h, 3, fill=1, stroke=0)

            # header row: pill badge + node index
            by = top - PAD - 10
            pw = c.stringWidth(badge, "Helvetica-Bold", 8.5) + 18
            c.setFillColor(accent)
            c.roundRect(x + PAD + 4, by - 3, pw, 17, 8.5, fill=1, stroke=0)
            c.setFillColor(colors.white)
            c.setFont("Helvetica-Bold", 8.5)
            c.drawString(x + PAD + 13, by + 1.5, badge)
            c.setFillColor(tcolor)
            c.setFont("Helvetica-Bold", 10)
            c.drawRightString(x + BOX_W - PAD, by + 1.5, "#%d" % i)
            # divider
            c.setStrokeColor(accent)
            c.setLineWidth(0.8)
            c.line(x + PAD, by - 12, x + BOX_W - PAD, by - 12)

            ry = by - 28
            for label, value in disp:
                if label is not None:
                    c.setFillColor(LABELC)
                    c.setFont("Helvetica", 9)
                    c.drawString(x + PAD + 6, ry, label.upper())
                c.setFillColor(INK)
                c.setFont("Helvetica", 10)
                c.drawString(x + PAD + 6 + LABEL_W, ry, value)
                ry -= LINE_H

            # connector to next layer, labelled with the amount that moved (fund-flow)
            if i < len(nodes) - 1:
                ax = width / 2.0
                c.setStrokeColor(MUTED)
                c.setLineWidth(1.4)
                c.line(ax, top - box_h, ax, top - box_h - 30)
                amt = fmt_money(nodes[i + 1].get("amount", "0"))
                if amt and amt not in ("Rs. 0.00", "0"):
                    pw = c.stringWidth(amt, "Helvetica-Bold", 8) + 16
                    c.setFillColor(colors.white)
                    c.setStrokeColor(LINEC)
                    c.setLineWidth(0.8)
                    c.roundRect(ax - pw / 2.0, top - box_h - 22, pw, 14, 7, fill=1, stroke=1)
                    c.setFillColor(LABELC)
                    c.setFont("Helvetica-Bold", 8)
                    c.drawCentredString(ax, top - box_h - 18, amt)
                ah = top - box_h - 30
                c.setStrokeColor(MUTED)
                c.line(ax, ah, ax - 4, ah + 6)
                c.line(ax, ah, ax + 4, ah + 6)
                y = top - box_h - 42
            else:
                y = top - box_h - 20

        # ---------- attestation block (government form) ----------
        if y < 110:
            c.showPage()
            y = begin_page(False)
        c.setFillColor(MUTED)
        c.setFont("Helvetica-Oblique", 7.5)
        c.drawString(MARGIN, y - 24, "Verified from FundTrail records as on date of generation.")
        c.setStrokeColor(INK)
        c.setLineWidth(0.8)
        c.line(width - MARGIN - 190, y - 22, width - MARGIN, y - 22)
        c.setFillColor(LABELC)
        c.setFont("Helvetica", 8.5)
        c.drawRightString(width - MARGIN, y - 34, "Signature of Investigating Officer")

        c.save()
        buffer.seek(0)

        filename = f"FundTrail_{ack_no}.pdf"
        log_usage("download_fundtrail_pdf", filename=filename, ack_no=ack_no)
        return send_file(buffer, mimetype=MIME_PDF, as_attachment=True, download_name=filename)

    except Exception as e:
        logger.exception(f"Error generating fundtrail PDF: {e}")
        return jsonify({"error": MSG_DOC_GENERATION_ERROR}), 500


# ─────────────────────────────────────────────────────────────────
#  Officer Analytics — shows only the current officer's own uploads
# ─────────────────────────────────────────────────────────────────
@app.route("/my_analytics", methods=["GET"])
@login_required
def my_analytics():
    """Analytics filtered to the logged-in officer's own uploads only."""
    me = session.get("username")
    try:
        log_usage("my_analytics")
    except Exception:
        pass

    my_files = (
        UploadedFile.query.options(defer(UploadedFile.data))
        .filter_by(uploader=me)
        .order_by(UploadedFile.upload_time.desc())
        .all()
    )
    my_file_ids = [f.id for f in my_files]

    # All my transactions
    my_txns = Transaction.query.filter(Transaction.upload_id.in_(my_file_ids)).all() if my_file_ids else []

    # KPIs
    total_files = len(my_files)
    total_txns = len(my_txns)
    total_amount = sum(t.amount or 0 for t in my_txns)
    disputed_amt = sum(t.disputed_amount or 0 for t in my_txns)
    hold_amt = sum(t.put_on_hold_amount or 0 for t in my_txns if t.put_on_hold_txn_id)
    hold_count = sum(1 for t in my_txns if t.put_on_hold_txn_id)
    unique_accts = len({t.to_account for t in my_txns if t.to_account})
    unique_banks = len({t.bank_name for t in my_txns if t.bank_name})
    max_layer = max((t.layer for t in my_txns if t.layer), default=0)
    atm_count = sum(1 for t in my_txns if t.atm_id)
    cheque_count = sum(1 for t in my_txns if t.cheque_no)

    # Per-file breakdown
    file_rows = []
    for f in my_files:
        ftxns = [t for t in my_txns if t.upload_id == f.id]
        ack_nos = list({t.ack_no for t in ftxns if t.ack_no})
        fa = sum(t.amount or 0 for t in ftxns)
        fd = sum(t.disputed_amount or 0 for t in ftxns)
        fh = sum(t.put_on_hold_amount or 0 for t in ftxns if t.put_on_hold_txn_id)
        fhc = sum(1 for t in ftxns if t.put_on_hold_txn_id)
        file_rows.append(
            {
                "id": f.id,
                "filename": f.filename,
                "upload_time": (to_ist(f.upload_time).strftime(DATETIME_DISPLAY_FORMAT) if f.upload_time else "—"),
                "txn_count": len(ftxns),
                "ack_nos": ack_nos,
                "total_amt": fa,
                "disputed_amt": fd,
                "hold_count": fhc,
                "hold_amt": fh,
            }
        )

    # Layer-wise
    layer_map = defaultdict(lambda: {"txns": 0, "amt": 0.0})
    for t in my_txns:
        if t.layer:
            layer_map[t.layer]["txns"] += 1
            layer_map[t.layer]["amt"] += t.amount or 0
    layer_dist = [(lyr, v["txns"], v["amt"]) for lyr, v in sorted(layer_map.items())]

    # Top banks
    bank_map = defaultdict(lambda: {"cnt": 0, "amt": 0.0})
    for t in my_txns:
        if t.bank_name:
            bank_map[t.bank_name]["cnt"] += 1
            bank_map[t.bank_name]["amt"] += t.amount or 0
    top_banks = sorted(bank_map.items(), key=lambda x: -x[1]["cnt"])[:10]
    top_banks = [(b, v["cnt"], v["amt"]) for b, v in top_banks]

    # State-wise
    state_map = defaultdict(lambda: {"cnt": 0, "amt": 0.0})
    for t in my_txns:
        if t.state:
            state_map[t.state]["cnt"] += 1
            state_map[t.state]["amt"] += t.amount or 0
    state_dist = sorted(state_map.items(), key=lambda x: -x[1]["cnt"])[:15]
    state_dist = [(s, v["cnt"], v["amt"]) for s, v in state_dist]

    # Top IFSC — skip junk values (nan/none/unknown/blank) so the table isn't polluted
    ifsc_map = defaultdict(lambda: {"cnt": 0, "bank": ""})
    _JUNK = {"", "nan", "none", "null", "unknown", "n/a", "na", "-", "—"}
    for t in my_txns:
        _ifsc = (t.ifsc_code or "").strip()
        if _ifsc and _ifsc.lower() not in _JUNK:
            ifsc_map[_ifsc]["cnt"] += 1
            _bank = (t.bank_name or "").strip()
            if _bank and _bank.lower() not in _JUNK:
                ifsc_map[_ifsc]["bank"] = _bank
    top_ifsc = sorted(ifsc_map.items(), key=lambda x: -x[1]["cnt"])[:10]
    top_ifsc = [(i, v["bank"], v["cnt"]) for i, v in top_ifsc]

    return render_template(
        "my_analytics.html",
        me=me,
        total_files=total_files,
        total_txns=total_txns,
        total_amount=total_amount,
        disputed_amt=disputed_amt,
        hold_amt=hold_amt,
        hold_count=hold_count,
        unique_accts=unique_accts,
        unique_banks=unique_banks,
        max_layer=max_layer,
        atm_count=atm_count,
        cheque_count=cheque_count,
        file_rows=file_rows,
        layer_dist=layer_dist,
        top_banks=top_banks,
        top_ifsc=top_ifsc,
        state_dist=state_dist,
    )


# ═════════════════════════════════════════════════════════════════
#  Enterprise feature routes
# ═════════════════════════════════════════════════════════════════

APP_VERSION = "2.0.0"


def _accessible_ack_nos():
    """Set of ACK numbers the current user may see (None = unrestricted admin)."""
    if is_admin():
        return None
    rows = (
        Complaint.query.filter(or_(Complaint.uploaded_by == current_user.id, Complaint.assigned_to == current_user.id))
        .with_entities(Complaint.ack_no)
        .all()
    )
    return {r[0] for r in rows}


def _integration_test_routes_enabled():
    return os.environ.get("ENABLE_INTEGRATION_TEST_ROUTES", "").lower() == "true"


def _require_integration_test_token():
    expected = os.environ.get("INTEGRATION_TEST_TOKEN", "")
    provided = request.headers.get("Authorization", "").removeprefix("Bearer ").strip()
    if not expected or not secrets.compare_digest(provided, expected):
        abort(404)


# ---------------------------------------------------------------------------
# 13. Routes: search, case workflow & health
# ---------------------------------------------------------------------------


@app.route("/healthz", methods=["GET"])
def healthz():
    """Unauthenticated liveness/readiness probe for monitors."""
    db_ok = True
    try:
        db.session.execute(text("SELECT 1"))
    except Exception:
        db_ok = False
    status = 200 if db_ok else 503
    return jsonify(
        {"status": "ok" if db_ok else "degraded", "db": "ok" if db_ok else "error", "version": APP_VERSION}
    ), status


@app.route("/__integration_test/sentry", methods=["POST"])
@csrf.exempt
def integration_test_sentry():
    """Disabled-by-default probe that intentionally raises for Sentry validation."""
    if not _integration_test_routes_enabled():
        abort(404)
    _require_integration_test_token()
    raise RuntimeError("FundTrail Sentry integration test exception")


@app.route("/__integration_test/resend", methods=["POST"])
@csrf.exempt
def integration_test_resend():
    """Disabled-by-default probe that sends one test email through Resend."""
    if not _integration_test_routes_enabled():
        abort(404)
    _require_integration_test_token()
    to_email = os.environ.get("RESEND_TEST_TO", "").strip()
    if not to_email:
        return jsonify({"status": "skipped", "reason": "RESEND_TEST_TO is not configured"}), 400
    result = send_integration_test_email(to_email)
    if result is None:
        return jsonify({"status": "skipped", "reason": "RESEND_API_KEY or resend package is not configured"}), 400
    return jsonify({"status": "sent", "to": to_email})


@app.route("/search", methods=["GET"])
@login_required
def global_search():
    """Find any ACK / account / transaction ID / bank across the
    cases the current user is allowed to see."""
    q = sanitize_text(request.args.get("q", ""), 80)
    results = {"cases": [], "transactions": []}
    if len(q) >= 3:
        allowed = _accessible_ack_nos()
        like = f"%{q}%"
        txn_q = Transaction.query.filter(
            or_(
                Transaction.ack_no.ilike(like),
                Transaction.to_account.ilike(like),
                Transaction.from_account.ilike(like),
                Transaction.txn_id.ilike(like),
                Transaction.bank_name.ilike(like),
            )
        )
        if allowed is not None:
            txn_q = txn_q.filter(Transaction.ack_no.in_(allowed))
        txns = txn_q.limit(100).all()
        results["transactions"] = txns
        case_q = Complaint.query.filter(Complaint.ack_no.ilike(like))
        if allowed is not None:
            case_q = case_q.filter(Complaint.ack_no.in_(allowed))
        results["cases"] = case_q.limit(50).all()
        log_usage("global_search", filename=q)
    return render_template("search_results.html", q=q, results=results)


@app.route("/repeat_accounts", methods=["GET"])
@login_required
@admin_required
def repeat_accounts():
    """Mule-account detection — beneficiary accounts that appear in
    more than one distinct case."""
    rows = (
        db.session.query(
            Transaction.to_account,
            func.count(func.distinct(Transaction.ack_no)).label("case_count"),
            func.count(Transaction.id).label("txn_count"),
            func.sum(Transaction.amount).label("total_amount"),
            func.max(Transaction.bank_name).label("bank_name"),
        )
        .filter(Transaction.to_account.isnot(None), Transaction.to_account != "")
        .group_by(Transaction.to_account)
        .having(func.count(func.distinct(Transaction.ack_no)) > 1)
        .order_by(desc("case_count"), desc("total_amount"))
        .limit(200)
        .all()
    )
    # Attach the case list per flagged account (bounded by the 200-row cap above).
    accounts = []
    for r in rows:
        acks = [
            a[0]
            for a in (
                db.session.query(func.distinct(Transaction.ack_no)).filter(Transaction.to_account == r.to_account).all()
            )
        ]
        accounts.append(
            {
                "account": r.to_account,
                "case_count": r.case_count,
                "txn_count": r.txn_count,
                "total_amount": r.total_amount or 0,
                "bank": r.bank_name,
                "ack_nos": acks,
            }
        )
    return render_template("repeat_accounts.html", accounts=accounts)


@app.route("/update_case_status", methods=["POST"])
@login_required
def update_case_status():
    """Move a case through its workflow (Open / Under Investigation /
    Closed). Allowed for the admin and the case's own officer."""
    ack_no = sanitize_text(request.form.get("ack_no", ""), 100)
    status = request.form.get("status", "").strip()
    if status not in CASE_STATUSES:
        flash("Invalid case status.", "danger")
        return redirect(request.referrer or url_for("role_home"))
    check_case_access(ack_no)
    complaint = Complaint.query.filter_by(ack_no=ack_no).first()
    if not complaint:
        flash(f"No case found for ACK {ack_no}.", "warning")
        return redirect(request.referrer or url_for("role_home"))
    complaint.status = status
    db.session.commit()
    log_usage("update_case_status", ack_no=ack_no, filename=status)
    flash(f'Case {ack_no} marked "{status}".', "success")
    return redirect(request.referrer or url_for("role_home"))


@app.route("/case_notes/<ack_no>", methods=["POST"])
@login_required
def add_case_note(ack_no):
    """Append a dated investigation note to a case."""
    check_case_access(ack_no)
    note = sanitize_text(request.form.get("note", ""), 2000)
    if not note:
        flash("Note cannot be empty.", "warning")
        return redirect(url_for("case_timeline", ack_no=ack_no))
    db.session.add(CaseNote(ack_no=str(ack_no).strip(), author=current_user.username, note=note))
    db.session.commit()
    log_usage("add_case_note", ack_no=ack_no)
    flash("Note added.", "success")
    return redirect(url_for("case_timeline", ack_no=ack_no))


@app.route("/case_timeline/<ack_no>", methods=["GET"])
@login_required
def case_timeline(ack_no):
    """Chronological case diary — notes + recorded actions."""
    check_case_access(ack_no)
    ack = str(ack_no).strip()
    complaint = Complaint.query.filter_by(ack_no=ack).first()
    notes = CaseNote.query.filter_by(ack_no=ack).order_by(CaseNote.created_at.desc()).all()
    events = UsageLog.query.filter_by(ack_no=ack).order_by(UsageLog.timestamp.desc()).limit(200).all()
    return render_template(
        "case_timeline.html", ack_no=ack, complaint=complaint, notes=notes, events=events, statuses=CASE_STATUSES
    )


@app.route("/download_letters_zip/<ack_no>", methods=["GET"])
@login_required
def download_letters_zip(ack_no):
    """Download every generated letter for a case as one ZIP."""
    check_case_access(ack_no)
    ack = secure_filename(str(ack_no).strip())
    base_dir = os.path.join(app.root_path, "generated_letters", ack)
    if not os.path.isdir(base_dir):
        flash("No generated letters found for this case yet.", "warning")
        return redirect(request.referrer or url_for("role_home"))
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for root, _dirs, files in os.walk(base_dir):
            for fname in files:
                full = os.path.join(root, fname)
                zf.write(full, os.path.relpath(full, base_dir))
    buf.seek(0)
    log_usage("download_letters_zip", ack_no=ack)
    return send_file(buf, mimetype="application/zip", as_attachment=True, download_name=f"Letters_{ack}.zip")


def _analytics_export_frames(uploader=None):
    """Build the DataFrames for the Excel export. When `uploader`
    is set, restrict to that officer's uploads (officer self-export)."""
    txn_q = db.session.query(
        Transaction.ack_no,
        Transaction.layer,
        Transaction.from_account,
        Transaction.to_account,
        Transaction.bank_name,
        Transaction.ifsc_code,
        Transaction.state,
        Transaction.amount,
        Transaction.disputed_amount,
        Transaction.put_on_hold_amount,
        Transaction.refund_status,
        Transaction.refund_amount,
        Transaction.txn_date,
    )
    if uploader:
        ids = [f.id for f in UploadedFile.query.with_entities(UploadedFile.id).filter_by(uploader=uploader)]
        txn_q = txn_q.filter(Transaction.upload_id.in_(ids or [-1]))
    txns = pd.DataFrame(txn_q.all())
    summary = (
        (
            txns.groupby("ack_no")
            .agg(
                transactions=("ack_no", "size"),
                total_amount=("amount", "sum"),
                held=("put_on_hold_amount", "sum"),
                refunded=("refund_amount", "sum"),
            )
            .reset_index()
        )
        if not txns.empty
        else pd.DataFrame()
    )
    return summary, txns


def _send_analytics_xlsx(uploader=None, label="analytics"):
    summary, txns = _analytics_export_frames(uploader)
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as xw:
        (summary if not summary.empty else pd.DataFrame({"info": ["no data"]})).to_excel(
            xw, sheet_name="Case Summary", index=False
        )
        (txns if not txns.empty else pd.DataFrame({"info": ["no data"]})).to_excel(
            xw, sheet_name="Transactions", index=False
        )
    buf.seek(0)
    log_usage("export_analytics_xlsx", filename=label)
    stamp = datetime.now().strftime("%Y%m%d_%H%M")
    return send_file(
        buf,
        as_attachment=True,
        download_name=f"FundTrail_{label}_{stamp}.xlsx",
        mimetype=MIME_XLSX,
    )


@app.route("/export_analytics_xlsx", methods=["GET"])
@login_required
@admin_required
def export_analytics_xlsx():
    """Admin-wide analytics export to Excel."""
    return _send_analytics_xlsx(label="all_cases")


@app.route("/export_my_analytics_xlsx", methods=["GET"])
@login_required
def export_my_analytics_xlsx():
    """Officer's own analytics export to Excel."""
    return _send_analytics_xlsx(uploader=session.get("username"), label="my_cases")


@app.route("/refund_dashboard", methods=["GET"])
@login_required
def refund_dashboard():
    """Per-case put-on-hold vs refunded amounts — the unit's
    headline recovery metric in one view."""
    allowed = _accessible_ack_nos()
    q = (
        db.session.query(
            Transaction.ack_no,
            func.sum(Transaction.put_on_hold_amount).label("held"),
            func.sum(Transaction.refund_amount).label("refunded"),
            func.count(Transaction.id).label("txns"),
        )
        .filter(Transaction.ack_no.isnot(None))
        .group_by(Transaction.ack_no)
    )
    if allowed is not None:
        q = q.filter(Transaction.ack_no.in_(allowed))
    rows = q.all()
    cases = []
    total_held = total_refunded = 0.0
    status_map = {c.ack_no: c.status for c in Complaint.query.all()}
    for r in rows:
        held = r.held or 0.0
        refunded = r.refunded or 0.0
        total_held += held
        total_refunded += refunded
        cases.append(
            {
                "ack_no": r.ack_no,
                "held": held,
                "refunded": refunded,
                "txns": r.txns,
                "status": status_map.get(r.ack_no, "Open"),
                "pct": (refunded / held * 100) if held else 0.0,
            }
        )
    cases.sort(key=lambda c: -c["held"])
    recovery_pct = (total_refunded / total_held * 100) if total_held else 0.0
    return render_template(
        "refund_dashboard.html",
        cases=cases,
        total_held=total_held,
        total_refunded=total_refunded,
        recovery_pct=recovery_pct,
    )


AUDIT_PAGE_SIZE = 50


@app.route("/audit_logs", methods=["GET"])
@login_required
@admin_required
def audit_logs():
    """Searchable, paginated viewer over the UsageLog audit trail."""
    page = max(request.args.get("page", 1, type=int), 1)
    f_user = sanitize_text(request.args.get("user", ""), 100)
    f_action = sanitize_text(request.args.get("action", ""), 100)
    f_ack = sanitize_text(request.args.get("ack", ""), 100)
    # Date/time range filter — values come from <input type="datetime-local"> ("from"/"to").
    f_from = sanitize_text(request.args.get("from", ""), 40)
    f_to = sanitize_text(request.args.get("to", ""), 40)

    def _parse_audit_dt(s):
        for fmt in ("%Y-%m-%dT%H:%M", "%Y-%m-%dT%H:%M:%S", "%Y-%m-%d %H:%M", "%Y-%m-%d"):
            try:
                return datetime.strptime(s, fmt)
            except ValueError:
                continue
        return None

    q = UsageLog.query
    if f_user:
        q = q.filter(UsageLog.username.ilike(f"%{f_user}%"))
    if f_action:
        q = q.filter(UsageLog.action.ilike(f"%{f_action}%"))
    if f_ack:
        q = q.filter(UsageLog.ack_no.ilike(f"%{f_ack}%"))
    _dt_from = _parse_audit_dt(f_from) if f_from else None
    _dt_to = _parse_audit_dt(f_to) if f_to else None
    if _dt_from:
        q = q.filter(UsageLog.timestamp >= _dt_from)
    if _dt_to:
        q = q.filter(UsageLog.timestamp <= _dt_to)
    total = q.count()
    logs = q.order_by(UsageLog.timestamp.desc()).offset((page - 1) * AUDIT_PAGE_SIZE).limit(AUDIT_PAGE_SIZE).all()
    pages = max((total + AUDIT_PAGE_SIZE - 1) // AUDIT_PAGE_SIZE, 1)
    actions = [a[0] for a in db.session.query(func.distinct(UsageLog.action)).all() if a[0]]
    return render_template(
        "audit_logs.html",
        logs=logs,
        page=page,
        pages=pages,
        total=total,
        f_user=f_user,
        f_action=f_action,
        f_ack=f_ack,
        f_from=f_from,
        f_to=f_to,
        actions=sorted(actions),
    )


@app.route("/admin_metrics", methods=["GET"])
@login_required
@admin_required
def admin_metrics():
    """Management metrics — uploads/week, case status mix,
    officer activity, and the hold-vs-refund recovery rate."""
    now = datetime.now(timezone.utc)
    weeks = []
    for i in range(11, -1, -1):
        start = now - timedelta(weeks=i + 1)
        end = now - timedelta(weeks=i)
        count = UploadedFile.query.filter(UploadedFile.upload_time >= start, UploadedFile.upload_time < end).count()
        weeks.append({"label": end.strftime("%d %b"), "count": count})

    status_counts = dict(db.session.query(Complaint.status, func.count(Complaint.id)).group_by(Complaint.status).all())
    held = db.session.query(func.sum(Transaction.put_on_hold_amount)).scalar() or 0.0
    refunded = db.session.query(func.sum(Transaction.refund_amount)).scalar() or 0.0
    officer_rows = (
        db.session.query(UploadedFile.uploader, func.count(UploadedFile.id)).group_by(UploadedFile.uploader).all()
    )
    return render_template(
        "admin_metrics.html",
        weeks=weeks,
        status_counts=status_counts,
        total_cases=Complaint.query.count(),
        total_officers=User.query.filter_by(role=ROLE_INVESTIGATIVE_OFFICER).count(),
        held=held,
        refunded=refunded,
        recovery_pct=(refunded / held * 100) if held else 0.0,
        officer_rows=officer_rows,
        app_version=APP_VERSION,
    )


@app.route("/help", methods=["GET"])
@login_required
def help_page():
    """Built-in workflow guide for new officers."""
    return render_template("help.html")


# ---------------------------------------------------------------------------
# 14. Entrypoint
# ---------------------------------------------------------------------------


if __name__ == "__main__":
    with app.app_context():
        db.create_all()
        ensure_transaction_columns()
        ensure_user_columns()
        ensure_complaint_columns()
        ensure_mrm_backfill()

    # Use environment variable to control debug mode
    debug_mode = os.environ.get("FLASK_DEBUG", "False").lower() == "true"
    # Port is configurable via the PORT env var so it can avoid conflicts — e.g.
    # macOS AirPlay Receiver (Control Center) occupies 5000 by default. Default 5000.
    port = int(os.environ.get("PORT", "5050"))
    app.run(debug=debug_mode, host="127.0.0.1", port=port)

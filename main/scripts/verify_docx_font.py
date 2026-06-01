from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def test_font_change():
    # Create a dummy document
    doc = Document()
    p = doc.add_paragraph("I am the investigating officer of the case mentioned in the subject. Old text.")
    
    # Simulate the logic
    target_text_start = "I am the investigating officer"
    formatted_l1_total = "₹23,00,000"
    new_paragraph_text = f"I am the investigating officer of the case mentioned in the subject. In this case, the victim has lost {formatted_l1_total}. Of the amount lost, partial amounts has been sent to the account mentioned below. Request you to provide the required details to proceed with the further investigation."
    
    found = False
    for p in doc.paragraphs:
        if p.text.strip().startswith(target_text_start):
            p.clear() # Remove existing runs
            run = p.add_run(new_paragraph_text)
            run.font.name = 'Bookman Old Style'
            run.font.size = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            found = True
            break
            
    if not found:
        print("Paragraph not found!")
        return

    # Verify
    p = doc.paragraphs[0]
    print(f"Text: {p.text}")
    print(f"Font Name: {p.runs[0].font.name}")
    print(f"Font Size: {p.runs[0].font.size.pt}")
    print(f"Alignment: {p.alignment}")
    
    assert p.text == new_paragraph_text
    assert p.runs[0].font.name == 'Bookman Old Style'
    assert p.runs[0].font.size.pt == 12.0
    assert p.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
    print("Test Passed!")

if __name__ == "__main__":
    test_font_change()


from docx import Document
import os

def inspect_template():
    template_name = 'Template for letter generation_suspect accounts.docx'
    if not os.path.exists(template_name):
        print(f"Template {template_name} not found.")
        return

    doc = Document(template_name)
    
    print("Paragraphs:")
    for p in doc.paragraphs:
        print(f"  - {p.text}")
            
    print("\nTables:")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    print(f"  - {p.text}")

if __name__ == "__main__":
    inspect_template()


import docx
import re

def inspect_docx(path):
    try:
        doc = docx.Document(path)
        print(f"--- Inspecting: {path} ---")
        
        placeholders = set()
        text_content = []

        for para in doc.paragraphs:
            text = para.text
            text_content.append(text)
            # Look for typical placeholders like <Name>, {{Name}}, etc.
            # Adjust regex based on what we see, but generally catching anything in brackets
            matches = re.findall(r'<[^>]+>', text)
            placeholders.update(matches)
        
        print(f"Found {len(placeholders)} unique placeholders: {placeholders}")
        print("\n--- First 10 paragraphs ---")
        for i, t in enumerate(text_content[:10]):
            print(f"{i}: {t}")
            
        print("\n--- Tables ---")
        for i, table in enumerate(doc.tables):
            print(f"Table {i} has {len(table.rows)} rows and {len(table.columns)} columns.")
            if len(table.rows) > 0:
                print(f"  Header row: {[cell.text for cell in table.rows[0].cells]}")

    except Exception as e:
        print(f"Error reading docx: {e}")

if __name__ == "__main__":
    inspect_docx(r"c:\Users\ADMIN\Downloads\FUND-TRAIL-TOOL-UPDATION-main\FUND-TRAIL-TOOL-UPDATION-main\Template for letter generation_victim account.docx")

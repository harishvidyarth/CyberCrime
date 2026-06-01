from flask import Blueprint, request, jsonify, render_template, send_file, current_app
from flask_login import login_required, current_user
from sqlalchemy import or_
from datetime import datetime
import io
import os
import zipfile
import logging
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from xhtml2pdf import pisa
from werkzeug.utils import secure_filename

from ..extensions import db
from ..models import Transaction
from ..utils.data_processing import format_indian_currency, ordinal
from ..utils.file_handling import resource_path
from ..utils.logging import log_usage

logger = logging.getLogger(__name__)

letters_bp = Blueprint('letters', __name__)

def get_layer_1_total(ack_no):
    try:
        l1_txns = Transaction.query.filter_by(ack_no=ack_no, layer=1).all()
        return sum(t.amount for t in l1_txns if t.amount) if l1_txns else 0.0
    except Exception as e:
        logger.error(f"Error calculating layer 1 total: {e}")
        return 0.0

@letters_bp.route('/generate_letter', methods=['POST'])
@login_required
def generate_letter():
    try:
        data = request.get_json() or {}
        ack_no = data.get('ack_no')
        account_number = data.get('account_number')
        letter_type = data.get('letter_type')

        context = {
            'ack_no': ack_no,
            'account_number': account_number,
            'officer_name': data.get('officer_name', ''),
            'officer_designation': data.get('officer_designation', ''),
            'officer_phone': data.get('officer_phone', ''),
            'officer_email': data.get('officer_email', ''),
            'letter_date': data.get('letter_date', datetime.now().strftime('%d-%m-%Y')),
            'crime_no': data.get('crime_no', ''),
            'ncrp_ack_no': data.get('ncrp_ack_no', ack_no or ''),
            'layer_label': 'Unknown layer',
            'layer_1_total': format_indian_currency(get_layer_1_total(ack_no))
        }

        template_name = 'letter_victim.html' if letter_type == 'victim' else 'letter_suspect.html'
        log_usage('generate_letter', filename=template_name, ack_no=ack_no)

        return render_template(template_name, **context)
    except Exception as e:
        logger.error(f"Error generating letter: {e}")
        return jsonify({'error': str(e)}), 500

@letters_bp.route('/generate_letter_pdf', methods=['POST'])
@login_required
def generate_letter_pdf():
    try:
        data = request.get_json() or {}
        ack_no = data.get('ack_no')
        account_number = data.get('account_number')
        letter_type = data.get('letter_type')

        context = {
            'ack_no': ack_no,
            'account_number': account_number,
            'officer_name': data.get('officer_name', ''),
            'officer_designation': data.get('officer_designation', ''),
            'officer_phone': data.get('officer_phone', ''),
            'officer_email': data.get('officer_email', ''),
            'letter_date': data.get('letter_date', datetime.now().strftime('%d-%m-%Y')),
            'crime_no': data.get('crime_no', ''),
            'ncrp_ack_no': data.get('ncrp_ack_no', ack_no or ''),
            'layer_1_total': format_indian_currency(get_layer_1_total(ack_no))
        }

        template_name = 'letter_victim.html' if letter_type == 'victim' else 'letter_suspect.html'

        try:
            t = Transaction.query.filter(Transaction.ack_no == ack_no, or_(Transaction.to_account == account_number, Transaction.account_number == account_number)).first()
            if t and getattr(t, 'layer', None):
                context['layer_label'] = f"{ordinal(t.layer)} layer"
        except Exception:
            pass

        html_content = render_template(template_name, **context)
        pdf_buffer = io.BytesIO()
        pisa_status = pisa.CreatePDF(io.BytesIO(html_content.encode('utf-8')), dest=pdf_buffer)

        if pisa_status.err: return jsonify({'error': 'PDF generation failed'}), 500

        folder_name = secure_filename(context['ncrp_ack_no'] or 'Unknown_ACK')
        base_dir = os.path.join(current_app.root_path, 'generated_letters')
        target_dir = os.path.join(base_dir, folder_name)
        os.makedirs(target_dir, exist_ok=True)
        
        safe_filename_val = secure_filename(f"{letter_type}_Letter_{account_number}.pdf")
        with open(os.path.join(target_dir, safe_filename_val), 'wb') as f:
            f.write(pdf_buffer.getvalue())

        pdf_buffer.seek(0)
        return send_file(pdf_buffer, mimetype='application/pdf', as_attachment=True, download_name=safe_filename_val)

    except Exception as e:
        logger.error(f"Error generating PDF letter: {e}")
        return jsonify({'error': str(e)}), 500

@letters_bp.route('/generate_letter_docx', methods=['POST'])
@login_required
def generate_letter_docx():
    try:
        data = request.get_json() or {}
        ack_no = data.get('ack_no')
        account_numbers = data.get('account_numbers', [])
        if not account_numbers and data.get('account_number'): account_numbers = [data.get('account_number')]
        if not account_numbers: return jsonify({'error': 'No account numbers provided'}), 400

        letter_type = data.get('letter_type')
        is_poh = data.get('is_poh', False)

        context = {
            'ack_no': ack_no,
            'officer_name': data.get('officer_name', '<Name>'),
            'officer_designation': data.get('officer_designation', '<Rank>'),
            'officer_phone': data.get('officer_phone', '<Phone>'),
            'officer_email': data.get('officer_email', '<email id>'),
            'letter_date': data.get('letter_date', datetime.now().strftime('%d-%m-%Y')),
            'crime_no': data.get('crime_no', '<Cr.No>'),
            'ncrp_ack_no': data.get('ncrp_ack_no', ack_no or '<Acknowledgement no>')
        }
        
        formatted_l1_total = format_indian_currency(get_layer_1_total(ack_no))

        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
            for account_number in account_numbers:
                if is_poh:
                    transactions = Transaction.query.filter(Transaction.ack_no == ack_no, Transaction.put_on_hold_txn_id != None, Transaction.put_on_hold_txn_id != '', or_(Transaction.account_number == account_number, Transaction.to_account == account_number)).all()
                else:
                    transactions = Transaction.query.filter_by(ack_no=ack_no, to_account=account_number).all()

                bank_name = "Unknown Bank"; amount_lost = 0.0; from_date = context['letter_date']
                if transactions:
                    bank_name = next((t.bank_name for t in transactions if t.bank_name), next((t.ifsc_code for t in transactions if t.ifsc_code), "Unknown Bank"))
                    if is_poh:
                        amount_lost = sum(t.put_on_hold_amount for t in transactions if t.put_on_hold_amount)
                        dates = [t.put_on_hold_date for t in transactions if t.put_on_hold_date]
                    else:
                        amount_lost = sum(t.amount for t in transactions if t.amount)
                        dates = [t.txn_date for t in transactions if t.txn_date]
                    if dates: from_date = min(dates)

                override_amt = data.get('disputed_amount')
                if override_amt:
                    try:
                        amount_lost = float(override_amt)
                    except ValueError:
                        pass

                layer_label = "Unknown layer"; date_minus_6_str = from_date
                if transactions:
                    for t in transactions:
                        if getattr(t, 'layer', None): layer_label = f"{ordinal(t.layer)} layer"; break
                    try:
                        date_minus_6_str = (pd.to_datetime(from_date, dayfirst=True) - timedelta(days=180)).strftime('%d-%m-%Y')
                    except: pass

                template_name = 'Template for letter generation_victim account.docx' if letter_type == 'victim' else 'Template for letter generation_suspect accounts.docx'
                template_path = resource_path(template_name)
                doc = Document(template_path) if os.path.exists(template_path) else Document()

                replacements = {
                    '<Bank Name>': bank_name, '<amount lost>': f"{amount_lost:,.2f}", '<total amount>': f"{amount_lost:,.2f}", '(corresponding layer)': layer_label,
                    '<layer>': layer_label, '<add all layer 1 amount>': formatted_l1_total, 'over all 1st layer total amount': formatted_l1_total, 'overall 1st layer total amount': formatted_l1_total,
                    '< Layer >': layer_label, '<Layer>': layer_label, '< layer >': layer_label, '<transaction date - 6 months>': date_minus_6_str, '<from date>': from_date,
                    '<Rank>': context['officer_designation'], '<Name>': context['officer_name'], '<Phno>': context['officer_phone'], '<email id>': context['officer_email'],
                    '<Cr.No>': context['crime_no'], '<Acknowledgement no>': context['ncrp_ack_no'], '<Account No>': account_number, '<To Date>': context['letter_date'],
                    '<current date>': context['letter_date'], '<Suspect Account Number>': str(account_number), '<Suspect Account>': str(account_number),
                }

                def replace_in_p(p):
                    for k, v in replacements.items():
                        if k in p.text:
                            for run in p.runs:
                                if k in run.text: run.text = run.text.replace(k, str(v))
                            if k in p.text: p.text = p.text.replace(k, str(v))

                for p in doc.paragraphs: replace_in_p(p)
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs: replace_in_p(p)

                target_start = "I am the investigating officer"
                new_p_text = f"I am the investigating officer of the case mentioned in the subject. In this case, the victim has lost {formatted_l1_total}. Of the amount lost, partial amounts has been sent to the account mentioned below. Request you to provide the required details to proceed with the further investigation."
                for p in doc.paragraphs:
                    if p.text.strip().startswith(target_start):
                        p.clear(); r = p.add_run(new_p_text); r.font.name = 'Bookman Old Style'; r.font.size = Pt(12); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; break

                for p in doc.paragraphs:
                    if "Suspect Account Txn Details" in p.text: p.text = f"({layer_label}) Suspect Account Txn Details"; break
                
                target_table = next((t for t in doc.tables if len(t.rows) > 0 and any(h in [c.text.strip() for c in t.rows[0].cells] for h in ["Suspect Account Number", "Transaction Id / UTR Number", "Victim Account Number"])), None)
                if target_table:
                    for i in range(len(target_table.rows) - 1, 0, -1):
                        tr = target_table.rows[i]._element; tr.getparent().remove(tr)
                else:
                    target_table = doc.add_table(rows=1, cols=6); target_table.style = 'Table Grid'
                    hdr = target_table.rows[0].cells; hs = ["S. No.", "Suspect Account Number", "Transaction Date", "Transaction Amount", "Transaction Id / UTR Number", "IFSC Code"]
                    for i, h in enumerate(hs):
                        hdr[i].text = h
                        for p in hdr[i].paragraphs:
                            for r in p.runs: r.bold = True
                
                for idx, t in enumerate(transactions, 1):
                    row = target_table.add_row().cells
                    row[0].text = str(idx)
                    if is_poh:
                        row[1].text = str(t.account_number or t.to_account or ''); row[2].text = str(t.put_on_hold_date or ''); row[3].text = str(t.put_on_hold_amount or ''); row[4].text = str(t.put_on_hold_txn_id or ''); row[5].text = str(t.ifsc_code or '')
                    else:
                        row[1].text = str(t.to_account or ''); row[2].text = str(t.txn_date or ''); row[3].text = str(t.amount or ''); row[4].text = str(t.txn_id or ''); row[5].text = str(t.ifsc_code or '')

                subf = 'suspect letter' if (is_poh or letter_type == 'suspect') else 'victim letter'
                safe_fn = secure_filename(f"{('Suspect_Account_Letter' if is_poh else letter_type + '_Letter')}_{account_number}.docx")
                
                target_dir = os.path.join(current_app.root_path, 'generated_letters', secure_filename(context['ncrp_ack_no'] or 'Unknown_ACK'), subf)
                os.makedirs(target_dir, exist_ok=True)
                doc.save(os.path.join(target_dir, safe_fn))
                zf.write(os.path.join(target_dir, safe_fn), arcname=os.path.join(subf, safe_fn))
        
        zip_buffer.seek(0)
        return send_file(zip_buffer, mimetype='application/zip', as_attachment=True, download_name=f"Letters_{ack_no}.zip")

    except Exception as e:
        logger.error(f"Error generating DOCX letter: {e}")
        return jsonify({'error': str(e)}), 500
